"""``sp.paper(data, question)`` — end-to-end "data → publication draft" pipeline.

The ``causal_workflow.CausalWorkflow`` object already strings together
``diagnose → recommend → estimate → robustness``. This module is the
agent-native top layer: take a DataFrame plus a *natural-language
question* and emit a near-publishable draft (markdown / LaTeX / Word) in
a single call.

Pipeline:

    1. Parse ``question`` (lightweight regex / token heuristic) to fill
       in any missing y/treatment/instrument/cutoff hints. Explicit
       arguments always win.
    2. Run :class:`CausalWorkflow` (diagnose → recommend → estimate →
       robustness).
    3. Build a structured EDA section from :func:`sp.sumstats` (or a
       graceful inline fallback).
    4. Render to the chosen format. Markdown leverages
       :meth:`CausalWorkflow._render_markdown`; LaTeX wraps the same
       content with ``\\section{}`` boilerplate; Word delegates to the
       built-in ``to_docx`` path on the result object.

Notes on design
---------------
- This is **orchestration only**. No numerical primitives are
  re-implemented here.
- The question parser is intentionally simple — it provides hints, never
  overrides the user's explicit kwargs. Agents can also fully bypass the
  parser by passing all relevant column args directly.
- LLM calls are not made by default. The pipeline only triggers an LLM
  oracle if the user passes ``dag=`` from a prior
  :func:`sp.llm_dag_constrained` run.
"""
from __future__ import annotations

import datetime as _dt
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd

from ..output._lineage import format_provenance, get_provenance


__all__ = ["paper", "PaperDraft", "parse_question"]


# --------------------------------------------------------------------- #
#  Question parser (heuristic)
# --------------------------------------------------------------------- #

# Words that strongly hint at a design when present near a column name.
_DESIGN_HINTS = {
    "did": "did",
    "difference-in-differences": "did",
    "diff-in-diff": "did",
    "parallel trends": "did",
    "rd": "rd",
    "discontinuity": "rd",
    "regression discontinuity": "rd",
    "iv": "iv",
    "instrument": "iv",
    "instrumental variable": "iv",
    "rct": "rct",
    "randomi": "rct",  # randomi(z|s)ation
    "experiment": "rct",
    "synthetic control": "synth",
    "event study": "did",
}


def parse_question(question: str, columns: List[str]) -> Dict[str, Any]:
    """Heuristic parse of a natural-language causal question.

    Returns a dict of *hints* the caller can fall back on when explicit
    column kwargs aren't provided. Never overrides explicit args.

    Parameters
    ----------
    question : str
        Natural-language question, e.g. ``"effect of training on wages"``.
    columns : list of str
        Columns of the dataset; the parser only proposes column names
        present in this list.

    Returns
    -------
    dict
        Possible keys: ``y``, ``treatment``, ``design``,
        ``instrument``, ``running_var``, ``cutoff``,
        ``raw_question``.
    """
    out: Dict[str, Any] = {"raw_question": question}
    if not isinstance(question, str) or not question.strip():
        return out
    q = question.lower()
    cols_lower = {c.lower(): c for c in columns}

    # Design hint
    for key, design in _DESIGN_HINTS.items():
        if key in q:
            out["design"] = design
            break

    # Pattern: "effect of X on Y" / "impact of X on Y" / "X on Y"
    m = re.search(
        r"(?:effect|impact|causal effect|relationship|influence)\s+"
        r"of\s+([a-z0-9_]+)\s+on\s+([a-z0-9_]+)",
        q,
    )
    if m:
        treat, y = m.group(1), m.group(2)
        if treat in cols_lower:
            out.setdefault("treatment", cols_lower[treat])
        if y in cols_lower:
            out.setdefault("y", cols_lower[y])

    # Pattern: "Y ~ X" / "Y = X"
    m = re.search(r"([a-z0-9_]+)\s*[~=]\s*([a-z0-9_]+)", q)
    if m:
        y, treat = m.group(1), m.group(2)
        if y in cols_lower:
            out.setdefault("y", cols_lower[y])
        if treat in cols_lower:
            out.setdefault("treatment", cols_lower[treat])

    # Pattern: "instrument <Z>" / "using <Z> as an instrument"
    m = re.search(
        r"(?:instrument(?:ing)?|using)\s+([a-z0-9_]+)\s+as\s+(?:an?\s+)?"
        r"instrument",
        q,
    )
    if m and m.group(1) in cols_lower:
        out["instrument"] = cols_lower[m.group(1)]
        out["design"] = "iv"

    # Pattern: "discontinuity at <c>" / "threshold <c>"
    m = re.search(
        r"(?:discontinuity|threshold|cutoff)\s+(?:at\s+)?(-?\d+\.?\d*)", q
    )
    if m:
        try:
            out["cutoff"] = float(m.group(1))
            out["design"] = "rd"
        except ValueError:
            pass
    # Pattern: "running variable <X>"
    m = re.search(r"running\s+variable\s+([a-z0-9_]+)", q)
    if m and m.group(1) in cols_lower:
        out["running_var"] = cols_lower[m.group(1)]
        out["design"] = "rd"

    return out


# --------------------------------------------------------------------- #
#  PaperDraft
# --------------------------------------------------------------------- #


@dataclass
class PaperDraft:
    """Draft causal-analysis report assembled by :func:`sp.paper`.

    Attributes
    ----------
    question : str
        The original natural-language question.
    sections : dict[str, str]
        Mapping ``section_title -> markdown_body``. Always includes at
        least: ``Question``, ``Data``, ``Identification``,
        ``Estimator``, ``Results``, ``Robustness``, ``References``.
    workflow : CausalWorkflow
        The underlying workflow object — exposes the raw fitted result
        (``workflow.result``), the diagnostics, the recommendation, etc.
    fmt : str
        Default output format (``markdown`` / ``tex`` / ``docx``).
    citations : list of str
        BibTeX-style entries collected from each estimator's ``cite()``.
    parsed_hints : dict
        What the question parser extracted, for transparency / debugging.
    """
    question: str
    sections: Dict[str, str]
    workflow: Any
    fmt: str
    citations: List[str] = field(default_factory=list)
    parsed_hints: Dict[str, Any] = field(default_factory=dict)

    # ----- rendering -------------------------------------------------- #

    def to_markdown(self) -> str:
        order = [
            "Question", "Data", "Identification",
            "Estimator", "Results", "Robustness", "References",
        ]
        chunks: List[str] = []
        for title in order:
            body = self.sections.get(title)
            if not body:
                continue
            chunks.append(f"## {title}\n\n{body.rstrip()}\n")
        # Append any extra sections in insertion order.
        for title, body in self.sections.items():
            if title not in order and body:
                chunks.append(f"## {title}\n\n{body.rstrip()}\n")
        return "\n".join(chunks)

    def to_tex(self) -> str:
        """Render to a LaTeX article skeleton.

        Each section becomes ``\\section{...}``; markdown bullet lists
        and code fences are translated to LaTeX equivalents.
        """
        body_lines: List[str] = []
        for title, body in self.sections.items():
            body_lines.append(f"\\section{{{_tex_escape(title)}}}")
            body_lines.append(_md_to_tex(body))
            body_lines.append("")
        bib = ""
        if self.citations:
            bib_items = "\n".join(f"\\bibitem{{r{i}}} {_tex_escape(c)}"
                                  for i, c in enumerate(self.citations))
            bib = (
                "\\begin{thebibliography}{99}\n"
                f"{bib_items}\n"
                "\\end{thebibliography}\n"
            )
        return (
            "\\documentclass{article}\n"
            "\\usepackage[utf8]{inputenc}\n"
            "\\usepackage{hyperref}\n"
            "\\title{Causal Analysis Draft}\n"
            "\\begin{document}\n"
            "\\maketitle\n\n"
            + "\n".join(body_lines)
            + "\n" + bib +
            "\\end{document}\n"
        )

    def to_docx(self, path: str) -> None:
        """Write a Word document to ``path``.

        Uses the workflow's already-fit result's ``to_docx`` if available;
        otherwise falls back to dropping a markdown file with a ``.docx``
        warning header (no python-docx hard dep).
        """
        try:
            import docx  # type: ignore
        except ImportError:
            # Fallback: write markdown to disk with a notice.
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(
                    "# (python-docx not installed; markdown fallback)\n\n"
                )
                fh.write(self.to_markdown())
            return
        doc = docx.Document()
        doc.add_heading("Causal Analysis Draft", level=0)
        for title, body in self.sections.items():
            doc.add_heading(title, level=1)
            for line in body.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
        doc.save(path)

    def to_qmd(
        self,
        *,
        title: str = "Causal Analysis Draft",
        author: Optional[str] = None,
        formats: Optional[List[str]] = None,
        bibliography: Optional[str] = None,
        csl: Optional[str] = None,
        include_provenance: bool = True,
    ) -> str:
        """Render to a Quarto (``.qmd``) document.

        Quarto is the publication-grade default: a single source compiles
        to PDF / HTML / DOCX / Beamer with cross-refs, citations (CSL),
        and embedded code chunks. ``sp.paper()`` already produces all
        the prose; this method just wraps it in the correct YAML
        frontmatter so ``quarto render paper.qmd`` Just Works.

        Parameters
        ----------
        title : str
            ``title:`` field in the YAML frontmatter.
        author : str, optional
            ``author:``. When omitted, no author line is emitted (Quarto
            handles that fine).
        formats : list of str, optional
            Output formats Quarto should support. Default
            ``["pdf", "html", "docx"]`` covers the common journal
            workflows. Pass e.g. ``["pdf", "beamer"]`` for slide decks.
        bibliography : str, optional
            Path Quarto should resolve for citation lookup, e.g.
            ``"paper.bib"``. When omitted, the YAML omits the field
            entirely (so Quarto won't error if no .bib file exists);
            but if ``self.citations`` is non-empty we default to
            ``"paper.bib"`` because :func:`sp.replication_pack` writes
            citations there alongside the rendered draft.
        csl : str, optional
            CSL style file (e.g. ``"american-economic-association.csl"``).
            Pure pass-through.
        include_provenance : bool, default True
            Append a Reproducibility appendix with
            :func:`format_provenance` when ``self.workflow.result`` carries
            a ``_provenance`` record.

        Returns
        -------
        str
            The complete ``.qmd`` document as a single string.

        Notes
        -----
        - The body sections are the same as :meth:`to_markdown` —
          standard markdown with ``## H2`` headers, which Quarto will
          render natively.
        - Code chunks are *not* injected by default. When the calling
          script wants the ``.qmd`` to re-execute the analysis on each
          render, pass it through :func:`sp.replication_pack` which
          writes both the ``.qmd`` and a ``code/script.py`` reproducer.
        """
        formats = formats or ["pdf", "html", "docx"]
        bib_path = bibliography
        if bib_path is None and self.citations:
            bib_path = "paper.bib"

        yaml_lines: List[str] = ["---", f"title: {_yaml_str(title)}"]
        if author:
            yaml_lines.append(f"author: {_yaml_str(author)}")
        yaml_lines.append(f"date: \"{_dt.date.today().isoformat()}\"")
        if self.question:
            yaml_lines.append(
                f"subtitle: {_yaml_str(self.question)}"
            )
        # ``format:`` block.
        if len(formats) == 1:
            yaml_lines.append(f"format: {formats[0]}")
        else:
            yaml_lines.append("format:")
            for f in formats:
                yaml_lines.append(f"  {f}: default")
        if bib_path:
            yaml_lines.append(f"bibliography: {_yaml_str(bib_path)}")
        if csl:
            # Accept short journal names ('aer' / 'qje' / ...) and resolve
            # them to the canonical .csl filename. Pre-existing .csl
            # paths pass through untouched.
            try:
                from ..output._bibliography import csl_filename
                resolved = csl_filename(csl)
            except Exception:
                resolved = csl
            yaml_lines.append(f"csl: {_yaml_str(resolved)}")
        # Provenance into YAML for machine-readable traceability.
        prov = self._workflow_provenance()
        if include_provenance and prov is not None:
            yaml_lines.append("statspai:")
            yaml_lines.append(
                f"  version: \"{prov.statspai_version}\""
            )
            yaml_lines.append(f"  run_id: \"{prov.run_id}\"")
            if prov.data_hash:
                yaml_lines.append(
                    f"  data_hash: \"{prov.data_hash}\""
                )
        yaml_lines.append("---")
        yaml = "\n".join(yaml_lines)

        # Body — identical section ordering to to_markdown().
        order = [
            "Question", "Data", "Identification",
            "Estimator", "Results", "Robustness", "References",
        ]
        chunks: List[str] = []
        for t in order:
            body = self.sections.get(t)
            if not body:
                continue
            chunks.append(f"## {t}\n\n{body.rstrip()}\n")
        for t, body in self.sections.items():
            if t not in order and body:
                chunks.append(f"## {t}\n\n{body.rstrip()}\n")

        # Reproducibility appendix.
        if include_provenance and prov is not None:
            chunks.append(
                "## Reproducibility {.appendix}\n\n"
                "```\n"
                f"{format_provenance(prov)}\n"
                "```\n"
            )

        return yaml + "\n\n" + "\n".join(chunks)

    def _workflow_provenance(self):
        wf = self.workflow
        if wf is None:
            return None
        result = getattr(wf, "result", None)
        if result is None:
            return None
        return get_provenance(result)

    def write(self, path: str) -> None:
        """Write the draft to disk in the format inferred from the path
        extension (``.md`` / ``.tex`` / ``.docx`` / ``.qmd``)."""
        lower = path.lower()
        if lower.endswith('.tex'):
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(self.to_tex())
        elif lower.endswith('.docx'):
            self.to_docx(path)
        elif lower.endswith('.qmd'):
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(self.to_qmd())
        else:
            with open(path, 'w', encoding='utf-8') as fh:
                fh.write(self.to_markdown())

    def summary(self) -> str:
        n = len(self.sections)
        return (
            "PaperDraft\n"
            "==========\n"
            f"  Question     : {self.question[:60]!r}"
            f"{'...' if len(self.question) > 60 else ''}\n"
            f"  Sections     : {n} ({', '.join(self.sections.keys())})\n"
            f"  Citations    : {len(self.citations)}\n"
            f"  Default fmt  : {self.fmt}\n"
        )

    def to_dict(self) -> Dict[str, Any]:
        return {
            'question': self.question,
            'sections': dict(self.sections),
            'parsed_hints': dict(self.parsed_hints),
            'citations': list(self.citations),
            'fmt': self.fmt,
        }


# --------------------------------------------------------------------- #
#  Helpers
# --------------------------------------------------------------------- #


def _yaml_str(value: str) -> str:
    """Quote a string safely for inclusion in a YAML scalar value.

    Always uses double-quotes and escapes any embedded ``"`` / ``\\``.
    Newlines are folded to a literal space (Quarto YAML headers don't
    play nicely with multi-line scalars in our context).
    """
    if value is None:
        return '""'
    s = str(value).replace("\n", " ").strip()
    # Backslash first, then double-quote.
    s = s.replace("\\", "\\\\").replace('"', '\\"')
    return f'"{s}"'


def _tex_escape(s: str) -> str:
    """Minimal LaTeX-escape for free-form text in section bodies."""
    if not isinstance(s, str):
        s = str(s)
    out = (s.replace("\\", r"\textbackslash{}")
           .replace("&", r"\&")
           .replace("%", r"\%")
           .replace("$", r"\$")
           .replace("#", r"\#")
           .replace("_", r"\_")
           .replace("{", r"\{")
           .replace("}", r"\}")
           .replace("~", r"\textasciitilde{}")
           .replace("^", r"\textasciicircum{}"))
    return out


def _md_to_tex(md: str) -> str:
    """Lightweight markdown → LaTeX translation for paper-section bodies.

    Handles: bold (**text**), bullet lists, fenced code blocks, inline
    code (`x`). Anything more elaborate falls through as escaped text.
    """
    out_lines: List[str] = []
    in_list = False
    in_code = False
    for ln in md.split("\n"):
        stripped = ln.rstrip()
        if stripped.startswith("```"):
            if in_code:
                out_lines.append(r"\end{verbatim}")
                in_code = False
            else:
                if in_list:
                    out_lines.append(r"\end{itemize}")
                    in_list = False
                out_lines.append(r"\begin{verbatim}")
                in_code = True
            continue
        if in_code:
            out_lines.append(stripped)
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            if not in_list:
                out_lines.append(r"\begin{itemize}")
                in_list = True
            item = stripped[2:].strip()
            out_lines.append(r"  \item " + _inline_md_to_tex(item))
            continue
        if in_list:
            out_lines.append(r"\end{itemize}")
            in_list = False
        if stripped == "":
            out_lines.append("")
        else:
            out_lines.append(_inline_md_to_tex(stripped))
    if in_list:
        out_lines.append(r"\end{itemize}")
    if in_code:
        out_lines.append(r"\end{verbatim}")
    return "\n".join(out_lines)


def _inline_md_to_tex(text: str) -> str:
    """Translate inline-markdown markers to LaTeX. Order matters: handle
    code spans before bold so we don't escape backticks inside code."""
    out = text
    # Code spans `...`
    out = re.sub(r"`([^`]+)`",
                 lambda m: r"\texttt{" + _tex_escape(m.group(1)) + "}",
                 out)
    # Bold **...**
    out = re.sub(r"\*\*([^*]+)\*\*",
                 lambda m: r"\textbf{" + _tex_escape(m.group(1)) + "}",
                 out)
    # Italic *...* (after bold so we don't double-match)
    out = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)",
                 lambda m: r"\emph{" + _tex_escape(m.group(1)) + "}",
                 out)
    # Anything left is plain text — escape special chars but leave
    # already-emitted LaTeX commands alone. The previous regexes have
    # produced LaTeX commands containing braces / backslashes; do a
    # cheap heuristic: only escape lines that don't already contain
    # ``\textbf`` / ``\texttt`` / ``\emph``.
    if (r"\textbf" not in out and r"\texttt" not in out
            and r"\emph" not in out):
        out = _tex_escape(out)
    return out


def _eda_block(data: pd.DataFrame, y: Optional[str],
               treatment: Optional[str],
               covariates: Optional[List[str]]) -> str:
    """Build a brief EDA markdown section (size, balance, missingness)."""
    lines: List[str] = []
    n_rows, n_cols = data.shape
    lines.append(f"- Sample size: **{n_rows:,}** rows, **{n_cols}** columns.")
    miss = data.isna().mean()
    miss = miss[miss > 0].sort_values(ascending=False)
    if not miss.empty:
        lines.append("- Missingness (top 5):")
        for col, frac in miss.head(5).items():
            lines.append(f"    - `{col}`: {frac*100:.1f}%")
    else:
        lines.append("- Missingness: none detected in the analysis frame.")
    if y and y in data.columns:
        ys = data[y].dropna()
        if pd.api.types.is_numeric_dtype(ys):
            lines.append(
                f"- Outcome `{y}`: "
                f"mean={ys.mean():.3f}, sd={ys.std():.3f}, "
                f"median={ys.median():.3f}, n={len(ys)}."
            )
    if treatment and treatment in data.columns:
        tr = data[treatment].dropna()
        if tr.nunique() <= 10:
            counts = tr.value_counts().sort_index()
            shares = (counts / counts.sum() * 100).round(1)
            lines.append(
                f"- Treatment `{treatment}` distribution: "
                + ", ".join(
                    f"{int(k) if k == int(k) else k}="
                    f"{int(counts.loc[k])} ({shares.loc[k]}%)"
                    for k in counts.index
                )
            )
        else:
            lines.append(
                f"- Treatment `{treatment}` (continuous): "
                f"mean={tr.mean():.3f}, sd={tr.std():.3f}."
            )
    # Optional covariate balance — only for binary treatment.
    if (treatment and treatment in data.columns
            and covariates and len(covariates) <= 8
            and data[treatment].nunique() == 2):
        try:
            grp = data.groupby(treatment)[covariates].mean()
            if grp.shape[0] == 2:
                lines.append("")
                lines.append("Mean covariates by treatment arm:")
                lines.append("")
                lines.append("| covariate | "
                             + " | ".join(str(g) for g in grp.index)
                             + " | std-diff |")
                lines.append("|---|" + "|".join(["---"] * grp.shape[0])
                             + "|---|")
                pooled_std = data[covariates].std()
                vals0 = grp.iloc[0]
                vals1 = grp.iloc[1]
                std_diff = (vals1 - vals0) / pooled_std.replace(0, np.nan)
                for c in covariates:
                    lines.append(
                        f"| {c} | {grp.iloc[0][c]:.3f} "
                        f"| {grp.iloc[1][c]:.3f} "
                        f"| {std_diff[c]:.3f} |"
                    )
        except Exception:
            pass
    return "\n".join(lines)


def _section_from_workflow(workflow) -> Dict[str, str]:
    """Extract Identification / Estimator / Results / Robustness sections
    from a fitted CausalWorkflow."""
    sections: Dict[str, str] = {}

    # Identification
    diag = workflow.diagnostics
    lines: List[str] = []
    lines.append(f"**Verdict**: {diag.verdict}")
    lines.append("")
    if diag.findings:
        for f in diag.findings:
            lines.append(f"- [{f.severity.upper()}] *{f.category}* — "
                         f"{f.message}")
            if f.suggestion:
                lines.append(f"    - Fix: {f.suggestion}")
    else:
        lines.append("No identification issues flagged.")
    sections["Identification"] = "\n".join(lines)

    # Estimator
    rec = workflow.recommendation
    lines = []
    if rec is not None and rec.recommendations:
        top = rec.recommendations[0]
        lines.append(f"- **Method**: {top['method']}")
        lines.append(f"- **Function**: `sp.{top['function']}()`")
        if top.get('reason'):
            lines.append(f"- **Rationale**: {top['reason']}")
        if top.get('assumptions'):
            lines.append("- **Key assumptions**: "
                         + ", ".join(top['assumptions']))
    else:
        lines.append("No estimator recommendation produced.")
    sections["Estimator"] = "\n".join(lines)

    # Results
    r = workflow.result
    lines = []
    if r is not None and hasattr(r, 'estimate') and hasattr(r, 'se'):
        try:
            est = float(r.estimate)
            se = float(r.se)
            lines.append(
                f"- **{getattr(r, 'estimand', 'Effect')}**: "
                f"{est:.4f} (SE = {se:.4f})"
            )
            ci = getattr(r, 'ci', None)
            if ci is not None and not isinstance(ci, (pd.DataFrame, pd.Series)):
                try:
                    lo, hi = float(ci[0]), float(ci[1])
                    lines.append(f"- **95% CI**: [{lo:.4f}, {hi:.4f}]")
                except Exception:
                    pass
            pv = getattr(r, 'pvalue', np.nan)
            if pd.notna(pv):
                lines.append(f"- **p-value**: {float(pv):.4f}")
            n_obs = getattr(r, 'n_obs', None)
            if n_obs is not None:
                lines.append(f"- **N obs**: {int(n_obs)}")
        except Exception:
            lines.append("Result available but not fully serialisable; "
                         "see `paper.workflow.result`.")
    elif r is not None and hasattr(r, 'params'):
        try:
            main = (workflow.treatment or list(r.params.index)[0])
            if main in r.params.index:
                est = float(r.params[main])
                se = float(r.std_errors[main])
                lines.append(f"- **{main}**: {est:.4f} (SE = {se:.4f})")
            else:
                lines.append("Coefficient table available; see "
                             "`paper.workflow.result.params`.")
        except Exception:
            lines.append("Result available but not fully serialisable.")
    else:
        lines.append("No fitted result available.")
    sections["Results"] = "\n".join(lines)

    # Robustness
    findings = workflow.robustness_findings or {}
    lines = []
    if findings:
        for k, v in findings.items():
            if isinstance(v, (int, float, np.integer, np.floating)):
                if isinstance(v, (int, np.integer)):
                    lines.append(f"- {k.replace('_', ' ').title()}: "
                                 f"{int(v)}")
                else:
                    lines.append(f"- {k.replace('_', ' ').title()}: "
                                 f"{float(v):.4f}")
            elif isinstance(v, dict):
                lines.append(f"- {k.replace('_', ' ').title()}:")
                for kk, vv in list(v.items())[:8]:
                    lines.append(f"    - {kk}: {vv}")
            else:
                lines.append(f"- {k.replace('_', ' ').title()}: {v}")
    else:
        lines.append("No robustness findings produced.")
    sections["Robustness"] = "\n".join(lines)

    return sections


# --------------------------------------------------------------------- #
#  Top-level entry point
# --------------------------------------------------------------------- #


def paper(
    data: pd.DataFrame,
    question: str,
    *,
    y: Optional[str] = None,
    treatment: Optional[str] = None,
    covariates: Optional[List[str]] = None,
    id: Optional[str] = None,
    time: Optional[str] = None,
    running_var: Optional[str] = None,
    instrument: Optional[str] = None,
    cutoff: Optional[float] = None,
    cohort: Optional[str] = None,
    cluster: Optional[str] = None,
    design: Optional[str] = None,
    dag=None,
    fmt: str = 'markdown',
    output_path: Optional[str] = None,
    include_eda: bool = True,
    include_robustness: bool = True,
    cite: bool = True,
    strict: bool = False,
) -> PaperDraft:
    """End-to-end "data → publication-draft" pipeline.

    Run :class:`CausalWorkflow` and assemble its outputs into a
    structured :class:`PaperDraft` (markdown / LaTeX / Word).

    Parameters
    ----------
    data : pd.DataFrame
        Analysis frame.
    question : str
        Natural-language causal question, e.g.
        ``"effect of training on wages"``. Used both to seed the
        question section of the draft and to fill in any missing
        column hints (``treatment`` / ``y``) when not given explicitly.
    y, treatment, covariates, id, time, running_var, instrument, cutoff, cohort, cluster, design, dag : optional
        Forwarded to :func:`sp.causal`. When omitted, the question
        parser tries to infer them from ``question``.
    fmt : {'markdown', 'tex', 'docx'}, default 'markdown'
        Default rendering format. The :class:`PaperDraft` always knows
        how to emit each format on demand via its ``.to_*()`` methods.
    output_path : str, optional
        When provided, write the rendered draft to disk in the format
        inferred from the path extension (``.md`` / ``.tex`` /
        ``.docx``).
    include_eda : bool, default True
        Include the Data section (descriptives + balance).
    include_robustness : bool, default True
        Include the Robustness section.
    cite : bool, default True
        Pull bibliography entries from the fitted result's ``cite()``
        method (when available).
    strict : bool, default False
        Forwarded to :func:`sp.causal` — when True, identification
        warnings escalate to errors.

    Returns
    -------
    PaperDraft

    Examples
    --------
    >>> import statspai as sp
    >>> draft = sp.paper(df, "effect of training on wages", design='did',
    ...                  treatment='trained', y='wage', time='year',
    ...                  id='worker_id')
    >>> print(draft.to_markdown()[:500])
    >>> draft.write('analysis.tex')

    Notes
    -----
    The question parser is purely additive — explicit kwargs always win.
    Pass everything you know; the parser fills in only what's missing.
    """
    if fmt not in {'markdown', 'tex', 'docx', 'qmd'}:
        raise ValueError(
            f"Unknown fmt={fmt!r}. Use 'markdown', 'tex', 'docx', or 'qmd'."
        )

    cols = list(data.columns)
    parsed = parse_question(question or "", cols)

    # Explicit args win; parser fills gaps.
    y_eff = y or parsed.get('y')
    t_eff = treatment or parsed.get('treatment')
    design_eff = design or parsed.get('design')
    instrument_eff = instrument or parsed.get('instrument')
    running_var_eff = running_var or parsed.get('running_var')
    cutoff_eff = cutoff if cutoff is not None else parsed.get('cutoff')

    if y_eff is None:
        raise ValueError(
            "Could not determine the outcome `y`. Pass `y=...` explicitly "
            "or include 'effect of X on Y' in the question."
        )

    from .causal_workflow import causal as _causal
    workflow = _causal(
        data,
        y=y_eff,
        treatment=t_eff,
        covariates=covariates,
        id=id,
        time=time,
        running_var=running_var_eff,
        instrument=instrument_eff,
        cutoff=cutoff_eff,
        cohort=cohort,
        cluster=cluster,
        design=design_eff,
        dag=dag,
        strict=strict,
    )
    # Drive the pipeline through to robustness, swallowing per-stage
    # failures into per-section fallback notes (the draft must always
    # produce something — agents shouldn't see a hard crash for one
    # bad estimator choice).
    pipeline_errors: List[str] = []
    for stage in ("diagnose", "recommend", "estimate"):
        try:
            getattr(workflow, stage)()
        except Exception as exc:  # pragma: no cover (defensive)
            pipeline_errors.append(
                f"`{stage}()` failed: {type(exc).__name__}: {exc}"
            )
            break
    if include_robustness:
        try:
            workflow.robustness()
        except Exception as exc:  # pragma: no cover (defensive)
            pipeline_errors.append(
                f"`robustness()` failed: {type(exc).__name__}: {exc}"
            )

    # Attach provenance to the workflow's result so downstream
    # ``replication_pack`` / Quarto appendix / table footers can pick
    # it up. Estimators that wire their own ``attach_provenance`` at
    # ``fit()`` end already populate ``_provenance``; ``overwrite=False``
    # preserves their (more specific) record.
    try:
        from ..output._lineage import attach_provenance as _attach_prov
        if workflow.result is not None:
            _attach_prov(
                workflow.result,
                function=f"sp.causal[{workflow.design or 'auto'}]",
                params={
                    "y": y_eff,
                    "treatment": t_eff,
                    "design": design_eff or workflow.design,
                    "instrument": instrument_eff,
                    "running_var": running_var_eff,
                    "cutoff": cutoff_eff,
                    "covariates": covariates,
                },
                data=data,
                overwrite=False,
            )
    except Exception:  # pragma: no cover — provenance must never break the draft
        pass

    sections: Dict[str, str] = {}

    # Question
    sections["Question"] = (
        f"> {question.strip() if question else '(no question supplied)'}\n\n"
        f"- **Outcome**: `{y_eff}`\n"
        + (f"- **Treatment**: `{t_eff}`\n" if t_eff else "")
        + (f"- **Design (auto-detected)**: `{workflow.design}`\n"
           if workflow.design else "")
    )

    # Data / EDA
    if include_eda:
        sections["Data"] = _eda_block(data, y_eff, t_eff, covariates)

    # Identification / Estimator / Results / Robustness
    sections.update(_section_from_workflow(workflow))

    if pipeline_errors:
        sections["Pipeline notes"] = "\n".join(
            f"- {e}" for e in pipeline_errors
        )

    # References
    citations: List[str] = []
    if cite and workflow.result is not None:
        cite_fn = getattr(workflow.result, 'cite', None)
        if callable(cite_fn):
            try:
                ref = cite_fn()
                if ref:
                    citations.append(str(ref))
            except Exception:
                pass
    sections["References"] = (
        "\n".join(f"- {c}" for c in citations)
        if citations else "_(No explicit citations attached — see "
        "`workflow.result.cite()` if available.)_"
    )

    draft = PaperDraft(
        question=question or "",
        sections=sections,
        workflow=workflow,
        fmt=fmt,
        citations=citations,
        parsed_hints=parsed,
    )

    if output_path is not None:
        draft.write(output_path)

    return draft
