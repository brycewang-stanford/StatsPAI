"""
Unified results class for all econometric models
"""

from typing import Dict, Any, Optional, List, Union
import pandas as pd
import numpy as np
from scipy import stats


class EconometricResults:
    """
    Unified results class for econometric models
    
    This class provides a consistent interface for accessing results
    from different econometric estimators, similar to R's broom package.
    """
    
    def __init__(
        self,
        params: pd.Series,
        std_errors: pd.Series,
        model_info: Dict[str, Any],
        data_info: Optional[Dict[str, Any]] = None,
        diagnostics: Optional[Dict[str, Any]] = None
    ):
        """
        Initialize results object
        
        Parameters
        ----------
        params : pd.Series
            Parameter estimates with variable names as index
        std_errors : pd.Series
            Standard errors with variable names as index
        model_info : Dict[str, Any]
            Model metadata (model type, estimation method, etc.)
        data_info : Dict[str, Any], optional
            Data metadata (sample size, variable names, etc.)
        diagnostics : Dict[str, Any], optional
            Model diagnostics (R-squared, F-statistics, etc.)
        """
        self.params = params
        self.std_errors = std_errors
        self.model_info = model_info
        self.data_info = data_info or {}
        self.diagnostics = diagnostics or {}
        
        # Compute derived statistics
        self._compute_statistics()
    
    def _compute_statistics(self):
        """Compute t-statistics, p-values, and confidence intervals"""
        self.tvalues = self.params / self.std_errors
        self.pvalues = 2 * (1 - stats.t.cdf(np.abs(self.tvalues), 
                                           self.data_info.get('df_resid', np.inf)))
        
        # 95% confidence intervals by default
        alpha = 0.05
        t_crit = stats.t.ppf(1 - alpha/2, self.data_info.get('df_resid', np.inf))
        self.conf_int_lower = self.params - t_crit * self.std_errors
        self.conf_int_upper = self.params + t_crit * self.std_errors
    
    def summary(self, alpha: float = 0.05) -> str:
        """
        Generate a summary table of results
        
        Parameters
        ----------
        alpha : float, default 0.05
            Significance level for confidence intervals
            
        Returns
        -------
        str
            Formatted summary table
        """
        # Create coefficients table
        coef_table = pd.DataFrame({
            'Coefficient': self.params,
            'Std. Error': self.std_errors,
            't-statistic': self.tvalues,
            'P>|t|': self.pvalues,
            f'[{alpha/2:.3f}': self.conf_int_lower,
            f'{1-alpha/2:.3f}]': self.conf_int_upper
        })
        
        # Format the output
        output = []
        output.append("=" * 80)
        output.append(f"Model: {self.model_info.get('model_type', 'Unknown')}")
        output.append(f"Method: {self.model_info.get('method', 'Unknown')}")
        if 'dependent_var' in self.data_info:
            output.append(f"Dependent Variable: {self.data_info['dependent_var']}")
        output.append("=" * 80)
        
        # Add coefficient table
        output.append(coef_table.to_string(float_format='%.4f'))
        
        # Add model diagnostics
        if self.diagnostics:
            output.append("")
            output.append("Model Diagnostics:")
            output.append("-" * 20)
            for key, value in self.diagnostics.items():
                if isinstance(value, (int, float)):
                    output.append(f"{key:20s}: {value:.4f}")
                else:
                    output.append(f"{key:20s}: {value}")
        
        output.append("=" * 80)
        return "\n".join(output)
    
    def conf_int(self, alpha: float = 0.05) -> pd.DataFrame:
        """
        Return confidence intervals for parameters
        
        Parameters
        ----------
        alpha : float, default 0.05
            Significance level
            
        Returns
        -------
        pd.DataFrame
            Confidence intervals
        """
        t_crit = stats.t.ppf(1 - alpha/2, self.data_info.get('df_resid', np.inf))
        lower = self.params - t_crit * self.std_errors
        upper = self.params + t_crit * self.std_errors
        
        return pd.DataFrame({
            f'{alpha/2:.3f}': lower,
            f'{1-alpha/2:.3f}': upper
        }, index=self.params.index)
    
    def predict(self, data: Optional[pd.DataFrame] = None) -> np.ndarray:
        """
        Generate predictions (to be implemented by specific models)
        
        Parameters
        ----------
        data : pd.DataFrame, optional
            Data for prediction
            
        Returns
        -------
        np.ndarray
            Predicted values
        """
        raise NotImplementedError("Prediction method not implemented for this model")
    
    def residuals(self) -> Optional[np.ndarray]:
        """
        Return model residuals if available
        
        Returns
        -------
        np.ndarray or None
            Residuals
        """
        return self.data_info.get('residuals')
    
    def fitted_values(self) -> Optional[np.ndarray]:
        """
        Return fitted values if available
        
        Returns
        -------
        np.ndarray or None
            Fitted values
        """
        return self.data_info.get('fitted_values')
    
    def to_docx(self, filename: str, title: Optional[str] = None):
        """
        Export results to a Word (.docx) document.

        Parameters
        ----------
        filename : str
            Output path (.docx).
        title : str, optional
            Table title. Defaults to model type.
        """
        _result_to_docx(self, filename, title)

    def _repr_html_(self) -> str:
        """Rich HTML display for Jupyter notebooks."""
        model_type = self.model_info.get('model_type', 'Unknown')
        method = self.model_info.get('method', '')
        dep_var = self.data_info.get('dependent_var', '')
        n_obs = self.data_info.get('nobs', '?')
        r2 = self.diagnostics.get('R-squared', None)
        f_stat = self.diagnostics.get('F-statistic', None)
        f_pv = self.diagnostics.get('F p-value', self.diagnostics.get('Prob (F-statistic)', None))

        def _s(pv):
            if pd.isna(pv): return ''
            if pv < 0.01: return '<span style="color:#E74C3C;">***</span>'
            if pv < 0.05: return '<span style="color:#E67E22;">**</span>'
            if pv < 0.1: return '<span style="color:#F39C12;">*</span>'
            return ''

        def _val(v):
            return f'{v:.4f}' if isinstance(v, float) else str(v)

        # CSS
        S = ('<style scoped>'
             '.sp-box{font-family:"Helvetica Neue",Arial,sans-serif;max-width:720px;border:1px solid #E5E7EB;border-radius:8px;overflow:hidden;margin:6px 0}'
             '.sp-hdr{background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;padding:12px 16px}'
             '.sp-hdr h3{margin:0;font-size:15px;font-weight:600;letter-spacing:0.3px}'
             '.sp-hdr .sp-sub{font-size:11px;color:#94A3B8;margin-top:2px}'
             '.sp-metrics{display:flex;gap:0;border-bottom:1px solid #E5E7EB}'
             '.sp-metric{flex:1;padding:10px 14px;text-align:center;border-right:1px solid #E5E7EB}'
             '.sp-metric:last-child{border-right:none}'
             '.sp-metric .sp-val{font-size:18px;font-weight:700;color:#1a1a2e}'
             '.sp-metric .sp-lab{font-size:10px;color:#94A3B8;text-transform:uppercase;letter-spacing:0.5px;margin-top:2px}'
             'table.sp-coef{width:100%;border-collapse:collapse;font-size:12px}'
             'table.sp-coef th{padding:6px 10px;text-align:right;font-weight:600;color:#64748B;border-bottom:2px solid #E5E7EB;font-size:11px}'
             'table.sp-coef th:first-child{text-align:left}'
             'table.sp-coef td{padding:5px 10px;text-align:right;border-bottom:1px solid #F1F5F9}'
             'table.sp-coef td:first-child{text-align:left;font-weight:500;color:#1a1a2e}'
             'table.sp-coef tr:hover{background:#F8FAFC}'
             '.sp-diag{display:grid;grid-template-columns:1fr 1fr;gap:0;border-top:1px solid #E5E7EB;font-size:11px}'
             '.sp-diag-item{padding:4px 14px;display:flex;justify-content:space-between;border-bottom:1px solid #F8FAFC}'
             '.sp-diag-item:nth-child(odd){border-right:1px solid #F1F5F9}'
             '.sp-diag-k{color:#94A3B8}.sp-diag-v{color:#334155;font-weight:500}'
             '.sp-foot{padding:6px 14px;font-size:10px;color:#94A3B8;border-top:1px solid #E5E7EB;display:flex;justify-content:space-between}'
             '</style>')

        h = [S, '<div class="sp-box">']

        # --- Header ---
        sub_parts = []
        if dep_var:
            sub_parts.append(f'Y = {dep_var}')
        if method:
            sub_parts.append(method)
        h.append(f'<div class="sp-hdr"><h3>{model_type}</h3>')
        if sub_parts:
            h.append(f'<div class="sp-sub">{" · ".join(sub_parts)}</div>')
        h.append('</div>')

        # --- Key Metrics Bar ---
        h.append('<div class="sp-metrics">')
        if r2 is not None:
            h.append(f'<div class="sp-metric"><div class="sp-val">{r2:.4f}</div><div class="sp-lab">R-squared</div></div>')
        if f_stat is not None:
            h.append(f'<div class="sp-metric"><div class="sp-val">{f_stat:.1f}</div><div class="sp-lab">F-statistic</div></div>')
        h.append(f'<div class="sp-metric"><div class="sp-val">{n_obs:,}</div><div class="sp-lab">Observations</div></div>')
        h.append(f'<div class="sp-metric"><div class="sp-val">{len(self.params)}</div><div class="sp-lab">Parameters</div></div>')
        h.append('</div>')

        # --- Coefficient Table ---
        h.append('<table class="sp-coef"><tr>')
        for col in ['', 'Coefficient', 'Std. Error', 't-stat', 'P&gt;|t|', '95% CI']:
            h.append(f'<th>{col}</th>')
        h.append('</tr>')

        for i, var in enumerate(self.params.index):
            coef = self.params.iloc[i]
            se = self.std_errors.iloc[i]
            t = self.tvalues.iloc[i] if isinstance(self.tvalues, pd.Series) else self.tvalues[i]
            pv = self.pvalues.iloc[i] if isinstance(self.pvalues, pd.Series) else self.pvalues[i]
            lo = self.conf_int_lower.iloc[i] if isinstance(self.conf_int_lower, pd.Series) else self.conf_int_lower[i]
            hi = self.conf_int_upper.iloc[i] if isinstance(self.conf_int_upper, pd.Series) else self.conf_int_upper[i]
            pv_color = '#DC2626' if pv < 0.01 else ('#EA580C' if pv < 0.05 else ('#D97706' if pv < 0.1 else '#64748B'))
            h.append(f'<tr><td>{var}</td>')
            h.append(f'<td>{coef:.4f} {_s(pv)}</td>')
            h.append(f'<td style="color:#94A3B8;">({se:.4f})</td>')
            h.append(f'<td>{t:.2f}</td>')
            h.append(f'<td style="color:{pv_color};font-weight:600;">{pv:.4f}</td>')
            h.append(f'<td style="color:#94A3B8;">[{lo:.3f}, {hi:.3f}]</td></tr>')
        h.append('</table>')

        # --- Diagnostics Grid ---
        diag_items = [(k, v) for k, v in self.diagnostics.items()
                      if isinstance(v, (int, float, str)) and k not in ('R-squared', 'F-statistic', 'F p-value', 'Prob (F-statistic)')]
        if diag_items:
            h.append('<div class="sp-diag">')
            for k, v in diag_items:
                h.append(f'<div class="sp-diag-item"><span class="sp-diag-k">{k}</span><span class="sp-diag-v">{_val(v)}</span></div>')
            h.append('</div>')

        # --- IV-specific: First-stage diagnostics ---
        iv_keys = [k for k in self.diagnostics if 'First-stage' in k or 'Hausman' in k or 'Partial' in k or 'Sargan' in k]
        if iv_keys:
            h.append('<details open style="border-top:1px solid #E5E7EB;"><summary style="padding:6px 14px;font-size:12px;'
                     'font-weight:600;color:#1a1a2e;cursor:pointer;">IV Diagnostics</summary>')
            h.append('<div class="sp-diag">')
            for k in iv_keys:
                v = self.diagnostics[k]
                h.append(f'<div class="sp-diag-item"><span class="sp-diag-k">{k}</span><span class="sp-diag-v">{_val(v)}</span></div>')
            h.append('</div></details>')

        # --- Footer ---
        h.append(f'<div class="sp-foot"><span>N = {n_obs:,}</span><span>* p&lt;0.1 &nbsp; ** p&lt;0.05 &nbsp; *** p&lt;0.01</span></div>')
        h.append('</div>')
        return '\n'.join(h)

    def __repr__(self) -> str:
        """String representation of results"""
        model_type = self.model_info.get('model_type', 'Unknown')
        n_params = len(self.params)
        n_obs = self.data_info.get('nobs', 'Unknown')
        return f"<EconometricResults: {model_type}, {n_params} parameters, {n_obs} observations>"


class CausalResult:
    """
    Unified result object for all causal inference methods in StatsPAI.

    All causal inference estimators (DID, RD, SCM, matching, etc.) return
    this object, providing a consistent interface for summaries, plots,
    and publication-quality output.

    Parameters
    ----------
    method : str
        Name of the estimation method (displayed in summary).
    estimand : str
        What is being estimated ('ATT', 'ATE', 'LATE').
    estimate : float
        Point estimate of the main treatment effect.
    se : float
        Standard error.
    pvalue : float
        Two-sided p-value for H0: effect = 0.
    ci : tuple of (float, float)
        Confidence interval (lower, upper).
    alpha : float
        Significance level used for CI.
    n_obs : int
        Number of observations.
    detail : pd.DataFrame, optional
        Detailed estimates (e.g., group-time ATTs).
    model_info : dict, optional
        Model metadata and aggregated results.
    _influence_funcs : np.ndarray, optional
        Influence function matrix (n_units, n_estimates) for joint inference.
    _citation_key : str, optional
        Key into the citation registry.
    """

    _CITATIONS: Dict[str, str] = {
        'did_2x2': (
            "@book{angrist2009mostly,\n"
            "  title={Mostly Harmless Econometrics: An Empiricist's Companion},\n"
            "  author={Angrist, Joshua D and Pischke, J{\\\"o}rn-Steffen},\n"
            "  year={2009},\n"
            "  publisher={Princeton University Press}\n"
            "}"
        ),
        'callaway_santanna': (
            "@article{callaway2021difference,\n"
            "  title={Difference-in-differences with multiple time periods},\n"
            "  author={Callaway, Brantly and Sant'Anna, Pedro H.C.},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={200--230},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'sun_abraham': (
            "@article{sun2021estimating,\n"
            "  title={Estimating dynamic treatment effects in event studies "
            "with heterogeneous treatment effects},\n"
            "  author={Sun, Liyang and Abraham, Sarah},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={175--199},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'rdrobust': (
            "@article{calonico2014robust,\n"
            "  title={Robust nonparametric confidence intervals for "
            "regression-discontinuity designs},\n"
            "  author={Calonico, Sebastian and Cattaneo, Matias D "
            "and Titiunik, Rocio},\n"
            "  journal={Econometrica},\n"
            "  volume={82},\n"
            "  number={6},\n"
            "  pages={2295--2326},\n"
            "  year={2014},\n"
            "  publisher={Wiley}\n"
            "}"
        ),
        'bacon_decomposition': (
            "@article{goodman2021difference,\n"
            "  title={Difference-in-differences with variation in treatment timing},\n"
            "  author={Goodman-Bacon, Andrew},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={254--277},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'did_multiplegt': (
            "@article{dechaisemartin2020two,\n"
            "  title={Two-Way Fixed Effects Estimators with "
            "Heterogeneous Treatment Effects},\n"
            "  author={de Chaisemartin, Cl{\\'e}ment and "
            "D'Haultf{\\oe}uille, Xavier},\n"
            "  journal={American Economic Review},\n"
            "  volume={110},\n"
            "  number={9},\n"
            "  pages={2964--2996},\n"
            "  year={2020}\n"
            "}"
        ),
        'stacked_did': (
            "@article{cengiz2019effect,\n"
            "  title={The Effect of Minimum Wages on Low-Wage Jobs},\n"
            "  author={Cengiz, Doruk and Dube, Arindrajit and "
            "Lindner, Attila and Zipperer, Ben},\n"
            "  journal={Quarterly Journal of Economics},\n"
            "  volume={134},\n"
            "  number={3},\n"
            "  pages={1405--1454},\n"
            "  year={2019},\n"
            "  publisher={Oxford University Press}\n"
            "}"
        ),
        'wooldridge_twfe': (
            "@unpublished{wooldridge2021two,\n"
            "  title={Two-Way Fixed Effects, the Two-Way Mundlak Regression, "
            "and Difference-in-Differences Estimators},\n"
            "  author={Wooldridge, Jeffrey M.},\n"
            "  year={2021},\n"
            "  note={Working paper, Michigan State University}\n"
            "}"
        ),
        'drdid': (
            "@article{santanna2020doubly,\n"
            "  title={Doubly Robust Difference-in-Differences Estimators},\n"
            "  author={Sant'Anna, Pedro H.C. and Zhao, Jun},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={219},\n"
            "  number={1},\n"
            "  pages={101--122},\n"
            "  year={2020},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'twfe_decomposition': (
            "@article{goodman2021difference,\n"
            "  title={Difference-in-differences with variation in treatment timing},\n"
            "  author={Goodman-Bacon, Andrew},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={254--277},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}\n"
            "@article{dechaisemartin2020two,\n"
            "  title={Two-Way Fixed Effects Estimators with "
            "Heterogeneous Treatment Effects},\n"
            "  author={de Chaisemartin, Cl{\\'e}ment and "
            "D'Haultf{\\oe}uille, Xavier},\n"
            "  journal={American Economic Review},\n"
            "  volume={110},\n"
            "  number={9},\n"
            "  pages={2964--2996},\n"
            "  year={2020}\n"
            "}"
        ),
    }

    def __init__(
        self,
        method: str,
        estimand: str,
        estimate: float,
        se: float,
        pvalue: float,
        ci: tuple,
        alpha: float,
        n_obs: int,
        detail: Optional[pd.DataFrame] = None,
        model_info: Optional[Dict[str, Any]] = None,
        _influence_funcs: Optional[np.ndarray] = None,
        _citation_key: Optional[str] = None,
    ):
        self.method = method
        self.estimand = estimand
        self.estimate = estimate
        self.se = se
        self.pvalue = pvalue
        self.ci = ci
        self.alpha = alpha
        self.n_obs = n_obs
        self.detail = detail
        self.model_info = model_info or {}
        self._influence_funcs = _influence_funcs
        self._citation_key = _citation_key

    # ------------------------------------------------------------------
    # Backward compatibility with EconometricResults
    # ------------------------------------------------------------------

    @property
    def params(self) -> pd.Series:
        """Treatment effect as a params Series (for outreg2 compatibility)."""
        return pd.Series({self.estimand: self.estimate})

    @property
    def std_errors(self) -> pd.Series:
        return pd.Series({self.estimand: self.se})

    @property
    def tvalues(self) -> pd.Series:
        t = self.estimate / self.se if self.se > 0 else np.nan
        return pd.Series({self.estimand: t})

    @property
    def pvalues(self) -> pd.Series:
        return pd.Series({self.estimand: self.pvalue})

    @property
    def diagnostics(self) -> Dict[str, Any]:
        return self.model_info

    @property
    def data_info(self) -> Dict[str, Any]:
        return {'nobs': self.n_obs}

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _stars(pvalue: float) -> str:
        """Significance stars."""
        if pd.isna(pvalue):
            return ""
        if pvalue < 0.01:
            return "***"
        if pvalue < 0.05:
            return "**"
        if pvalue < 0.1:
            return "*"
        return ""

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------

    def summary(self, alpha: Optional[float] = None) -> str:
        """
        Generate a formatted text summary of the causal estimation results.

        Parameters
        ----------
        alpha : float, optional
            Override significance level for display.

        Returns
        -------
        str
        """
        alpha = alpha or self.alpha
        lines: List[str] = []

        lines.append("=" * 78)
        lines.append(f"  {self.method}")
        lines.append("=" * 78)
        lines.append("")

        stars = self._stars(self.pvalue)
        lines.append(f"  {self.estimand}:      {self.estimate: .6f} {stars}")
        lines.append(f"  Std. Error:  ({self.se:.6f})")
        pct = int(100 * (1 - alpha))
        lines.append(f"  [{pct}% CI]:    [{self.ci[0]:.6f},  {self.ci[1]:.6f}]")
        lines.append(f"  P-value:     {self.pvalue:.4f}")
        lines.append("")

        # Event study coefficients
        if 'event_study' in self.model_info:
            es = self.model_info['event_study']
            lines.append("-" * 78)
            lines.append("  Event Study Coefficients")
            lines.append("-" * 78)
            for _, row in es.iterrows():
                e = int(row['relative_time'])
                att = row['att']
                se_v = row['se']
                pv = row.get('pvalue', np.nan)
                s = self._stars(pv)
                lines.append(
                    f"  e = {e:>3d}  |  {att:>10.4f}  ({se_v:.4f}) {s}"
                )
            lines.append("")

        # Detailed estimates
        if self.detail is not None and len(self.detail) > 0:
            if 'att' in self.detail.columns:
                # Causal inference format (group-time ATTs)
                cols = [c for c in ['group', 'time', 'att', 'se',
                                    'ci_lower', 'ci_upper', 'pvalue']
                        if c in self.detail.columns]
                title_str = "Group-Time ATT Estimates"
            elif 'method' in self.detail.columns and 'estimate' in self.detail.columns:
                # RD-style inference table
                cols = [c for c in ['method', 'estimate', 'se', 'z',
                                    'pvalue', 'ci_lower', 'ci_upper']
                        if c in self.detail.columns]
                title_str = "Inference"
            elif 'coefficient' in self.detail.columns:
                # Regression format (variable / coefficient / se)
                cols = [c for c in ['variable', 'coefficient', 'se',
                                    'tstat', 'pvalue']
                        if c in self.detail.columns]
                title_str = "Regression Coefficients"
            else:
                cols = list(self.detail.columns)
                title_str = "Detailed Estimates"
            lines.append("-" * 78)
            lines.append(f"  {title_str}")
            lines.append("-" * 78)
            lines.append(
                self.detail[cols].to_string(index=False, float_format='%.4f')
            )
            lines.append("")

        # Pre-trend test
        if 'pretrend_test' in self.model_info:
            pt = self.model_info['pretrend_test']
            lines.append("-" * 78)
            lines.append(
                f"  Pre-trend Test: chi2({pt['df']}) = {pt['statistic']:.4f}, "
                f"p-value = {pt['pvalue']:.4f}"
            )
            lines.append("")

        # Model info footer
        lines.append("-" * 78)
        lines.append(f"  Observations:    {self.n_obs:,}")
        _skip = {'event_study', 'pretrend_test', 'aggregations',
                 'cohort_sizes', 'influence_funcs_matrix',
                 'conventional', 'robust'}
        for key, val in self.model_info.items():
            if key in _skip or isinstance(val, (pd.DataFrame, np.ndarray,
                                                dict, list)):
                continue
            label = key.replace('_', ' ').title()
            lines.append(f"  {label}:    {val}")
        lines.append("=" * 78)
        lines.append("  * p<0.1, ** p<0.05, *** p<0.01")

        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Plotting
    # ------------------------------------------------------------------

    def event_study_plot(
        self,
        ax=None,
        title: Optional[str] = None,
        color: str = '#2C3E50',
        ci_alpha: float = 0.15,
        figsize: tuple = (10, 6),
        **kwargs,
    ):
        """
        Plot event study coefficients with confidence intervals.

        Parameters
        ----------
        ax : matplotlib.axes.Axes, optional
        title : str, optional
        color : str
        ci_alpha : float
        figsize : tuple

        Returns
        -------
        (fig, ax)
        """
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            raise ImportError(
                "matplotlib required for plotting. "
                "Install: pip install matplotlib"
            )

        if 'event_study' not in self.model_info:
            raise ValueError(
                "No event study estimates. Use a staggered DID method."
            )

        es = self.model_info['event_study'].copy()

        if ax is None:
            fig, ax = plt.subplots(figsize=figsize)
        else:
            fig = ax.get_figure()

        e = es['relative_time'].values
        att = es['att'].values
        lo = es['ci_lower'].values
        hi = es['ci_upper'].values

        ax.fill_between(e, lo, hi, alpha=ci_alpha, color=color)
        ax.scatter(e, att, color=color, s=40, zorder=5)
        ax.plot(e, att, color=color, linewidth=1, alpha=0.7, zorder=4)
        ax.errorbar(
            e, att,
            yerr=[att - lo, hi - att],
            fmt='none', color=color, capsize=3, linewidth=1, zorder=3,
        )

        ax.axhline(y=0, color='gray', linestyle='--', linewidth=0.8)
        ax.axvline(
            x=-0.5, color='#E74C3C', linestyle=':',
            linewidth=1, alpha=0.5, label='Treatment onset',
        )

        ax.set_xlabel('Periods Relative to Treatment', fontsize=11)
        ax.set_ylabel('Estimated Effect', fontsize=11)
        ax.set_title(title or f'Event Study: {self.method}', fontsize=13)
        ax.tick_params(labelsize=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend(fontsize=9, frameon=False)
        fig.tight_layout()
        return fig, ax

    def plot(self, type: str = 'auto', **kwargs):
        """
        Generate appropriate visualisation based on model type.

        Parameters
        ----------
        type : str
            'auto' (recommended), 'event_study', 'coefplot',
            'trajectory', 'gap', 'both', 'weights', 'placebo',
            'placebo_gap', 'conformal', 'staggered', 'factors',
            'balance', 'density', 'original', 'pointwise',
            'cumulative', 'all'.

        Returns
        -------
        (fig, ax) or (fig, axes)

        Notes
        -----
        In 'auto' mode, the plot type is selected by method:

        - DID/event study → event study plot
        - Synthetic Control (all variants) → trajectory + gap (both)
        - Conformal SCM → conformal period-level CI plot
        - Staggered SCM → cohort-level ATT comparison
        - Causal Impact → 3-panel (original + pointwise + cumulative)
        - Matching → Love plot (covariate balance)
        - RD → coefplot (use ``rdplot()`` for binned scatter)
        - Other → coefplot
        """
        method_lower = self.method.lower()

        # All synth-related plot types handled by unified synthplot
        _synth_types = {
            'trajectory', 'gap', 'both', 'weights', 'placebo',
            'placebo_gap', 'placebo_dist', 'conformal', 'staggered',
            'factors', 'loadings', 'compare',
        }

        if type == 'auto':
            # Event study (DID)
            if 'event_study' in self.model_info:
                return self.event_study_plot(**kwargs)

            # Synthetic Control — detect ALL variants
            if self._is_synth_result():
                from ..synth.plots import synthplot
                # Pick best auto type based on variant
                if 'period_results' in self.model_info:
                    return synthplot(self, type='conformal', **kwargs)
                if 'cohort_effects' in self.model_info:
                    return synthplot(self, type='staggered', **kwargs)
                return synthplot(self, type='both', **kwargs)

            # Causal Impact
            if 'causal impact' in method_lower or \
               'intervention_time' in self.model_info:
                from ..causal_impact.impact import impactplot
                return impactplot(self, type='all', **kwargs)

            # Matching (has balance/SMD table)
            if self.detail is not None and 'smd' in getattr(
                    self.detail, 'columns', []):
                from ..matching.match import balanceplot
                return balanceplot(self, **kwargs)

            return self._coefplot(**kwargs)

        # Explicit type overrides
        if type == 'event_study':
            return self.event_study_plot(**kwargs)
        if type in _synth_types:
            from ..synth.plots import synthplot
            return synthplot(self, type=type, **kwargs)
        if type in ('original', 'pointwise', 'cumulative', 'all'):
            from ..causal_impact.impact import impactplot
            return impactplot(self, type=type, **kwargs)
        if type == 'balance':
            from ..matching.match import balanceplot
            return balanceplot(self, **kwargs)
        return self._coefplot(**kwargs)

    def _is_synth_result(self) -> bool:
        """Check if this result is from any synthetic control variant."""
        mi = self.model_info
        # Direct markers from various synth variants
        synth_keys = {
            'gap_table',        # classic, demeaned, robust
            'Y_obs',            # sdid
            'trajectory',       # gsynth
            'factors_pre',      # gsynth
            'cohort_effects',   # staggered
            'period_results',   # conformal
        }
        if synth_keys & set(mi.keys()):
            return True
        # Augmented SCM
        if mi.get('model_type', '').startswith('Synthetic'):
            return True
        # Method name check
        m = self.method.lower()
        return any(kw in m for kw in (
            'synthetic', 'synth', 'sdid', 'gsynth', 'staggered',
            'conformal', 'augmented', 'ascm', 'demeaned', 'de-meaned',
            'de-trended', 'unconstrained', 'factor',
        ))

    def _coefplot(self, ax=None, figsize=(8, 5), **kwargs):
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            raise ImportError("matplotlib required for plotting")

        if ax is None:
            fig, ax = plt.subplots(figsize=figsize)
        else:
            fig = ax.get_figure()

        ax.errorbar(
            0, self.estimate,
            yerr=[[self.estimate - self.ci[0]], [self.ci[1] - self.estimate]],
            fmt='o', color='#2C3E50', capsize=5, markersize=8,
        )
        ax.axhline(y=0, color='gray', linestyle='--', linewidth=0.8)
        ax.set_xlim(-1, 1)
        ax.set_xticks([0])
        ax.set_xticklabels([self.estimand])
        ax.set_ylabel('Estimated Effect')
        ax.set_title(f'{self.method}: {self.estimand}')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        fig.tight_layout()
        return fig, ax

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    def to_latex(self, caption: Optional[str] = None,
                 label: Optional[str] = None) -> str:
        """Generate a LaTeX table of the results."""
        caption = caption or f'{self.method} Results'
        label = label or 'tab:causal_result'

        lines = [
            '\\begin{table}[htbp]',
            '\\centering',
            f'\\caption{{{caption}}}',
            f'\\label{{{label}}}',
        ]

        if self.detail is not None and len(self.detail) > 0:
            # Detect table format: causal (group/time/att) vs regression (variable/coefficient)
            if 'att' in self.detail.columns:
                cols = [c for c in ['group', 'time', 'att', 'se', 'pvalue']
                        if c in self.detail.columns]
                coef_col, star_col = 'att', 'att'
            elif 'method' in self.detail.columns and 'estimate' in self.detail.columns:
                cols = [c for c in ['method', 'estimate', 'se', 'pvalue']
                        if c in self.detail.columns]
                coef_col, star_col = 'estimate', 'estimate'
            elif 'coefficient' in self.detail.columns:
                cols = [c for c in ['variable', 'coefficient', 'se', 'pvalue']
                        if c in self.detail.columns]
                coef_col, star_col = 'coefficient', 'coefficient'
            else:
                cols = list(self.detail.columns)
                coef_col, star_col = None, None

            n_cols = len(cols)
            spec = 'l' + 'c' * (n_cols - 1)
            lines.append(f'\\begin{{tabular}}{{{spec}}}')
            lines.append('\\hline\\hline')
            hdr = {'group': 'Group', 'time': 'Time', 'att': 'ATT',
                   'se': 'Std.\\ Error', 'pvalue': 'P-value',
                   'variable': 'Variable', 'coefficient': 'Coefficient',
                   'tstat': 't-stat'}
            lines.append(' & '.join(hdr.get(c, c) for c in cols) + ' \\\\')
            lines.append('\\hline')
            for _, row in self.detail.iterrows():
                vals = []
                for c in cols:
                    v = row[c]
                    if isinstance(v, float):
                        s = self._stars(row.get('pvalue', np.nan)) if c == star_col else ''
                        vals.append(f'{v:.4f}{s}')
                    else:
                        vals.append(str(int(v)) if isinstance(v, (int, np.integer)) else str(v))
                lines.append(' & '.join(vals) + ' \\\\')
        else:
            lines.append('\\begin{tabular}{lc}')
            lines.append('\\hline\\hline')
            lines.append(
                f'{self.estimand} & '
                f'{self.estimate:.4f}{self._stars(self.pvalue)} \\\\'
            )
            lines.append(f'& ({self.se:.4f}) \\\\')

        lines += [
            '\\hline',
            f'Observations & {self.n_obs:,} \\\\',
            '\\hline\\hline',
            '\\end{tabular}',
            '\\begin{tablenotes}',
            '\\footnotesize',
            '\\item Standard errors in parentheses.',
            '\\item * p<0.1, ** p<0.05, *** p<0.01',
            '\\end{tablenotes}',
            '\\end{table}',
        ]
        return '\n'.join(lines)

    def cite(self) -> str:
        """Return BibTeX citation for the method."""
        key = self._citation_key or self.method.lower().replace(' ', '_')
        if key in self._CITATIONS:
            return self._CITATIONS[key]
        for k, v in self._CITATIONS.items():
            if k in key or key in k:
                return v
        return f"% No citation registered for method: {self.method}"

    def pretrend_test(self) -> Dict[str, Any]:
        """Return pre-trend test results (DID methods)."""
        if 'pretrend_test' not in self.model_info:
            raise ValueError("Pre-trend test not available for this method.")
        return self.model_info['pretrend_test']

    def to_docx(self, filename: str, title: Optional[str] = None):
        """
        Export results to a Word (.docx) document.

        Parameters
        ----------
        filename : str
            Output path (.docx).
        title : str, optional
            Table title. Defaults to method name.
        """
        _result_to_docx(self, filename, title)

    def _repr_html_(self) -> str:
        """Rich HTML display for Jupyter notebooks — model-specific layouts."""
        mi = self.model_info
        pct = int(100 * (1 - self.alpha))
        stars_raw = self._stars(self.pvalue)

        def _s(pv):
            if pd.isna(pv): return ''
            if pv < 0.01: return '<span style="color:#DC2626">***</span>'
            if pv < 0.05: return '<span style="color:#EA580C">**</span>'
            if pv < 0.1: return '<span style="color:#D97706">*</span>'
            return ''

        # Significance-based accent color
        if self.pvalue < 0.01:
            accent, accent_bg = '#059669', '#ECFDF5'  # green
        elif self.pvalue < 0.05:
            accent, accent_bg = '#2563EB', '#EFF6FF'  # blue
        elif self.pvalue < 0.1:
            accent, accent_bg = '#D97706', '#FFFBEB'  # amber
        else:
            accent, accent_bg = '#64748B', '#F8FAFC'  # gray

        # Shared CSS (scoped)
        S = ('<style scoped>'
             '.sp-box{font-family:"Helvetica Neue",Arial,sans-serif;max-width:720px;'
             'border:1px solid #E5E7EB;border-radius:8px;overflow:hidden;margin:6px 0}'
             '.sp-hdr{background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;padding:12px 16px}'
             '.sp-hdr h3{margin:0;font-size:15px;font-weight:600;letter-spacing:0.3px}'
             '.sp-hdr .sp-sub{font-size:11px;color:#94A3B8;margin-top:2px}'
             '.sp-effect{display:flex;align-items:center;gap:16px;padding:14px 16px;border-bottom:1px solid #E5E7EB}'
             '.sp-effect-num{font-size:28px;font-weight:700;letter-spacing:-0.5px}'
             '.sp-effect-meta{font-size:12px;color:#64748B;line-height:1.6}'
             '.sp-effect-badge{display:inline-block;padding:2px 8px;border-radius:10px;font-size:10px;font-weight:600;letter-spacing:0.3px}'
             '.sp-metrics{display:flex;gap:0;border-bottom:1px solid #E5E7EB}'
             '.sp-metric{flex:1;padding:8px 12px;text-align:center;border-right:1px solid #E5E7EB}'
             '.sp-metric:last-child{border-right:none}'
             '.sp-metric .sp-val{font-size:15px;font-weight:700;color:#1a1a2e}'
             '.sp-metric .sp-lab{font-size:9px;color:#94A3B8;text-transform:uppercase;letter-spacing:0.5px;margin-top:1px}'
             '.sp-section{border-top:1px solid #E5E7EB}'
             '.sp-section summary{padding:8px 14px;font-size:12px;font-weight:600;color:#1a1a2e;cursor:pointer;'
             'list-style:none;display:flex;align-items:center;gap:6px}'
             '.sp-section summary::before{content:"\\25B6";font-size:8px;color:#94A3B8;transition:transform 0.2s}'
             '.sp-section[open] summary::before{transform:rotate(90deg)}'
             'table.sp-tbl{width:100%;border-collapse:collapse;font-size:11px}'
             'table.sp-tbl th{padding:4px 10px;text-align:right;font-weight:600;color:#64748B;border-bottom:1px solid #E5E7EB;font-size:10px}'
             'table.sp-tbl th:first-child{text-align:left}'
             'table.sp-tbl td{padding:4px 10px;text-align:right;border-bottom:1px solid #F8FAFC}'
             'table.sp-tbl td:first-child{text-align:left;font-weight:500;color:#334155}'
             'table.sp-tbl tr:hover{background:#F8FAFC}'
             '.sp-grid{display:grid;grid-template-columns:1fr 1fr;gap:0;font-size:11px}'
             '.sp-grid-item{padding:4px 14px;display:flex;justify-content:space-between;border-bottom:1px solid #F8FAFC}'
             '.sp-grid-item:nth-child(odd){border-right:1px solid #F1F5F9}'
             '.sp-gk{color:#94A3B8}.sp-gv{color:#334155;font-weight:500}'
             '.sp-bar{height:6px;border-radius:3px;margin-top:2px}'
             '.sp-foot{padding:6px 14px;font-size:10px;color:#94A3B8;border-top:1px solid #E5E7EB;display:flex;justify-content:space-between}'
             '</style>')

        h = [S, '<div class="sp-box">']

        # ── Header ──
        h.append(f'<div class="sp-hdr"><h3>{self.method}</h3>')
        sub = f'{self.estimand}'
        if mi.get('rd_type'):
            sub += f' · {mi["rd_type"]} RD'
        elif mi.get('distance'):
            sub += f' · {mi["distance"]} {mi.get("method", "")}'
        h.append(f'<div class="sp-sub">{sub}</div></div>')

        # ── Main Effect Card ──
        sig_label = '< 0.01' if self.pvalue < 0.01 else ('< 0.05' if self.pvalue < 0.05 else ('< 0.10' if self.pvalue < 0.1 else f'= {self.pvalue:.3f}'))
        h.append(f'<div class="sp-effect" style="background:{accent_bg};">')
        h.append(f'<div class="sp-effect-num" style="color:{accent};">{self.estimate:.4f}</div>')
        h.append(f'<div class="sp-effect-meta">')
        h.append(f'<span class="sp-effect-badge" style="background:{accent};color:white;">{stars_raw or "n.s."}</span> &nbsp; p {sig_label}<br>')
        h.append(f'SE = {self.se:.4f} &nbsp;&nbsp; {pct}% CI [{self.ci[0]:.4f}, {self.ci[1]:.4f}]')
        h.append(f'</div></div>')

        # ── Model-Specific Metric Bars ──
        h.append('<div class="sp-metrics">')
        h.append(f'<div class="sp-metric"><div class="sp-val">{self.n_obs:,}</div><div class="sp-lab">Observations</div></div>')

        if self._is_synth_result():
            # SC metrics
            for key, label in [('n_donors', 'Donors'), ('n_pre_periods', 'Pre-periods'), ('n_post_periods', 'Post-periods')]:
                if key in mi:
                    h.append(f'<div class="sp-metric"><div class="sp-val">{mi[key]}</div><div class="sp-lab">{label}</div></div>')
            if 'pre_treatment_rmse' in mi:
                h.append(f'<div class="sp-metric"><div class="sp-val">{mi["pre_treatment_rmse"]:.3f}</div><div class="sp-lab">Pre-RMSE</div></div>')
        elif mi.get('rd_type') is not None:
            # RD metrics
            for key, label in [('n_effective_left', 'N Left (eff.)'), ('n_effective_right', 'N Right (eff.)'),
                               ('bandwidth_h', 'Bandwidth')]:
                if key in mi:
                    v = mi[key]
                    vs = f'{v:.3f}' if isinstance(v, float) else str(v)
                    h.append(f'<div class="sp-metric"><div class="sp-val">{vs}</div><div class="sp-lab">{label}</div></div>')
        elif self.detail is not None and 'smd' in getattr(self.detail, 'columns', []):
            # Matching metrics
            for key, label in [('n_treated', 'Treated'), ('n_control', 'Control'), ('n_matches', 'Matches')]:
                if key in mi:
                    h.append(f'<div class="sp-metric"><div class="sp-val">{mi[key]}</div><div class="sp-lab">{label}</div></div>')
        h.append('</div>')

        # ── SC: Donor Weights ──
        if self._is_synth_result() and self.detail is not None and 'weight' in getattr(self.detail, 'columns', []):
            weights_df = self.detail[self.detail['weight'] > 0.001].sort_values('weight', ascending=False)
            if len(weights_df) > 0:
                max_w = weights_df['weight'].max()
                h.append('<details class="sp-section" open><summary>Donor Weights</summary>')
                h.append('<div style="padding:4px 14px 8px;">')
                for _, row in weights_df.iterrows():
                    unit_name = row.get('unit', row.iloc[0])
                    w = row['weight']
                    pct_w = (w / max_w) * 100
                    h.append(f'<div style="display:flex;align-items:center;gap:8px;margin:3px 0;font-size:11px;">'
                             f'<span style="width:60px;color:#334155;font-weight:500;">Unit {unit_name}</span>'
                             f'<div style="flex:1;background:#F1F5F9;border-radius:3px;height:8px;">'
                             f'<div class="sp-bar" style="width:{pct_w:.0f}%;background:{accent};"></div></div>'
                             f'<span style="width:50px;text-align:right;color:#64748B;">{w:.3f}</span></div>')
                h.append('</div></details>')

        # ── SC: Gap Table ──
        if 'gap_table' in mi and isinstance(mi['gap_table'], pd.DataFrame):
            gap = mi['gap_table']
            h.append('<details class="sp-section"><summary>Period-by-Period Effects</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for c in gap.columns:
                h.append(f'<th>{c}</th>')
            h.append('</tr>')
            for _, row in gap.iterrows():
                h.append('<tr>')
                for c in gap.columns:
                    v = row[c]
                    h.append(f'<td>{v:.4f}</td>' if isinstance(v, float) else f'<td>{v}</td>')
                h.append('</tr>')
            h.append('</table></details>')

        # ── RD: Conventional vs Robust Inference ──
        if self.detail is not None and 'method' in getattr(self.detail, 'columns', []) and 'estimate' in getattr(self.detail, 'columns', []):
            h.append('<details class="sp-section" open><summary>Inference Comparison</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for col in ['Method', 'Estimate', 'Std. Err.', 'z', 'p-value', 'CI']:
                h.append(f'<th>{col}</th>')
            h.append('</tr>')
            for _, row in self.detail.iterrows():
                meth = row.get('method', '')
                est = row.get('estimate', np.nan)
                se_v = row.get('se', np.nan)
                z_v = row.get('z', np.nan)
                pv = row.get('pvalue', np.nan)
                lo = row.get('ci_lower', np.nan)
                hi = row.get('ci_upper', np.nan)
                pvc = '#DC2626' if pv < 0.01 else ('#EA580C' if pv < 0.05 else ('#D97706' if pv < 0.1 else '#64748B'))
                bold = 'font-weight:600;' if 'Robust' in str(meth) else ''
                h.append(f'<tr style="{bold}">')
                h.append(f'<td>{meth}</td><td>{est:.4f} {_s(pv)}</td>')
                h.append(f'<td style="color:#94A3B8;">({se_v:.4f})</td><td>{z_v:.2f}</td>')
                h.append(f'<td style="color:{pvc};font-weight:600;">{pv:.4f}</td>')
                h.append(f'<td style="color:#94A3B8;">[{lo:.4f}, {hi:.4f}]</td></tr>')
            h.append('</table></details>')

        # ── RD: Design Parameters ──
        if mi.get('rd_type') is not None:
            h.append('<details class="sp-section"><summary>Design Parameters</summary><div class="sp-grid">')
            rd_params = [('cutoff', 'Cutoff'), ('polynomial_p', 'Poly Order (p)'), ('polynomial_q', 'Bias Poly (q)'),
                         ('kernel', 'Kernel'), ('bwselect', 'BW Selection'), ('bandwidth_h', 'Bandwidth (h)'),
                         ('bandwidth_b', 'Bias BW (b)'), ('n_left', 'N Left'), ('n_right', 'N Right'),
                         ('n_effective_left', 'N Eff. Left'), ('n_effective_right', 'N Eff. Right')]
            for key, label in rd_params:
                if key in mi:
                    v = mi[key]
                    vs = f'{v:.4f}' if isinstance(v, float) else str(v)
                    h.append(f'<div class="sp-grid-item"><span class="sp-gk">{label}</span><span class="sp-gv">{vs}</span></div>')
            h.append('</div></details>')

        # ── Matching: Covariate Balance ──
        if self.detail is not None and 'smd' in getattr(self.detail, 'columns', []):
            h.append('<details class="sp-section" open><summary>Covariate Balance</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for col in ['Variable', 'Treated', 'Control', 'SMD', '']:
                h.append(f'<th>{col}</th>')
            h.append('</tr>')
            for _, row in self.detail.iterrows():
                var = row.get('variable', row.iloc[0])
                mt = row.get('mean_treated', np.nan)
                mc = row.get('mean_control', np.nan)
                smd = row.get('smd', np.nan)
                smd_abs = abs(smd) if not pd.isna(smd) else 0
                bar_color = '#059669' if smd_abs < 0.1 else ('#D97706' if smd_abs < 0.25 else '#DC2626')
                bar_w = min(smd_abs / 0.5 * 100, 100)
                h.append(f'<tr><td>{var}</td>')
                h.append(f'<td>{mt:.2f}</td><td>{mc:.2f}</td>')
                h.append(f'<td style="color:{bar_color};font-weight:600;">{smd:.3f}</td>')
                h.append(f'<td style="width:80px;"><div style="background:#F1F5F9;border-radius:3px;height:6px;">'
                         f'<div style="width:{bar_w:.0f}%;height:6px;border-radius:3px;background:{bar_color};"></div>'
                         f'</div></td></tr>')
            h.append('</table>')
            # Balance threshold annotation
            h.append('<div style="padding:4px 10px;font-size:10px;color:#94A3B8;">'
                     'SMD: <span style="color:#059669">|d|&lt;0.1 balanced</span> · '
                     '<span style="color:#D97706">0.1-0.25 borderline</span> · '
                     '<span style="color:#DC2626">&gt;0.25 imbalanced</span></div>')
            h.append('</details>')

        # ── Matching: Design info ──
        if mi.get('distance') and 'smd' in getattr(self.detail, 'columns', []) if self.detail is not None else False:
            h.append('<details class="sp-section"><summary>Matching Parameters</summary><div class="sp-grid">')
            match_params = [('distance', 'Distance'), ('method', 'Method'), ('estimand', 'Estimand'),
                            ('n_treated', 'N Treated'), ('n_control', 'N Control'), ('n_matches', 'Matches'),
                            ('caliper', 'Caliper'), ('replace', 'With Replacement'), ('bias_correction', 'Bias Correction')]
            for key, label in match_params:
                if key in mi:
                    v = mi[key]
                    vs = f'{v:.4f}' if isinstance(v, float) else str(v)
                    h.append(f'<div class="sp-grid-item"><span class="sp-gk">{label}</span><span class="sp-gv">{vs}</span></div>')
            h.append('</div></details>')

        # ── Event study coefficients ──
        if 'event_study' in mi:
            es = mi['event_study']
            h.append('<details class="sp-section" open><summary>Event Study Coefficients</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for col in ['Period', 'ATT', 'Std. Err.', 'CI', 'p-value']:
                h.append(f'<th>{col}</th>')
            h.append('</tr>')
            for _, row in es.iterrows():
                e = int(row['relative_time'])
                att = row['att']
                se_v = row['se']
                pv = row.get('pvalue', np.nan)
                lo = row.get('ci_lower', np.nan)
                hi = row.get('ci_upper', np.nan)
                pvc = '#DC2626' if pv < 0.01 else ('#EA580C' if pv < 0.05 else ('#D97706' if pv < 0.1 else '#64748B'))
                bg = 'background:#FFFBEB;' if e == 0 else ''
                h.append(f'<tr style="{bg}">')
                h.append(f'<td style="font-weight:600;">e = {e}</td>')
                h.append(f'<td>{att:.4f} {_s(pv)}</td>')
                h.append(f'<td style="color:#94A3B8;">({se_v:.4f})</td>')
                ci_str = f'[{lo:.4f}, {hi:.4f}]' if not pd.isna(lo) else ''
                h.append(f'<td style="color:#94A3B8;">{ci_str}</td>')
                pv_str = f'{pv:.4f}' if not pd.isna(pv) else ''
                h.append(f'<td style="color:{pvc};font-weight:600;">{pv_str}</td></tr>')
            h.append('</table></details>')

        # ── Generic detail table (group-time ATTs, etc.) ──
        if (self.detail is not None and len(self.detail) > 0
                and 'event_study' not in mi
                and 'smd' not in getattr(self.detail, 'columns', [])
                and not ('method' in getattr(self.detail, 'columns', []) and 'estimate' in getattr(self.detail, 'columns', []))
                and not ('weight' in getattr(self.detail, 'columns', []) and self._is_synth_result())):
            if 'att' in self.detail.columns:
                cols = [c for c in ['group', 'time', 'att', 'se', 'pvalue'] if c in self.detail.columns]
                title_str = 'Group-Time ATT Estimates'
            else:
                cols = list(self.detail.columns)
                title_str = 'Detail'
            h.append(f'<details class="sp-section"><summary>{title_str}</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for c in cols:
                h.append(f'<th>{c}</th>')
            h.append('</tr>')
            max_rows = 20
            for _, row in self.detail[cols].head(max_rows).iterrows():
                h.append('<tr>')
                for c in cols:
                    v = row[c]
                    h.append(f'<td>{v:.4f}</td>' if isinstance(v, float) else f'<td>{v}</td>')
                h.append('</tr>')
            if len(self.detail) > max_rows:
                h.append(f'<tr><td colspan="{len(cols)}" style="text-align:center;color:#94A3B8;">... {len(self.detail)-max_rows} more</td></tr>')
            h.append('</table></details>')

        # ── SC: Model metadata ──
        if self._is_synth_result():
            h.append('<details class="sp-section"><summary>Model Parameters</summary><div class="sp-grid">')
            sc_params = [('treatment_time', 'Treatment Time'), ('treated_unit', 'Treated Unit'),
                         ('penalization', 'Penalization'), ('pre_treatment_mspe', 'Pre-MSPE'),
                         ('pre_treatment_rmse', 'Pre-RMSE'), ('n_placebos', 'Placebo Tests')]
            for key, label in sc_params:
                if key in mi:
                    v = mi[key]
                    vs = f'{v:.4f}' if isinstance(v, float) else str(v)
                    h.append(f'<div class="sp-grid-item"><span class="sp-gk">{label}</span><span class="sp-gv">{vs}</span></div>')
            h.append('</div></details>')

        # ── Footer ──
        h.append(f'<div class="sp-foot"><span>N = {self.n_obs:,}</span><span>* p&lt;0.1 &nbsp; ** p&lt;0.05 &nbsp; *** p&lt;0.01</span></div>')
        h.append('</div>')
        return '\n'.join(h)

    def __repr__(self) -> str:
        s = self._stars(self.pvalue)
        return (
            f"<CausalResult: {self.method}, {self.estimand} = "
            f"{self.estimate:.4f}{s}, SE = {self.se:.4f}, n = {self.n_obs:,}>"
        )

    def __str__(self) -> str:
        return self.summary()


# ======================================================================
# Shared Word export helper
# ======================================================================

def _result_to_docx(result, filename: str, title: Optional[str] = None):
    """
    Export a single EconometricResults or CausalResult to Word (.docx).

    Produces a one-model table with coefficients, SEs, stars, and
    diagnostics in APA format.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx required for Word export. "
            "Install: pip install python-docx"
        )

    if not filename.endswith('.docx'):
        filename += '.docx'

    doc = Document()

    # Title
    if title is None:
        title = getattr(result, 'method', None) or result.model_info.get('model_type', 'Results')
    p = doc.add_paragraph()
    run = p.add_run(str(title))
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build rows from result
    params = result.params
    std_errors = result.std_errors
    pvalues = result.pvalues if hasattr(result, 'pvalues') else None

    def _stars(pv):
        if pv is None or (isinstance(pv, float) and np.isnan(pv)):
            return ''
        if pv < 0.01: return '***'
        if pv < 0.05: return '**'
        if pv < 0.1: return '*'
        return ''

    # Table: Variable | Coefficient | SE
    rows_data = []
    if isinstance(params, pd.Series):
        for var in params.index:
            coef = params[var]
            se = std_errors[var] if var in std_errors.index else np.nan
            if pvalues is not None and isinstance(pvalues, pd.Series) and var in pvalues.index:
                pv = float(pvalues[var])
            else:
                pv = np.nan
            rows_data.append((str(var), f'{coef:.4f}{_stars(pv)}', f'({se:.4f})'))

    n_rows = len(rows_data) * 2 + 1  # coef rows + SE rows + header
    # Actually: each variable gets 2 rows (coef, SE)
    table = doc.add_table(rows=1 + len(rows_data) * 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    table.rows[0].cells[0].text = ''
    table.rows[0].cells[1].text = title
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data
    row_idx = 1
    for var, coef_str, se_str in rows_data:
        # Coefficient row
        table.rows[row_idx].cells[0].text = var
        table.rows[row_idx].cells[1].text = coef_str
        for cell in table.rows[row_idx].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell == table.rows[row_idx].cells[1] else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
        row_idx += 1

        # SE row
        table.rows[row_idx].cells[0].text = ''
        table.rows[row_idx].cells[1].text = se_str
        for para in table.rows[row_idx].cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
        row_idx += 1

    # Diagnostics as additional paragraph
    diag = getattr(result, 'diagnostics', {})
    n_obs = None
    if hasattr(result, 'data_info') and isinstance(result.data_info, dict):
        n_obs = result.data_info.get('nobs')
    if hasattr(result, 'n_obs'):
        n_obs = result.n_obs

    diag_lines = []
    if n_obs is not None:
        diag_lines.append(f'Observations: {n_obs:,}')
    if isinstance(diag, dict):
        for k, v in diag.items():
            if isinstance(v, (int, float)) and not isinstance(v, bool):
                diag_lines.append(f'{k}: {v:.4f}' if isinstance(v, float) else f'{k}: {v:,}')

    if diag_lines:
        dp = doc.add_paragraph()
        for line in diag_lines:
            run = dp.add_run(line + '\n')
            run.font.size = Pt(8)

    # Significance note
    np_ = doc.add_paragraph()
    run = np_.add_run('* p<0.1, ** p<0.05, *** p<0.01')
    run.italic = True
    run.font.size = Pt(8)

    doc.save(filename)
