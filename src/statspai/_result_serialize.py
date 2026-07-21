"""Shared agent-native result protocol for domain result dataclasses.

Design principle 3 ("统一结果对象") asks every result object to be usable
through one entry point by a human *and* an agent. The flagship
`CausalResult` / `EconometricResults` already expose `.to_dict()` /
`.to_latex()` / `.cite()`; the lighter domain result dataclasses (negative
controls, proximal, four-way mediation, network spillover, ITS, Rosenbaum
bounds, longitudinal TMLE, QTE, robustness, transport, … and the BCF / MR
extensions) historically only had `.summary()`.

`ResultProtocolMixin` gives any such dataclass the full export protocol in one
line of inheritance, reusing the existing `core.results._to_jsonable` converter
so the serialisation logic lives in exactly one place (CLAUDE.md §4):

* ``to_dict()``     — JSON-safe ``{field: value}`` (numpy / pandas / NaN aware).
* ``to_latex()``    — a compact ``booktabs`` table of the scalar fields.
* ``to_markdown()`` — the same table as GitHub-flavoured markdown.
* ``to_excel()``    — a two-column Field/Value ``.xlsx`` sheet (openpyxl).
* ``to_word()``     — the class's own ``to_docx`` when defined, else a
  Field/Value ``.docx`` table (python-docx).
* ``cite()``     — the **verified** paper.bib key(s) the estimator is based on
  (set via the ``_citation_keys`` class attribute), or an honest placeholder.
  Zero-hallucination (CLAUDE.md §10): keys are pointers into ``paper.bib`` — the
  single source of truth — never a generated citation string. Resolve a key to
  full BibTeX with ``sp.bibtex(keys=[...])``.

Subclass methods always win: the mixin's defaults only run when the host class
does not define the method itself (ordinary Python MRO).
"""

from __future__ import annotations

import warnings
from dataclasses import fields, is_dataclass
from typing import Any, ClassVar, Dict, Tuple

from .core.results import _to_jsonable


def result_to_dict(obj: Any) -> Dict[str, Any]:
    """Return a JSON-safe ``{field: value}`` mapping for a result object.

    Every field is passed through ``core.results._to_jsonable`` so numpy
    scalars/arrays, pandas Series/DataFrames and NaN/Inf are converted to
    JSON-friendly values (``json.dumps`` never raises on the result).

    Dataclasses use their declared fields; plain result classes fall back
    to their public, non-callable instance attributes **plus** the public
    ``property`` descriptors declared on the class, so the export protocol
    also covers the handful of non-dataclass result types.

    The property sweep matters: a result class may keep all of its state in
    private attributes (``self._tables``) and expose it exclusively through
    read-only properties (``KMResult.survival_table``). Harvesting only
    ``__dict__`` returned ``{}`` for such classes, which silently produced
    empty .xlsx / .docx exports instead of failing loudly (CLAUDE.md §7).
    """
    if is_dataclass(obj):
        return {f.name: _to_jsonable(getattr(obj, f.name)) for f in fields(obj)}
    attrs = getattr(obj, "__dict__", None)
    if attrs is None:
        raise TypeError(
            "result_to_dict expects a result dataclass or an object with "
            f"instance attributes, got {type(obj).__name__}."
        )
    out: Dict[str, Any] = {
        k: _to_jsonable(v)
        for k, v in attrs.items()
        if not k.startswith("_") and not callable(v)
    }
    for klass in type(obj).__mro__:
        for name, member in vars(klass).items():
            if name.startswith("_") or name in out:
                continue
            if not isinstance(member, property):
                continue
            # Deliberately not guarded: a property that raises is a real
            # defect in the result class, and swallowing it here would
            # reproduce exactly the silent-empty-export bug this sweep
            # exists to fix (CLAUDE.md §7).
            out[name] = _to_jsonable(getattr(obj, name))
    return out


def _fmt_scalar(v: Any) -> str:
    if isinstance(v, bool):
        return r"\text{" + str(v) + "}"
    if isinstance(v, int):
        return str(v)
    if isinstance(v, float):
        return f"{v:.4g}"
    return str(v).replace("_", r"\_").replace("%", r"\%").replace("&", r"\&")


def _accepts_second_positional(fn: Any) -> bool:
    """True when ``fn`` can be called as ``fn(a, b)``.

    Bespoke ``to_docx`` renderers come in both ``(filename)`` and
    ``(filename, title)`` shapes; calling the former with a title raises
    ``TypeError``. Unintrospectable callables are assumed to accept it, so
    the historical two-argument call is preserved.
    """
    import inspect

    try:
        sig = inspect.signature(fn)
    except (TypeError, ValueError):
        return True
    positional = 0
    for param in sig.parameters.values():
        if param.kind is param.VAR_POSITIONAL:
            return True
        if param.kind in (param.POSITIONAL_ONLY, param.POSITIONAL_OR_KEYWORD):
            positional += 1
    return positional >= 2


def _warn_if_no_rows(obj: Any, n_rows: int) -> None:
    """Warn when a result yields nothing exportable (CLAUDE.md §7).

    Writing a header-only .xlsx / .docx and reporting success is a silent
    degradation: the caller gets a file that looks valid and contains no
    data. Surface it instead.
    """
    if n_rows:
        return
    warnings.warn(
        f"{type(obj).__name__} produced no exportable fields; the export "
        "will contain only a header row. This usually means the result "
        "class keeps its state in private attributes without public "
        "properties, so nothing could be harvested.",
        RuntimeWarning,
        stacklevel=3,
    )


def result_export_rows(obj: Any) -> "list[Tuple[str, Any]]":
    """``(field, value)`` rows for tabular export of a result dataclass.

    Scalars pass through natively (so Excel keeps numeric cells); list /
    dict / array fields are summarised as ``[N items]`` rather than dumped,
    mirroring :func:`result_to_latex`. ``None`` fields are dropped.

    Warns (rather than silently writing an empty file) when the result
    exposes nothing exportable.
    """
    rows: list[Tuple[str, Any]] = []
    for key, val in result_to_dict(obj).items():
        if val is None:
            continue
        if isinstance(val, (list, dict)):
            n = len(val)
            rows.append((key, f"[{n} item{'s' if n != 1 else ''}]"))
        else:
            rows.append((key, val))
    _warn_if_no_rows(obj, len(rows))
    return rows


def result_to_latex(
    obj: Any, *, caption: str | None = None, label: str | None = None
) -> str:
    """Render the *scalar* fields of a result dataclass as a booktabs table.

    Array / DataFrame / nested fields are summarised as their shape rather than
    dumped, so the table stays publication-sized. Mirrors the ``to_latex``
    other result objects expose, so ``sp.<est>(...).to_latex()`` works
    uniformly.
    """
    d = result_to_dict(obj)
    rows = []
    for key, val in d.items():
        if val is None:
            continue
        if isinstance(val, (list, dict)):
            n = len(val)
            disp = f"[{n} item{'s' if n != 1 else ''}]"
        else:
            disp = _fmt_scalar(val)
        rows.append((str(key).replace("_", r"\_"), disp))
    _warn_if_no_rows(obj, len(rows))
    cap = caption or type(obj).__name__
    body = " \\\\\n".join(f"  {k} & {v}" for k, v in rows)
    lines = [
        "\\begin{table}[htbp]",
        "\\centering",
        f"\\caption{{{cap}}}",
    ]
    if label:
        lines.append(f"\\label{{{label}}}")
    lines += [
        "\\begin{tabular}{lr}",
        "\\toprule",
        "Field & Value \\\\",
        "\\midrule",
        body + " \\\\" if body else "",
        "\\bottomrule",
        "\\end{tabular}",
        "\\end{table}",
    ]
    return "\n".join(lines)


class ResultProtocolMixin:
    """Mixin giving a result dataclass ``to_dict`` / ``to_latex`` / ``cite``.

    Set ``_citation_keys`` to the paper.bib key(s) of the method's canonical
    reference(s) — verified to exist in ``paper.bib`` (CLAUDE.md §10) — to make
    ``cite()`` return them; leave it empty for textbook methods with no single
    canonical paper.
    """

    #: Verified paper.bib key(s) for the estimator (see CLAUDE.md §10).
    _citation_keys: ClassVar[Tuple[str, ...]] = ()

    def to_dict(self) -> Dict[str, Any]:
        """JSON-safe dict of every field (agent-native serialization)."""
        return result_to_dict(self)

    def to_latex(self, *, caption: str | None = None, label: str | None = None) -> str:
        """A compact booktabs table of the result's scalar fields."""
        return result_to_latex(self, caption=caption, label=label)

    def to_markdown(self) -> str:
        """GitHub-flavoured markdown table of the result's scalar fields."""
        rows = result_export_rows(self)
        lines = [
            f"### {type(self).__name__}",
            "",
            "| Field | Value |",
            "| --- | --- |",
        ]
        for key, val in rows:
            disp = f"{val:.6g}" if isinstance(val, float) else str(val)
            # A literal newline ends the table row and orphans the rest of
            # the value as body text, silently truncating the table; a bare
            # pipe opens a spurious column.
            disp = disp.replace("|", r"\|").replace("\r\n", " ").replace("\n", " ")
            disp = disp.replace("\r", " ")
            lines.append(f"| {key} | {disp} |")
        return "\n".join(lines)

    def to_word(self, filename: str, title: str | None = None) -> str:
        """§3 Word (.docx) export.

        Uses the class's own ``to_docx`` when it defines one (bespoke
        renderers win); otherwise writes a two-column Field/Value table of
        the scalar fields via ``python-docx`` (a core dependency).

        ``to_docx`` implementations differ in whether they accept a title
        (``to_docx(filename)`` vs ``to_docx(filename, title)``), so the
        title is only forwarded when the callee actually accepts it.
        """
        fn = getattr(self, "to_docx", None)
        if fn is not None and not getattr(fn, "__isabstractmethod__", False):
            if title is None or not _accepts_second_positional(fn):
                fn(filename)
            else:
                fn(filename, title)
            return filename

        from docx import Document

        doc = Document()
        doc.add_heading(title or type(self).__name__, level=1)
        rows = result_export_rows(self)
        table = doc.add_table(rows=len(rows) + 1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for i, (key, val) in enumerate(rows, start=1):
            table.rows[i].cells[0].text = str(key)
            table.rows[i].cells[1].text = (
                f"{val:.6g}" if isinstance(val, float) else str(val)
            )
        doc.save(filename)
        return filename

    def to_excel(self, path: str, **kwargs: Any) -> str:
        """§3 Excel (.xlsx) export.

        Writes the scalar fields as a two-column Field/Value sheet via
        pandas/openpyxl (core dependencies). Subclasses with a bespoke
        ``to_excel`` shadow this default automatically (Python MRO).

        .. note::
           Until 2026-07 this method delegated to ``getattr(self,
           "to_excel")`` — i.e. *itself* — so any subclass that did not
           override it recursed forever. It now writes a real workbook.
        """
        import pandas as pd

        rows = result_export_rows(self)
        frame = pd.DataFrame(rows, columns=["Field", "Value"])
        sheet = kwargs.pop("sheet_name", type(self).__name__[:31])
        frame.to_excel(path, index=False, sheet_name=sheet, **kwargs)
        return path

    def cite(self, format: str = "keys") -> Any:
        """Return the estimator's verified paper.bib citation key(s).

        Parameters
        ----------
        format : {"keys", "json"}, default ``"keys"``
            ``"keys"`` returns a newline-joined string of bib keys (or an
            honest placeholder); ``"json"`` returns a structured dict.

        Notes
        -----
        Zero-hallucination (CLAUDE.md §10): the keys point into ``paper.bib``;
        resolve them to full BibTeX via ``sp.bibtex(keys=[...])``. A method
        with no single canonical paper honestly returns a placeholder rather
        than a fabricated reference.
        """
        keys = tuple(getattr(self, "_citation_keys", ()) or ())
        if format == "json":
            return {
                "citation_keys": list(keys),
                "source": "paper.bib",
                "resolve_with": "sp.bibtex(keys=[...])",
            }
        if format != "keys":
            raise ValueError(f"format must be 'keys' or 'json'; got {format!r}")
        if not keys:
            return (
                f"No single canonical reference registered for "
                f"{type(self).__name__}; see the estimator docstring."
            )
        return "\n".join(keys)
