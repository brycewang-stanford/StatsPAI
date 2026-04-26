"""
Descriptive statistics and balance tables.

Provides:
- sumstats(): Summary statistics table (Table 1 in every empirical paper)
- balance_table(): Pre-treatment balance check for matching/DID

Equivalent to Stata's ``summarize``, ``tabstat``, ``balancetable``.
"""

from typing import Optional, List, Dict, Any, Union
import numpy as np
import pandas as pd


def sumstats(
    data: pd.DataFrame,
    vars: Optional[List[str]] = None,
    by: Optional[str] = None,
    stats: Optional[List[str]] = None,
    output: str = 'text',
    title: str = 'Summary Statistics',
    fmt: str = '%.3f',
    labels: Optional[Dict[str, str]] = None,
    by_labels: Optional[Dict[Any, str]] = None,
) -> Union[str, pd.DataFrame]:
    """
    Generate descriptive statistics table.

    Parameters
    ----------
    data : pd.DataFrame
        Input data.
    vars : list of str, optional
        Variables to summarize. Defaults to all numeric columns.
    by : str, optional
        Group-by variable for stratified summaries.
    stats : list of str, optional
        Statistics to compute. Default: ['n', 'mean', 'sd', 'min', 'p25',
        'median', 'p75', 'max'].
        Available: 'n', 'mean', 'sd', 'min', 'max', 'p25', 'median',
        'p75', 'p10', 'p90'.
    output : str, default 'text'
        'text', 'latex', 'html', 'dataframe', or filepath (.xlsx/.docx).
    title : str
        Table title.
    fmt : str
        Number format.
    labels : dict, optional
        Variable labels: ``{'x1': 'Education (years)'}``.
    by_labels : dict, optional
        Group labels for ``by=`` panel headers, e.g.
        ``{0: 'Control', 1: 'Treated'}``. When ``by`` is binary 0/1
        and no ``by_labels`` is supplied, ``Control``/``Treated`` is
        auto-applied so academic Table 1 reads correctly out of the box.

    Returns
    -------
    str or pd.DataFrame

    Examples
    --------
    >>> sp.sumstats(df, vars=['wage', 'edu', 'exp'], by='female')
    >>> sp.sumstats(df, output='table1.docx')
    """
    # Select variables
    if vars is None:
        vars = list(data.select_dtypes(include=[np.number]).columns)
    if stats is None:
        stats = ['n', 'mean', 'sd', 'min', 'p25', 'median', 'p75', 'max']

    stat_funcs = {
        'n': ('N', lambda s: int(s.count())),
        'mean': ('Mean', lambda s: s.mean()),
        'sd': ('Std. Dev.', lambda s: s.std()),
        'min': ('Min', lambda s: s.min()),
        'max': ('Max', lambda s: s.max()),
        'p10': ('P10', lambda s: s.quantile(0.1)),
        'p25': ('P25', lambda s: s.quantile(0.25)),
        'median': ('Median', lambda s: s.median()),
        'p75': ('P75', lambda s: s.quantile(0.75)),
        'p90': ('P90', lambda s: s.quantile(0.9)),
    }

    if by is None:
        df_result = _compute_stats(data, vars, stats, stat_funcs, fmt, labels)
    else:
        groups = sorted(data[by].dropna().unique())
        # Auto Control/Treated for binary 0/1 unless caller supplied labels.
        if by_labels is None:
            unique_set = set(groups)
            if unique_set <= {0, 1, 0.0, 1.0} and len(unique_set) >= 1:
                by_labels = {
                    0: "Control", 1: "Treated",
                    0.0: "Control", 1.0: "Treated",
                }
        panels = {}
        for g in groups:
            label = by_labels.get(g, str(g)) if by_labels else str(g)
            subset = data[data[by] == g]
            panels[label] = _compute_stats(
                subset, vars, stats, stat_funcs, fmt, labels
            )
        # Stack panels
        df_result = pd.concat(panels, axis=1)

    return _format_output(df_result, output, title, stats)


def _compute_stats(data, vars, stats, stat_funcs, fmt, labels):
    """Compute statistics for each variable."""
    rows = {}
    for var in vars:
        if var not in data.columns:
            continue
        s = data[var].dropna()
        display = labels.get(var, var) if labels else var
        row = {}
        for stat in stats:
            if stat in stat_funcs:
                label, fn = stat_funcs[stat]
                val = fn(s)
                if stat == 'n':
                    row[label] = str(int(val))
                else:
                    row[label] = fmt % val if not np.isnan(val) else ''
        rows[display] = row
    return pd.DataFrame(rows).T


def balance_table(
    data: pd.DataFrame,
    treat: str,
    covariates: List[str],
    output: str = 'text',
    title: str = 'Balance Table',
    fmt: str = '%.3f',
    labels: Optional[Dict[str, str]] = None,
    test: str = 'ttest',
) -> Union[str, pd.DataFrame]:
    """
    Generate a balance table comparing treated and control groups.

    Standard Table 1 for matching, DID, and RCT papers.

    Parameters
    ----------
    data : pd.DataFrame
        Input data.
    treat : str
        Binary treatment variable (0/1).
    covariates : list of str
        Variables to check balance on.
    output : str, default 'text'
        'text', 'latex', 'html', 'dataframe', or filepath.
    title : str
        Table title.
    fmt : str
        Number format.
    labels : dict, optional
        Variable labels.
    test : str, default 'ttest'
        Test for difference: 'ttest' or 'ranksum'.

    Returns
    -------
    str or pd.DataFrame

    Examples
    --------
    >>> sp.balance_table(df, treat='treated',
    ...                  covariates=['age', 'edu', 'income'],
    ...                  output='balance.docx')
    """
    from scipy import stats as sp_stats

    T = data[treat].values
    treated = data[T == 1]
    control = data[T == 0]

    rows = []
    for var in covariates:
        if var not in data.columns:
            continue
        display = labels.get(var, var) if labels else var

        t_vals = treated[var].dropna()
        c_vals = control[var].dropna()

        mean_t = t_vals.mean()
        mean_c = c_vals.mean()
        sd_t = t_vals.std()
        sd_c = c_vals.std()

        # Standardized mean difference
        sd_pooled = np.sqrt((sd_t**2 + sd_c**2) / 2)
        smd = (mean_t - mean_c) / sd_pooled if sd_pooled > 0 else 0

        # Test
        if test == 'ttest':
            stat, pval = sp_stats.ttest_ind(t_vals, c_vals, equal_var=False)
        else:
            stat, pval = sp_stats.ranksums(t_vals, c_vals)

        rows.append({
            'Variable': display,
            'Treated Mean': fmt % mean_t,
            'Treated SD': fmt % sd_t,
            'Control Mean': fmt % mean_c,
            'Control SD': fmt % sd_c,
            'Diff': fmt % (mean_t - mean_c),
            'SMD': fmt % smd,
            'p-value': '%.3f' % pval,
        })

    df_result = pd.DataFrame(rows).set_index('Variable')

    # Add N row
    n_row = pd.DataFrame([{
        'Variable': 'N',
        'Treated Mean': str(len(treated)),
        'Treated SD': '',
        'Control Mean': str(len(control)),
        'Control SD': '',
        'Diff': '',
        'SMD': '',
        'p-value': '',
    }]).set_index('Variable')
    df_result = pd.concat([df_result, n_row])

    return _format_output(df_result, output, title, None)


# ======================================================================
# Output formatting
# ======================================================================

def _format_output(df, output, title, stats_list):
    """Route to the appropriate output format."""
    if output == 'dataframe':
        return df
    elif output.endswith('.xlsx'):
        _sumstats_to_excel(df, output, title)
        return f"Exported to: {output}"
    elif output.endswith('.docx'):
        _sumstats_to_word(df, output, title)
        return f"Exported to: {output}"
    elif output == 'latex':
        return _sumstats_to_latex(df, title)
    elif output == 'html':
        return _sumstats_to_html(df, title)
    else:
        return _sumstats_to_text(df, title)


def _sumstats_to_text(df, title):
    """Plain text table."""
    lines = []
    if title:
        lines.append(title)
    lines.append('=' * 80)
    lines.append(df.to_string())
    lines.append('=' * 80)
    return '\n'.join(lines)


def _sumstats_to_latex(df, title):
    """LaTeX table."""
    latex = df.to_latex(caption=title, label='tab:sumstats')
    return latex


def _sumstats_to_html(df, title):
    """HTML table."""
    html = f'<h3>{title}</h3>\n' if title else ''
    html += df.to_html()
    return html


def _sumstats_to_excel(df, filename, title):
    """Export sumstats DataFrame as a book-tab Excel table.

    Delegates layout + borders to ``_excel_style.render_dataframe_to_xlsx``
    so the visual conventions stay aligned with every other StatsPAI
    xlsx writer (Times New Roman, top/mid/bottom rules, merged
    panel headers for ``by=`` MultiIndex columns).
    """
    from ._excel_style import render_dataframe_to_xlsx

    render_dataframe_to_xlsx(
        df,
        filename,
        title=title,
        sheet_name="Summary",
        index_label="Variable",
    )


def _sumstats_to_word(df, filename, title):
    """Export to Word (.docx) in AER/QJE book-tab style."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx required. Install: pip install python-docx")

    from ._aer_style import (
        apply_word_booktab_rules,
        style_word_table_typography,
    )

    doc = Document()

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    n_rows = len(df) + 1
    n_cols = len(df.columns) + 1  # +1 for index
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header
    table.rows[0].cells[0].text = "Variable"
    for j, col in enumerate(df.columns, 1):
        table.rows[0].cells[j].text = str(col)
    # Body
    for i, (idx, row) in enumerate(df.iterrows()):
        table.rows[i + 1].cells[0].text = str(idx)
        for j, val in enumerate(row, 1):
            table.rows[i + 1].cells[j].text = str(val)

    style_word_table_typography(
        table, header_rows=(0,),
        header_pt=10, body_pt=9,
        align_first_col="left", align_data_cols="center",
    )
    apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)

    doc.save(filename)
