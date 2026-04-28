"""
Publication-quality regression tables and balance tables.

Provides ``regtable()`` for unified regression output across formats,
and ``mean_comparison()`` for balance / summary statistics tables.

Usage
-----
>>> import statspai as sp
>>> m1 = sp.regress("y ~ x1", data=df)
>>> m2 = sp.regress("y ~ x1 + x2", data=df)
>>> sp.regtable(m1, m2)
>>> sp.regtable(m1, m2, output="latex", filename="table1.tex")
>>>
>>> sp.mean_comparison(df, variables=["age", "income"], group="treated")
"""

from __future__ import annotations

import warnings
from collections import OrderedDict
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union  # noqa: F401

import numpy as np
import pandas as pd
from scipy import stats as sp_stats


# ---------------------------------------------------------------------------
# Re-use extraction helpers from estimates module
# ---------------------------------------------------------------------------

from .estimates import (
    _ci_bounds,
    _extract_model_data,
    _ModelData,
    _format_stars,
    _fmt_val,
    _fmt_int,
    _latex_escape,
    _html_escape,
    _STAT_ALIASES,
    _STAT_DISPLAY,
)
from ._diagnostics import extract_diagnostic_rows
from ._journals import get_template, list_templates, star_note_for
from ._repro import build_repro_note


# Bracket styles cycled through when rendering ``multi_se`` extra SE rows.
# The four bracket pairs are deliberately Markdown-safe: a fourth ``||`` pair
# would collide with GFM pipe-table delimiters and break the row, so we use
# guillemets ``«»`` instead. The primary SE always uses parentheses.
_MULTI_SE_BRACKETS = (("[", "]"), ("{", "}"), ("⟨", "⟩"), ("«", "»"))


def _resolve_multi_se(
    multi_se: Optional[Dict[str, Sequence[Any]]],
    n_models: int,
) -> List[Tuple[str, List[Dict[str, float]]]]:
    """Validate and normalise a ``multi_se`` argument.

    Returns a list of ``(label, [per-model dict-of-var->se, ...])`` tuples
    in user-supplied order. Each per-model entry is a plain ``dict``
    keyed by coefficient name; missing variables yield empty cells.
    """
    if not multi_se:
        return []
    out: List[Tuple[str, List[Dict[str, float]]]] = []
    for label, per_model in multi_se.items():
        if not isinstance(label, str) or not label:
            raise ValueError(f"multi_se keys must be non-empty strings, got {label!r}.")
        seq = list(per_model) if per_model is not None else []
        if len(seq) != n_models:
            raise ValueError(
                f"multi_se[{label!r}] has {len(seq)} entries but there are "
                f"{n_models} models."
            )
        normalized: List[Dict[str, float]] = []
        for entry in seq:
            if entry is None:
                normalized.append({})
                continue
            if isinstance(entry, pd.Series):
                normalized.append({str(k): float(v) for k, v in entry.items()
                                   if v is not None and not pd.isna(v)})
            elif isinstance(entry, dict):
                normalized.append({str(k): float(v) for k, v in entry.items()
                                   if v is not None and not pd.isna(v)})
            else:
                raise TypeError(
                    f"multi_se[{label!r}] entries must be pandas.Series or "
                    f"dict, got {type(entry).__name__}."
                )
        out.append((label, normalized))
    return out


# ═══════════════════════════════════════════════════════════════════════════
# RegtableResult
# ═══════════════════════════════════════════════════════════════════════════

class RegtableResult:
    """Rich result object for regression tables with multi-format export."""

    def __init__(
        self,
        panels: List["_PanelData"],
        *,
        panel_labels: Optional[List[str]],
        model_labels: List[str],
        dep_var_labels: Optional[List[str]],
        coef_labels: Optional[Dict[str, str]],
        keep: Optional[List[str]],
        drop: Optional[List[str]],
        order: Optional[List[str]],
        se_type: str,
        stars: bool,
        star_levels: Tuple[float, ...],
        fmt: str,
        title: Optional[str],
        notes: Optional[List[str]],
        add_rows: Optional[Dict[str, List[str]]],
        stats: Optional[List[str]],
        output: str = "text",
        alpha: float = 0.05,
        multi_se: Optional[List[Tuple[str, List[Dict[str, float]]]]] = None,
        se_label: Optional[str] = None,
        template: Optional[str] = None,
        quarto_label: Optional[str] = None,
        quarto_caption: Optional[str] = None,
        eform_flags: Optional[List[bool]] = None,
        column_spanners: Optional[List[Tuple[str, int]]] = None,
    ):
        if not 0.0 < alpha < 1.0:
            raise ValueError(f"alpha must be in (0, 1), got {alpha!r}")
        self.panels = panels
        self.panel_labels = panel_labels
        self.model_labels = model_labels
        self.dep_var_labels = dep_var_labels
        self.coef_labels = coef_labels or {}
        self.keep = keep
        self.drop = set(drop) if drop else set()
        self.order = order
        self.se_type = se_type
        self.show_stars = stars
        self.star_levels = star_levels
        self.fmt = fmt
        self.title = title
        self.notes = notes or []
        self.add_rows = add_rows or {}
        self.requested_stats = stats or ["N", "R2", "adj_R2", "F"]
        self.alpha = float(alpha)
        self.n_models = sum(len(p.models) for p in panels)
        self.multi_se = multi_se or []
        # Override the SE-row footer label (e.g. "Robust standard errors"
        # via QJE preset). When None, the label is derived from se_type.
        self._se_label_override = se_label
        # Journal template name (informational; resolution already happened
        # at the regtable() call site).
        self.template = template
        # Quarto cross-reference metadata. When ``quarto_label`` is set,
        # ``to_quarto()`` (and ``to_markdown(quarto=True)``) emit a Quarto
        # ``: caption {#tbl-<label>}`` line so the table can be referenced
        # via ``@tbl-<label>`` in the manuscript.
        self.quarto_label = quarto_label
        self.quarto_caption = quarto_caption
        # Controls which renderer __str__ uses. Jupyter still gets HTML via
        # _repr_html_ regardless, so output='latex' in a notebook still renders
        # pretty HTML — users who want the LaTeX source call to_latex() or
        # print(result).
        self._output = output

        # Per-model eform flag (length n_models). When True for a model,
        # the rendered point estimate becomes exp(b), the SE becomes
        # exp(b)·SE(b) (delta method), and CI bounds become (exp(lo), exp(hi))
        # of the original endpoints. t and p-values are unchanged because
        # H0: b=0 is equivalent to H0: exp(b)=1.
        if eform_flags is None:
            eform_flags = [False] * self.n_models
        if len(eform_flags) != self.n_models:
            raise ValueError(
                f"eform_flags has {len(eform_flags)} entries but "
                f"there are {self.n_models} models."
            )
        self.eform_flags = [bool(f) for f in eform_flags]

        # Column spanners: a list of ``(label, span)`` tuples where the
        # spans must sum to n_models. Renders as a multi-row header above
        # the model labels (LaTeX ``\multicolumn``, HTML ``colspan``,
        # text-mode centered ASCII).
        if column_spanners is not None:
            total_span = sum(int(s) for _, s in column_spanners)
            if total_span != self.n_models:
                raise ValueError(
                    f"column_spanners total span = {total_span} but there are "
                    f"{self.n_models} models. The spans must partition the "
                    f"columns exactly (consecutive grouping)."
                )
        self.column_spanners = (
            [(str(lbl), int(s)) for lbl, s in column_spanners]
            if column_spanners
            else None
        )

        # Resolve stat keys once
        self._stat_keys = self._resolve_stat_keys()

    # --- helpers -----------------------------------------------------------

    def _resolve_stat_keys(self) -> List[str]:
        keys: List[str] = []
        for s in self.requested_stats:
            canonical = _STAT_ALIASES.get(s, s)
            if canonical not in keys:
                keys.append(canonical)
        return keys

    def _resolve_vars(self, models: List[_ModelData]) -> List[str]:
        seen: OrderedDict[str, None] = OrderedDict()
        for m in models:
            for v in m.params.index:
                seen[v] = None
        all_vars = list(seen)

        if self.keep is not None:
            keep_set = set(self.keep)
            all_vars = [v for v in all_vars if v in keep_set]
        if self.drop:
            all_vars = [v for v in all_vars if v not in self.drop]
        if self.order:
            ordered: List[str] = []
            remaining = list(all_vars)
            for v in self.order:
                if v in remaining:
                    ordered.append(v)
                    remaining.remove(v)
            ordered.extend(remaining)
            all_vars = ordered
        return all_vars

    def _model_eform(self, flat_idx: int) -> bool:
        """Return whether model at flat position ``flat_idx`` uses eform."""
        if not self.eform_flags:
            return False
        if 0 <= flat_idx < len(self.eform_flags):
            return self.eform_flags[flat_idx]
        return False

    def _coef_cell(self, model: _ModelData, var: str, flat_idx: int = 0) -> str:
        if var not in model.params.index:
            return ""
        val = float(model.params[var])
        if self._model_eform(flat_idx) and np.isfinite(val):
            val = float(np.exp(val))
        txt = _fmt_val(val, self.fmt)
        if self.show_stars and var in model.pvalues.index:
            txt += _format_stars(model.pvalues[var], self.star_levels)
        return txt

    def _se_cell(self, model: _ModelData, var: str, flat_idx: int = 0) -> str:
        if var not in model.params.index:
            return ""
        eform = self._model_eform(flat_idx)
        if self.se_type == "ci":
            lo_v, hi_v = _ci_bounds(model, var, self.alpha)
            if eform:
                if not (lo_v is None or pd.isna(lo_v)):
                    lo_v = float(np.exp(lo_v))
                if not (hi_v is None or pd.isna(hi_v)):
                    hi_v = float(np.exp(hi_v))
            lo = _fmt_val(lo_v, self.fmt)
            hi = _fmt_val(hi_v, self.fmt)
            return f"[{lo}, {hi}]"
        if self.se_type == "t":
            # t = b / SE(b); the null H0: exp(b)=1 is equivalent to b=0,
            # so the t and p statistics are unchanged under eform.
            return f"({_fmt_val(model.tvalues.get(var, np.nan), self.fmt)})"
        if self.se_type == "p":
            return f"({_fmt_val(model.pvalues.get(var, np.nan), self.fmt)})"
        # default: standard error (delta method under eform)
        se_v = model.std_errors.get(var, np.nan)
        if eform and not pd.isna(se_v):
            b = float(model.params[var])
            if np.isfinite(b):
                se_v = float(np.exp(b)) * float(se_v)
        return f"({_fmt_val(se_v, self.fmt)})"

    def _has_any_eform(self) -> bool:
        return any(self.eform_flags)

    def _se_label(self) -> str:
        if self._se_label_override is not None and self.se_type == "se":
            return self._se_label_override
        if self.se_type == "ci":
            level = (1.0 - self.alpha) * 100.0
            return f"{level:g}% CI"
        return {"t": "t-statistics", "p": "p-values"}.get(
            self.se_type, "Standard errors"
        )

    def _multi_se_cell(
        self,
        per_model: Dict[str, float],
        var: str,
        bracket_idx: int,
        model: Optional[_ModelData] = None,
        flat_idx: int = 0,
    ) -> str:
        """Render the bracket-wrapped extra-SE cell for one model column.

        When eform is active for the column, the extra SE (bootstrap /
        cluster / jackknife / etc.) is rescaled by ``exp(b)`` via the
        same delta-method approximation used for the primary SE — keeps
        the cells reading on a single scale.
        """
        if var not in per_model:
            return ""
        try:
            val = float(per_model[var])
        except (TypeError, ValueError):
            return ""
        if not np.isfinite(val):
            return ""
        if (
            self._model_eform(flat_idx)
            and model is not None
            and var in model.params.index
        ):
            b = float(model.params[var])
            if np.isfinite(b):
                val = float(np.exp(b)) * val
        lo, hi = _MULTI_SE_BRACKETS[bracket_idx % len(_MULTI_SE_BRACKETS)]
        return f"{lo}{_fmt_val(val, self.fmt)}{hi}"

    def _stat_cell(self, model: _ModelData, key: str) -> str:
        val = model.stats.get(key)
        if val is None:
            return ""
        if key == "N":
            return _fmt_int(val)
        return _fmt_val(float(val), "%.3f")

    def _star_note(self) -> str:
        # Delegate to ``star_note_for`` so the renderer's footer line and the
        # journal-template ``notes_default`` lines use the same strict-first
        # convention ("*** p<0.01, ** p<0.05, * p<0.10"). Reduces drift.
        return star_note_for(self.star_levels)

    def _all_models_flat(self) -> List[_ModelData]:
        out: List[_ModelData] = []
        for p in self.panels:
            out.extend(p.models)
        return out

    # ═══════════════════════════════════════════════════════════════════════
    # TEXT
    # ═══════════════════════════════════════════════════════════════════════

    def _text_panel(
        self,
        models: List[_ModelData],
        var_list: List[str],
        col_w: int,
        label_w: int,
        panel_idx: int = 0,
    ) -> List[str]:
        lines: List[str] = []
        base_idx = sum(len(p.models) for p in self.panels[:panel_idx])
        for var in var_list:
            label = self.coef_labels.get(var, var)
            row = f"{label:<{label_w}}"
            for off, m in enumerate(models):
                row += f"{self._coef_cell(m, var, base_idx + off):>{col_w}}"
            lines.append(row)
            # SE row
            row2 = " " * label_w
            for off, m in enumerate(models):
                row2 += f"{self._se_cell(m, var, base_idx + off):>{col_w}}"
            lines.append(row2)
            # Extra SE rows from multi_se. Each label maps to one entry per
            # model across the WHOLE table, so we slice into the panel using
            # cumulative model offsets.
            for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                row3 = " " * label_w
                for off, m in enumerate(models):
                    per_model = per_model_list[base_idx + off]
                    row3 += f"{self._multi_se_cell(per_model, var, ext_idx, m, base_idx + off):>{col_w}}"
                lines.append(row3)
            lines.append("")  # blank between vars
        return lines

    def to_text(self) -> str:
        col_w = 14
        all_models = self._all_models_flat()
        all_vars_set: set = set()
        for p in self.panels:
            all_vars_set.update(v for m in p.models for v in m.params.index)
        label_names = [self.coef_labels.get(v, v) for v in all_vars_set]
        stat_names = [_STAT_DISPLAY.get(k, k) for k in self._stat_keys]
        add_row_names = list(self.add_rows.keys())
        max_label = max(
            (len(n) for n in label_names + stat_names + add_row_names),
            default=10,
        )
        label_w = max(max_label + 2, 18)
        total_w = label_w + col_w * len(all_models) + 2

        thick = "\u2501" * total_w
        thin = "\u2500" * total_w
        lines: List[str] = []

        if self.title:
            lines.append(f"  {self.title}")
            lines.append("")

        lines.append(thick)

        # Column spanners: a header row above the model-label row, with
        # each label centered over its block of columns.
        if self.column_spanners:
            span_row = " " * label_w
            for lbl, span in self.column_spanners:
                block_w = col_w * span
                span_row += f"{lbl:^{block_w}}"
            lines.append(span_row)
            # Thin underline beneath each spanner block (cmidrule analog)
            rule_row = " " * label_w
            for _, span in self.column_spanners:
                block_w = col_w * span
                rule_row += " " + "─" * (block_w - 2) + " "
            lines.append(rule_row)

        # Header: model labels
        hdr = " " * label_w
        for lbl in self.model_labels:
            hdr += f"{lbl:>{col_w}}"
        lines.append(hdr)

        # Dep-var row
        if self.dep_var_labels:
            dvr = " " * label_w
            for dv in self.dep_var_labels:
                dvr += f"{dv:>{col_w}}"
            lines.append(dvr)

        lines.append(thick)

        # Panels
        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(f"  {self.panel_labels[pi]}")
                lines.append(thin)

            var_list = self._resolve_vars(panel.models)
            lines.extend(self._text_panel(panel.models, var_list, col_w, label_w, panel_idx=pi))

            if multi and pi < len(self.panels) - 1:
                lines.append(thin)

        lines.append(thick)

        # Add rows (Controls, FE, etc.)
        for row_label, row_vals in self.add_rows.items():
            row = f"{row_label:<{label_w}}"
            for i, m in enumerate(all_models):
                val = row_vals[i] if i < len(row_vals) else ""
                row += f"{val:>{col_w}}"
            lines.append(row)

        if self.add_rows:
            lines.append(thick)

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            row = f"{disp:<{label_w}}"
            for m in all_models:
                row += f"{self._stat_cell(m, key):>{col_w}}"
            lines.append(row)

        lines.append(thick)

        # Notes
        lines.append(f"{self._se_label()} in parentheses")
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            lines.append(f"{label} in {lo}…{hi}")
        if self._has_any_eform():
            lines.append(self._eform_note())
        if self.show_stars:
            lines.append(self._star_note())
        for note in self.notes:
            lines.append(note)

        return "\n".join(lines)

    def _eform_note(self) -> str:
        """Footer note explaining the eform transformation."""
        if all(self.eform_flags):
            return ("Coefficients reported as exp(b); standard errors via "
                    "delta method (exp(b)·SE). Stars from p-values of the "
                    "untransformed estimates.")
        cols = [i + 1 for i, f in enumerate(self.eform_flags) if f]
        col_str = ", ".join(f"({c})" for c in cols)
        return (f"Columns {col_str} report exp(b) (delta-method SE); "
                f"other columns report b. Stars from p-values of the "
                f"untransformed estimates.")

    # ═══════════════════════════════════════════════════════════════════════
    # HTML (also _repr_html_)
    # ═══════════════════════════════════════════════════════════════════════

    def to_html(self) -> str:
        all_models = self._all_models_flat()
        ncols = len(all_models) + 1
        lines: List[str] = []
        lines.append(
            '<table class="regtable" style="border-collapse:collapse; '
            'font-family:\'Times New Roman\', serif; font-size:13px; min-width:500px;">'
        )

        if self.title:
            lines.append(
                f'<caption style="font-weight:bold; font-size:14px; '
                f'margin-bottom:8px; caption-side:top;">'
                f'{_html_escape(self.title)}</caption>'
            )

        # Header
        lines.append("<thead>")

        # Column spanners (above the model-label row) when set
        if self.column_spanners:
            lines.append("<tr>")
            lines.append(
                '<th style="text-align:left; border-top:3px solid black; '
                'padding:4px 8px;"></th>'
            )
            for lbl, span in self.column_spanners:
                lines.append(
                    f'<th colspan="{span}" style="text-align:center; '
                    f'border-top:3px solid black; border-bottom:1px solid #999; '
                    f'padding:4px 12px;">{_html_escape(lbl)}</th>'
                )
            lines.append("</tr>")

        lines.append("<tr>")
        # When spanners present, the top rule already lives on the spanner row
        top_rule = "" if self.column_spanners else "border-top:3px solid black; "
        lines.append(
            f'<th style="text-align:left; {top_rule}'
            'border-bottom:1px solid black; padding:4px 8px;"></th>'
        )
        for lbl in self.model_labels:
            lines.append(
                f'<th style="text-align:center; {top_rule}'
                f'border-bottom:1px solid black; padding:4px 12px;">'
                f'{_html_escape(lbl)}</th>'
            )
        lines.append("</tr>")

        # Dep-var row
        if self.dep_var_labels:
            lines.append("<tr>")
            lines.append('<th style="text-align:left; padding:2px 8px;"></th>')
            for dv in self.dep_var_labels:
                lines.append(
                    f'<th style="text-align:center; padding:2px 12px; '
                    f'font-style:italic; font-weight:normal;">'
                    f'{_html_escape(dv)}</th>'
                )
            lines.append("</tr>")

        lines.append("</thead>")
        lines.append("<tbody>")

        # Panels
        multi = len(self.panels) > 1
        model_idx = 0
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(
                    f'<tr><td colspan="{ncols}" style="text-align:left; '
                    f'font-weight:bold; padding:6px 8px 2px 8px; '
                    f'border-top:1px solid #999;">'
                    f'{_html_escape(self.panel_labels[pi])}</td></tr>'
                )

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = _html_escape(self.coef_labels.get(var, var))
                lines.append("<tr>")
                lines.append(
                    f'<td style="text-align:left; padding:1px 8px;">{label}</td>'
                )
                # Empty cells for models NOT in this panel
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            lines.append(
                                f'<td style="text-align:center; padding:1px 12px;">'
                                f'{_html_escape(self._coef_cell(m, var, flat_idx))}</td>'
                            )
                        else:
                            lines.append(
                                '<td style="text-align:center; padding:1px 12px;"></td>'
                            )
                        flat_idx += 1
                lines.append("</tr>")
                # SE row
                lines.append("<tr>")
                lines.append("<td></td>")
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            lines.append(
                                f'<td style="text-align:center; padding:0 12px; '
                                f'color:#555; font-size:12px;">'
                                f'{_html_escape(self._se_cell(m, var, flat_idx))}</td>'
                            )
                        else:
                            lines.append(
                                '<td style="text-align:center; padding:0 12px;"></td>'
                            )
                        flat_idx += 1
                lines.append("</tr>")
                # Extra SE rows from multi_se
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    lines.append("<tr>")
                    lines.append("<td></td>")
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                cell = self._multi_se_cell(
                                    per_model, var, ext_idx, m, base_idx + off
                                )
                                lines.append(
                                    f'<td style="text-align:center; padding:0 12px; '
                                    f'color:#777; font-size:12px;">'
                                    f'{_html_escape(cell)}</td>'
                                )
                            else:
                                lines.append(
                                    '<td style="text-align:center; padding:0 12px;"></td>'
                                )
                    lines.append("</tr>")

        # Separator
        lines.append(
            f'<tr><td colspan="{ncols}" '
            f'style="border-top:1px solid black; padding:0;"></td></tr>'
        )

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            lines.append("<tr>")
            lines.append(
                f'<td style="text-align:left; padding:1px 8px;">'
                f'{_html_escape(row_label)}</td>'
            )
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{_html_escape(val)}</td>'
                )
            lines.append("</tr>")

        if self.add_rows:
            lines.append(
                f'<tr><td colspan="{ncols}" '
                f'style="border-top:1px solid #aaa; padding:0;"></td></tr>'
            )

        # Stats
        for key in self._stat_keys:
            disp = _html_escape(_STAT_DISPLAY.get(key, key))
            lines.append("<tr>")
            lines.append(
                f'<td style="text-align:left; padding:1px 8px;">{disp}</td>'
            )
            for m in all_models:
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{self._stat_cell(m, key)}</td>'
                )
            lines.append("</tr>")

        # Bottom border
        lines.append(
            f'<tr><td colspan="{ncols}" '
            f'style="border-top:3px solid black; padding:0;"></td></tr>'
        )

        lines.append("</tbody>")

        # Notes
        lines.append("<tfoot>")
        note_text = f"{self._se_label()} in parentheses"
        lines.append(
            f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
            f'padding:4px 8px 0 8px;">{_html_escape(note_text)}</td></tr>'
        )
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            multi_note = f"{label} in {lo}…{hi}"
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(multi_note)}</td></tr>'
            )
        if self._has_any_eform():
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(self._eform_note())}</td></tr>'
            )
        if self.show_stars:
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(self._star_note())}</td></tr>'
            )
        for note in self.notes:
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(note)}</td></tr>'
            )
        lines.append("</tfoot>")
        lines.append("</table>")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # LaTeX
    # ═══════════════════════════════════════════════════════════════════════

    def to_latex(self) -> str:
        all_models = self._all_models_flat()
        n_cols = len(all_models) + 1
        col_spec = "l" + "c" * len(all_models)
        lines: List[str] = []
        lines.append("\\begin{table}[htbp]")
        lines.append("\\centering")
        if self.title:
            lines.append(f"\\caption{{{_latex_escape(self.title)}}}")
        lines.append(f"\\begin{{tabular}}{{{col_spec}}}")
        lines.append("\\hline\\hline")

        # Column spanners (above the model-label row) when set
        if self.column_spanners:
            cells_sp: List[str] = [""]
            cmidrules: List[str] = []
            cur_col = 2  # LaTeX column 1 is the row-label, models start at 2
            for lbl, span in self.column_spanners:
                cells_sp.append(
                    f"\\multicolumn{{{span}}}{{c}}{{{_latex_escape(lbl)}}}"
                )
                cmidrules.append(
                    f"\\cmidrule(lr){{{cur_col}-{cur_col + span - 1}}}"
                )
                cur_col += span
            lines.append(" & ".join(cells_sp) + " \\\\")
            lines.append("".join(cmidrules))

        # Header
        hdr = " & ".join(
            [""] + [_latex_escape(n) for n in self.model_labels]
        ) + " \\\\"
        lines.append(hdr)

        # Dep-var
        if self.dep_var_labels:
            dvr = " & ".join(
                [""] + [f"\\textit{{{_latex_escape(dv)}}}" for dv in self.dep_var_labels]
            ) + " \\\\"
            lines.append(dvr)

        lines.append("\\hline")

        # Panels
        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(
                    f"\\multicolumn{{{n_cols}}}{{l}}"
                    f"{{\\textbf{{{_latex_escape(self.panel_labels[pi])}}}}}"
                    " \\\\"
                )
                lines.append("\\hline")

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = _latex_escape(self.coef_labels.get(var, var))
                cells: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells.append(_latex_escape(self._coef_cell(m, var, flat_idx)))
                        else:
                            cells.append("")
                        flat_idx += 1
                lines.append(f"{label} & " + " & ".join(cells) + " \\\\")
                # SE row
                cells2: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells2.append(_latex_escape(self._se_cell(m, var, flat_idx)))
                        else:
                            cells2.append("")
                        flat_idx += 1
                lines.append(" & " + " & ".join(cells2) + " \\\\")
                # Extra SE rows (multi_se)
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    cells_ext: List[str] = []
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                cells_ext.append(
                                    _latex_escape(self._multi_se_cell(
                                        per_model, var, ext_idx, m, base_idx + off
                                    ))
                                )
                            else:
                                cells_ext.append("")
                    lines.append(" & " + " & ".join(cells_ext) + " \\\\")

            if multi and pi < len(self.panels) - 1:
                lines.append("\\hline")

        lines.append("\\hline")

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            cells_ar: List[str] = []
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                cells_ar.append(_latex_escape(val))
            lines.append(
                f"{_latex_escape(row_label)} & " + " & ".join(cells_ar) + " \\\\"
            )

        if self.add_rows:
            lines.append("\\hline")

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            if key == "R-squared":
                disp = "R$^2$"
            elif key == "Adj. R-squared":
                disp = "Adj. R$^2$"
            else:
                disp = _latex_escape(disp)
            cells_s = [self._stat_cell(m, key) for m in all_models]
            lines.append(f"{disp} & " + " & ".join(cells_s) + " \\\\")

        lines.append("\\hline\\hline")

        # Notes
        note_line = f"{self._se_label()} in parentheses"
        lines.append(
            f"\\multicolumn{{{n_cols}}}{{l}}"
            f"{{\\footnotesize {_latex_escape(note_line)}}} \\\\"
        )
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            multi_note = f"{label} in {lo}…{hi}"
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(multi_note)}}} \\\\"
            )
        if self._has_any_eform():
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(self._eform_note())}}} \\\\"
            )
        if self.show_stars:
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(self._star_note())}}} \\\\"
            )
        for note in self.notes:
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(note)}}} \\\\"
            )

        lines.append("\\end{tabular}")
        lines.append("\\end{table}")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Markdown
    # ═══════════════════════════════════════════════════════════════════════

    def to_markdown(self, *, quarto: bool = False) -> str:
        """Render the table as Markdown.

        Parameters
        ----------
        quarto : bool, default False
            When ``True``, append a Quarto cross-reference caption block
            of the form ``: <caption> {#tbl-<label>}`` so the table can be
            referenced via ``@tbl-<label>`` in the manuscript. Requires
            ``quarto_label`` to have been set on the result (typically via
            ``regtable(..., quarto_label="main")``). Equivalent to calling
            :meth:`to_quarto`.
        """
        if quarto:
            return self.to_quarto()
        all_models = self._all_models_flat()
        lines: List[str] = []
        if self.title:
            lines.append(f"**{self.title}**")
            lines.append("")

        # Column spanners — Markdown has no native colspan, so we render
        # each spanner block as repeated label cells (one per column it
        # covers) inside a bold formatting band. Pandoc / GFM renderers
        # show this as a centered visual group.
        if self.column_spanners:
            spanner_cells: List[str] = []
            for lbl, span in self.column_spanners:
                spanner_cells.extend([f"**{lbl}**"] * span)
            lines.append(
                "| |"
                + "|".join(f" {c} " for c in spanner_cells)
                + "|"
            )

        # Header
        hdr = "| |" + "|".join(f" {n} " for n in self.model_labels) + "|"
        sep = "|---|" + "|".join("---:" for _ in self.model_labels) + "|"
        lines.append(hdr)
        lines.append(sep)

        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(
                    f"| **{self.panel_labels[pi]}** |"
                    + "|".join(" " for _ in self.model_labels)
                    + "|"
                )

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = self.coef_labels.get(var, var)
                cells: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells.append(self._coef_cell(m, var, flat_idx))
                        else:
                            cells.append("")
                        flat_idx += 1
                lines.append(f"| {label} |" + "|".join(f" {c} " for c in cells) + "|")
                # SE row
                cells2: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells2.append(self._se_cell(m, var, flat_idx))
                        else:
                            cells2.append("")
                        flat_idx += 1
                lines.append("| |" + "|".join(f" {c} " for c in cells2) + "|")
                # Extra SE rows from multi_se
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    cells3: List[str] = []
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                cells3.append(self._multi_se_cell(
                                    per_model, var, ext_idx, m, base_idx + off
                                ))
                            else:
                                cells3.append("")
                    lines.append("| |" + "|".join(f" {c} " for c in cells3) + "|")

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            cells_ar: List[str] = []
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                cells_ar.append(val)
            lines.append(
                f"| {row_label} |" + "|".join(f" {c} " for c in cells_ar) + "|"
            )

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            cells_s = [self._stat_cell(m, key) for m in all_models]
            lines.append(f"| {disp} |" + "|".join(f" {c} " for c in cells_s) + "|")

        lines.append("")
        lines.append(f"*{self._se_label()} in parentheses*")
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            lines.append(f"*{label} in {lo}…{hi}*")
        if self._has_any_eform():
            lines.append(f"*{self._eform_note()}*")
        if self.show_stars:
            lines.append(f"*{self._star_note()}*")
        for note in self.notes:
            lines.append(f"*{note}*")

        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Quarto
    # ═══════════════════════════════════════════════════════════════════════

    def to_quarto(self) -> str:
        """Render as a Quarto-cross-referenceable Markdown table.

        Builds on :meth:`to_markdown` and appends a Quarto caption block
        of the form::

            : <caption> {#tbl-<label>}

        which lets the manuscript reference the table via
        ``@tbl-<label>``. The ``tbl-`` prefix is auto-prepended when the
        user passes a bare ``quarto_label="main"``.

        Behaviour
        ---------
        - ``quarto_label`` is required. Without it, ``ValueError`` is
          raised — Quarto cross-references need an id.
        - ``quarto_caption`` falls back to ``title`` when not provided.
          If neither is set, a generic ``"Regression results"`` is used
          and a warning is emitted.
        - The leading title line is dropped (the caption block replaces
          it) to avoid duplicating the heading.
        """
        if not self.quarto_label:
            raise ValueError(
                "to_quarto() requires quarto_label to be set. "
                "Pass quarto_label='main' (or similar) to regtable()."
            )

        raw_label = str(self.quarto_label).strip()
        label = raw_label if raw_label.startswith("tbl-") else f"tbl-{raw_label}"

        if self.quarto_caption:
            caption = str(self.quarto_caption)
        elif self.title:
            caption = str(self.title)
        else:
            warnings.warn(
                "to_quarto(): no quarto_caption or title provided; "
                "using default 'Regression results'. Quarto cross-refs "
                "render better with an explicit caption.",
                UserWarning,
                stacklevel=2,
            )
            caption = "Regression results"

        saved_title = self.title
        try:
            self.title = None
            body = self.to_markdown()
        finally:
            self.title = saved_title

        body = body.rstrip()
        return f"{body}\n\n: {caption} {{#{label}}}\n"

    # ═══════════════════════════════════════════════════════════════════════
    # DataFrame
    # ═══════════════════════════════════════════════════════════════════════

    def to_dataframe(self) -> pd.DataFrame:
        """Return the table as a pandas DataFrame."""
        all_models = self._all_models_flat()
        records: List[Dict[str, str]] = []

        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                row_ph: Dict[str, str] = {"": self.panel_labels[pi]}
                for n in self.model_labels:
                    row_ph[n] = ""
                records.append(row_ph)

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = self.coef_labels.get(var, var)
                row: Dict[str, str] = {"": label}
                mi = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        col_name = self.model_labels[mi]
                        if gi == pi:
                            row[col_name] = self._coef_cell(m, var, mi)
                        else:
                            row[col_name] = ""
                        mi += 1
                records.append(row)
                # SE row
                row2: Dict[str, str] = {"": ""}
                mi = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        col_name = self.model_labels[mi]
                        if gi == pi:
                            row2[col_name] = self._se_cell(m, var, mi)
                        else:
                            row2[col_name] = ""
                        mi += 1
                records.append(row2)
                # Extra SE rows (multi_se)
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    row3: Dict[str, str] = {"": ""}
                    mi = 0
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            col_name = self.model_labels[mi]
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                row3[col_name] = self._multi_se_cell(
                                    per_model, var, ext_idx, m, base_idx + off
                                )
                            else:
                                row3[col_name] = ""
                            mi += 1
                    records.append(row3)

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            row_ar: Dict[str, str] = {"": row_label}
            for i, lbl in enumerate(self.model_labels):
                row_ar[lbl] = row_vals[i] if i < len(row_vals) else ""
            records.append(row_ar)

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            row_s: Dict[str, str] = {"": disp}
            for i, m in enumerate(all_models):
                row_s[self.model_labels[i]] = self._stat_cell(m, key)
            records.append(row_s)

        df = pd.DataFrame(records)
        df = df.set_index("")
        df.index.name = None
        return df

    # ═══════════════════════════════════════════════════════════════════════
    # Excel
    # ═══════════════════════════════════════════════════════════════════════

    def to_excel(self, filename: str) -> None:
        """Export table to Excel as a strict book-tab three-line table.

        Uses the shared ``_excel_style`` primitives so the visual output
        is byte-aligned with ``sumstats``, ``tab``, ``paper_tables``,
        ``collection``, ``modelsummary`` and ``outreg2``: thick top rule
        above the column header, thin mid rule between header and body,
        thick bottom rule below the last data row, Times New Roman
        throughout.
        """
        try:
            import openpyxl
            from openpyxl.styles import Alignment, Font
        except ImportError:
            warnings.warn(
                "openpyxl is required for Excel export. "
                "Install with: pip install openpyxl"
            )
            return

        from ._excel_style import (
            BODY_PT, HEADER_PT, NOTES_PT, TIMES,
            apply_booktab_borders, autofit_columns, write_title,
        )

        df = self.to_dataframe()
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Regression Table"

        header_font = Font(bold=True, name=TIMES, size=HEADER_PT)
        body_font = Font(name=TIMES, size=BODY_PT)
        notes_font = Font(italic=True, name=TIMES, size=NOTES_PT)
        center = Alignment(horizontal="center")
        left = Alignment(horizontal="left")

        n_cols = len(df.columns) + 1  # +1 for the row-label column
        row_idx = 1
        if self.title:
            row_idx = write_title(ws, row_idx, n_cols, self.title)

        # Header row
        header_top_row = row_idx
        c0 = ws.cell(row=row_idx, column=1, value="")
        c0.font = header_font
        for j, col in enumerate(df.columns, 2):
            cell = ws.cell(row=row_idx, column=j, value=col)
            cell.font = header_font
            cell.alignment = center
        header_bot_row = row_idx
        row_idx += 1

        # Body rows
        body_top_row = row_idx
        for idx, row_data in df.iterrows():
            c0 = ws.cell(row=row_idx, column=1, value=str(idx))
            c0.font = body_font
            c0.alignment = left
            for j, val in enumerate(row_data, 2):
                cell = ws.cell(row=row_idx, column=j, value=str(val))
                cell.font = body_font
                cell.alignment = center
            row_idx += 1
        body_bot_row = row_idx - 1

        apply_booktab_borders(
            ws,
            header_top_row=header_top_row,
            header_bot_row=header_bot_row,
            body_top_row=body_top_row,
            body_bot_row=body_bot_row,
            n_cols=n_cols,
        )

        # Notes — emit the same lines that to_text/to_html/to_latex/to_word
        # emit so users who pass multi_se / repro / notes do not lose them
        # when exporting to Excel.
        note_row = body_bot_row + 1
        ws.cell(
            row=note_row, column=1,
            value=f"{self._se_label()} in parentheses"
        ).font = notes_font
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            note_row += 1
            ws.cell(row=note_row, column=1,
                    value=f"{label} in {lo}…{hi}").font = notes_font
        if self.show_stars:
            note_row += 1
            ws.cell(row=note_row, column=1,
                    value=self._star_note()).font = notes_font
        for note in self.notes:
            note_row += 1
            ws.cell(row=note_row, column=1, value=note).font = notes_font

        autofit_columns(ws, n_cols, max_width=25)
        wb.save(filename)

    # ═══════════════════════════════════════════════════════════════════════
    # Word (docx)
    # ═══════════════════════════════════════════════════════════════════════

    def to_word(self, filename: str) -> None:
        """Export table to Word (.docx) file in AER/QJE book-tab style.

        The exported document follows economics-journal conventions:
        a heavy top rule, thin mid rule below the header, heavy bottom
        rule above notes, and **no** internal vertical borders. Body
        text is Times New Roman 10pt; the notes paragraph is 8pt italic.
        """
        try:
            from docx import Document
        except ImportError:
            warnings.warn(
                "python-docx is required for Word export. "
                "Install with: pip install python-docx"
            )
            return

        from ._aer_style import (
            apply_word_booktab_rules,
            style_word_table_typography,
            add_word_notes_paragraph,
        )

        doc = Document()
        if self.title:
            doc.add_heading(self.title, level=2)

        df = self.to_dataframe()
        n_rows = len(df) + 1
        n_cols = len(df.columns) + 1
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.autofit = True

        # Populate header
        header_row = table.rows[0]
        header_row.cells[0].text = ""
        for j, col in enumerate(df.columns, 1):
            header_row.cells[j].text = str(col)
        # Populate body
        for i, (idx, row_data) in enumerate(df.iterrows(), 1):
            table.rows[i].cells[0].text = str(idx)
            for j, val in enumerate(row_data, 1):
                table.rows[i].cells[j].text = str(val)

        style_word_table_typography(table, header_rows=(0,))
        apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)

        # Notes (italic, 8pt)
        note_lines = [f"{self._se_label()} in parentheses"]
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            note_lines.append(f"{label} in {lo}…{hi}")
        if self._has_any_eform():
            note_lines.append(self._eform_note())
        if self.show_stars:
            note_lines.append(self._star_note())
        note_lines.extend(self.notes)
        add_word_notes_paragraph(doc, "\n".join(note_lines))

        doc.save(filename)

    # ═══════════════════════════════════════════════════════════════════════
    # Save (auto-detect from extension)
    # ═══════════════════════════════════════════════════════════════════════

    def save(self, filename: str) -> None:
        """Auto-detect format from file extension and save."""
        path = Path(filename)
        ext = path.suffix.lower()

        if ext in (".xlsx", ".xls"):
            self.to_excel(filename)
        elif ext == ".docx":
            self.to_word(filename)
        elif ext == ".tex":
            path.write_text(self.to_latex(), encoding="utf-8")
        elif ext in (".html", ".htm"):
            path.write_text(self.to_html(), encoding="utf-8")
        elif ext == ".md":
            path.write_text(self.to_markdown(), encoding="utf-8")
        elif ext == ".qmd":
            path.write_text(self.to_quarto(), encoding="utf-8")
        elif ext == ".csv":
            self.to_dataframe().to_csv(filename)
        else:
            path.write_text(self.to_text(), encoding="utf-8")

    # ═══════════════════════════════════════════════════════════════════════
    # Dunder methods
    # ═══════════════════════════════════════════════════════════════════════

    def _render(self, fmt: str) -> str:
        return {
            "text": self.to_text,
            "latex": self.to_latex,
            "tex": self.to_latex,
            "html": self.to_html,
            "markdown": self.to_markdown,
            "md": self.to_markdown,
            "quarto": self.to_quarto,
            "qmd": self.to_quarto,
        }.get(fmt, self.to_text)()

    def __str__(self) -> str:
        return self._render(self._output)

    def __repr__(self) -> str:
        return self.__str__()

    def _repr_html_(self) -> str:
        return self.to_html()


# ═══════════════════════════════════════════════════════════════════════════
# _PanelData: lightweight container for a group of models
# ═══════════════════════════════════════════════════════════════════════════

class _PanelData:
    __slots__ = ("models",)

    def __init__(self, models: List[_ModelData]):
        self.models = models


# ═══════════════════════════════════════════════════════════════════════════
# regtable() — the main public API
# ═══════════════════════════════════════════════════════════════════════════

def regtable(
    *args,
    panel_labels: Optional[List[str]] = None,
    coef_labels: Optional[Dict[str, str]] = None,
    dep_var_labels: Optional[List[str]] = None,
    model_labels: Optional[List[str]] = None,
    keep: Optional[Sequence[str]] = None,
    drop: Optional[Sequence[str]] = None,
    order: Optional[Sequence[str]] = None,
    stats: Optional[Sequence[str]] = None,
    se_type: str = "se",
    stars: bool = True,
    star_levels: Optional[Tuple[float, ...]] = None,
    fmt: str = "%.3f",
    output: str = "text",
    filename: Optional[str] = None,
    title: Optional[str] = None,
    notes: Optional[List[str]] = None,
    add_rows: Optional[Dict[str, List[str]]] = None,
    alpha: float = 0.05,
    template: Optional[str] = None,
    diagnostics: Union[str, bool] = "auto",
    multi_se: Optional[Dict[str, Sequence[Any]]] = None,
    repro: Union[bool, Dict[str, Any], None] = None,
    quarto_label: Optional[str] = None,
    quarto_caption: Optional[str] = None,
    eform: Union[bool, Sequence[bool]] = False,
    column_spanners: Optional[Sequence[Tuple[str, int]]] = None,
    coef_map: Optional[Dict[str, str]] = None,
    consistency_check: bool = True,
) -> RegtableResult:
    """
    Unified publication-quality regression table.

    Accepts model results as positional arguments. If the first argument
    is a list, each list is treated as a separate panel.

    Parameters
    ----------
    *args : model results or lists of model results
        ``EconometricResults``, ``CausalResult``, or any duck-typed object
        with ``params`` / ``std_errors`` attributes. Pass multiple lists
        to create a multi-panel table.
    panel_labels : list of str, optional
        Labels for each panel (e.g., ``["Panel A: Wages", "Panel B: Hours"]``).
    coef_labels : dict, optional
        Rename variables: ``{"education": "Years of Education"}``.
    dep_var_labels : list of str, optional
        Dependent variable labels shown below column headers.
    model_labels : list of str, optional
        Column header labels. Defaults to ``(1), (2), ...``.
    keep : list of str, optional
        Only show these variables.
    drop : list of str, optional
        Hide these variables.
    order : list of str, optional
        Reorder variables.
    stats : list of str, optional
        Summary statistics. Defaults to ``["N", "R2", "adj_R2", "F"]``.
    se_type : str, default ``"se"``
        What to show beneath coefficients: ``"se"``, ``"t"``, ``"p"``,
        or ``"ci"`` for confidence intervals.
    stars : bool, default True
        Append significance stars.
    star_levels : tuple, default ``(0.10, 0.05, 0.01)``
        Thresholds for ``*``, ``**``, ``***``.
    fmt : str, default ``"%.3f"``
        Format string for numeric values. Pass any C-style format
        (``"%.0f"``, ``"%.4f"``, ...) for fixed precision, or
        ``"auto"`` for magnitude-adaptive precision (recommended when
        a single table mixes dollar-magnitude coefficients like
        ``1521`` with elasticity-magnitude coefficients like ``0.288``
        — fixed ``"%.0f"`` would round the latter to ``0``).
    output : str, default ``"text"``
        Controls what ``str(result)`` / ``repr(result)`` / ``print(result)``
        returns — one of ``"text"``, ``"latex"``, ``"html"``, ``"markdown"``,
        ``"quarto"``, ``"word"``, ``"excel"``. In Jupyter, ``_repr_html_``
        always renders HTML regardless of this setting.
    filename : str, optional
        Save the table to this file path. The format is chosen from the
        **file extension** (``.tex``/``.html``/``.md``/``.qmd``/``.docx``/``.xlsx``),
        independently of ``output=``. Pass a matching extension and
        ``output=`` to avoid surprises.
    quarto_label : str, optional
        Quarto cross-reference id. Pass ``"main"`` to make the table
        referenceable as ``@tbl-main`` from the manuscript prose. The
        ``tbl-`` prefix is auto-prepended when missing. Required for
        ``to_quarto()`` and ``output="quarto"``.
    quarto_caption : str, optional
        Caption rendered alongside the Quarto cross-ref id. Falls back
        to ``title`` when omitted; if both are absent, a generic
        ``"Regression results"`` is used and a warning is emitted.
    title : str, optional
        Table title / caption.
    notes : list of str, optional
        Additional notes beneath the table.
    add_rows : dict, optional
        Custom rows: ``{"Controls": ["No", "Yes", "Yes"]}``. User-provided
        rows take precedence over auto-extracted diagnostic rows with the
        same label.
    alpha : float, default 0.05
        Significance level used when ``se_type='ci'``. Displayed CI is
        ``(1 - alpha) * 100``%. With ``alpha=0.05`` (default) the bounds
        come from the model's stored 95% CI; for any other ``alpha`` the
        bounds are recomputed as ``b ± crit · se``, using the
        t-distribution when ``df_resid`` is known, else the standard
        normal.
    template : str, optional
        Journal preset name. One of ``"aer"``, ``"qje"``, ``"econometrica"``,
        ``"restat"``, ``"jf"``, ``"aeja"``, ``"jpe"``, ``"restud"``. When
        set, fills in defaults for ``star_levels``, the SE-row footer label
        (e.g. QJE → "Robust standard errors"), the default ``stats``
        selection (e.g. JF/AEJA include Adj. R²), and any extra notes —
        but every explicit kwarg you pass still wins. See
        :data:`statspai.output._journals.JOURNALS`.
    diagnostics : {'auto', 'off'} or bool, default ``'auto'``
        Auto-extract publication-quality diagnostic rows from the result
        objects:

        - **FE / Cluster indicators** — one row per distinct fixed effect
          variable (AER style: ``"Firm FE: Yes/No"``, ``"Year FE: Yes/No"``;
          interactions render as ``"Firm × Year FE"``), plus
          ``"Cluster SE: <var>"``. Falls back to a single
          ``"Fixed Effects: Yes/No"`` row when FE metadata is present but
          unparseable.
        - **IV** — first-stage F (Olea-Pflueger / KP), Hansen-J p.
        - **DiD** — pre-trend p-value, treated-group count.
        - **RD** — bandwidth, kernel, polynomial order.

        ``"auto"`` (and ``True``) emit only rows where at least one column
        produces a non-empty cell; ``False`` / ``"off"`` disables all
        auto-extraction. User-supplied ``add_rows`` always override.
    multi_se : dict, optional
        Stack additional SE specifications under the primary SE row.
        Keys are display labels (e.g. ``"Cluster SE"``, ``"Bootstrap SE"``)
        and values are sequences of :class:`pandas.Series` or dicts
        (one per model column) mapping coefficient names to SE values.
        Bracket styles cycle ``[]``/``{}``/``⟨⟩``/``||``. Footer notes
        record each label automatically.
    repro : bool or dict, optional
        Append a reproducibility metadata note (StatsPAI version, optional
        seed and data hash, timestamp) as the last footer line. ``True``
        emits the version + timestamp only. Pass a dict to record more:
        ``{"data": df, "seed": 42, "extra": "git@<sha>"}``.
    eform : bool or list of bool, default ``False``
        Report exponentiated coefficients — odds ratios for ``logit`` /
        ``probit``, incidence-rate ratios for ``poisson``, hazard ratios
        for Cox-style models. Standard errors use the delta method
        (``exp(b)·SE(b)``), CI bounds are ``(exp(lo), exp(hi))`` of the
        original endpoints, and t / p values are unchanged because
        ``H_0: b=0`` is equivalent to ``H_0: exp(b)=1``. Pass a per-model
        list (length matches ``n_models``) to mix transformed and
        untransformed columns (e.g. logit + OLS in the same table). A
        footer note transparently flags which columns are exponentiated.
    column_spanners : list of (label, span), optional
        Multi-row header above the model labels — each tuple groups
        ``span`` consecutive columns under ``label``. Spans must
        partition all model columns (sum equals ``n_models``). Renders
        as ``\\multicolumn{n}{c}{label}`` + ``\\cmidrule`` in LaTeX,
        ``colspan="n"`` in HTML, repeated bold cells in Markdown,
        and centered ASCII in text. Word and Excel exports inherit
        ``to_dataframe()``'s flat column model and currently omit the
        spanner row — use the LaTeX or HTML output for paper-grade
        spanners. Mirrors Stata ``mgroups()`` and R ``modelsummary``'s
        ``group`` argument. Example: ``column_spanners=[("OLS", 2),
        ("IV", 2)]`` over four models.
    coef_map : dict, optional
        Single-shot rename + reorder + drop. Mirrors R
        ``modelsummary``'s ``coef_map``: pass an ordered dict whose
        keys are coefficient names to **keep** (in display order) and
        values are the rendered labels. Variables not in ``coef_map``
        are dropped. Mutually exclusive with ``coef_labels`` /
        ``keep`` / ``drop`` / ``order`` — pass either the unified map
        or the legacy four-parameter spec.
    consistency_check : bool, default True
        When two or more columns are passed and their sample sizes
        differ, emit a ``UserWarning``. Reviewer red flag — disable by
        setting ``False`` (or annotate with ``notes=[...]``) when the
        N-mismatch is intentional (IV first stage on a subsample,
        RD bandwidth restriction, etc.).

    Returns
    -------
    RegtableResult
        Object with ``.to_text()``, ``.to_latex()``, ``.to_html()``,
        ``.to_markdown()``, ``.to_excel(filename)``, ``.to_word(filename)``,
        ``.to_dataframe()``, ``.save(filename)`` methods.
        Renders as rich HTML in Jupyter notebooks via ``_repr_html_()``.

    Examples
    --------
    >>> import statspai as sp
    >>> m1 = sp.regress("y ~ x1", data=df)
    >>> m2 = sp.regress("y ~ x1 + x2", data=df)
    >>> sp.regtable(m1, m2)
    >>> sp.regtable(m1, m2, output="latex", filename="table1.tex")
    >>> sp.regtable([m1, m2], [m3, m4],
    ...     panel_labels=["Panel A: OLS", "Panel B: IV"])
    >>>
    >>> # Logit odds ratios
    >>> sp.regtable(sp.logit("y ~ x", data=df), eform=True)
    >>>
    >>> # IV three-block table with column spanners
    >>> sp.regtable(
    ...     ols1, ols2, iv1, iv2,
    ...     column_spanners=[("OLS", 2), ("IV", 2)],
    ...     stats=["N", "R2", "depvar_mean", "depvar_sd"],
    ... )
    >>>
    >>> # Unified coef_map (rename + order + drop in one shot)
    >>> sp.regtable(m1, m2, coef_map={
    ...     "x2": "Education",
    ...     "x1": "Experience",
    ...     "Intercept": "Constant",
    ... })
    """
    if not args:
        raise ValueError("At least one model result is required.")

    _VALID_OUTPUTS = {
        "text", "latex", "tex", "html", "markdown", "md",
        "quarto", "qmd", "word", "excel",
    }
    if output not in _VALID_OUTPUTS:
        raise ValueError(
            f"output={output!r} is invalid. Must be one of: "
            f"{sorted(_VALID_OUTPUTS)}"
        )

    # --- Resolve journal template (sets defaults; explicit kwargs win) ---
    se_label_override: Optional[str] = None
    template_notes: List[str] = []
    if template is not None:
        preset = get_template(template)
        if star_levels is None:
            star_levels = tuple(preset["star_levels"])
        if stats is None:
            stats = list(preset["stats"])
        if se_type == "se":
            se_label_override = preset.get("se_label")
        # Footer notes from the template are appended *after* user notes,
        # except we skip the boilerplate "stars" / "SE in parentheses"
        # lines because the renderer emits those itself.
        for line in preset.get("notes_default", ()):
            low = line.lower()
            if "in parenthes" in low or "p<0" in low:
                continue
            template_notes.append(line)

    if star_levels is None:
        star_levels = (0.10, 0.05, 0.01)

    # --- coef_map shortcut (mirrors R modelsummary's three-in-one) ----
    # When set, it simultaneously renames + reorders + drops via a single
    # ordered dict. Conflicts with the legacy keep/drop/order/coef_labels
    # parameters are rejected up front because resolving them is ambiguous
    # and silent precedence would surprise users.
    if coef_map is not None:
        if coef_labels is not None:
            raise ValueError(
                "Pass either coef_map or coef_labels, not both. coef_map is "
                "the unified shortcut (rename + order + drop); coef_labels "
                "only renames."
            )
        if keep is not None or drop is not None or order is not None:
            raise ValueError(
                "coef_map already defines the keep / order behaviour "
                "(via its insertion order and key set). Drop the explicit "
                "keep/drop/order arguments when using coef_map."
            )
        coef_labels = dict(coef_map)
        keep = list(coef_map.keys())
        order = list(coef_map.keys())

    # --- Detect panel structure ---
    # If first arg is a list, treat each positional arg as a panel
    if isinstance(args[0], list):
        raw_panels = list(args)
        flat_results = [r for raw in raw_panels for r in raw]
    else:
        raw_panels = [list(args)]
        flat_results = list(args)

    # Extract model data per panel
    panels: List[_PanelData] = []
    total_models = 0
    for raw in raw_panels:
        model_data_list = [_extract_model_data(r) for r in raw]
        panels.append(_PanelData(model_data_list))
        total_models += len(model_data_list)

    # --- Resolve eform flags (one bool per flat model position) -------
    if isinstance(eform, bool):
        eform_flags = [eform] * total_models
    else:
        eform_seq = list(eform)
        if len(eform_seq) != total_models:
            raise ValueError(
                f"eform list has {len(eform_seq)} entries but there are "
                f"{total_models} models."
            )
        eform_flags = [bool(f) for f in eform_seq]

    # --- Consistency checks: warn on N-mismatch -----------------------
    # Mixed sample sizes across columns is a Reviewer red flag. We don't
    # *block* (sometimes mixing is intentional — IV first stage on a
    # subsample, RD bandwidth restriction); we warn so the user puts an
    # explicit footnote.
    if consistency_check and total_models >= 2:
        ns = []
        for p in panels:
            for m in p.models:
                n_val = m.stats.get("N")
                if n_val is not None:
                    try:
                        ns.append(int(n_val))
                    except (TypeError, ValueError):
                        pass
        if len(ns) >= 2 and (max(ns) - min(ns) > 0):
            warnings.warn(
                f"Sample sizes differ across columns (range "
                f"{min(ns):,}–{max(ns):,}). If this is intentional (e.g. "
                f"IV first stage on a subsample), add a footnote via "
                f"notes=[...]; otherwise re-fit on a common sample.",
                UserWarning,
                stacklevel=2,
            )

    # Default model labels
    if model_labels is None:
        model_labels = [f"({i + 1})" for i in range(total_models)]
    elif len(model_labels) != total_models:
        raise ValueError(
            f"model_labels has {len(model_labels)} entries but "
            f"there are {total_models} models."
        )

    # Validate dep_var_labels length
    if dep_var_labels is not None and len(dep_var_labels) != total_models:
        raise ValueError(
            f"dep_var_labels has {len(dep_var_labels)} entries but "
            f"there are {total_models} models."
        )

    # --- Auto-extract diagnostic rows ----------------------------------
    if diagnostics in (False, "off"):
        auto_rows: Dict[str, List[str]] = {}
    else:
        auto_rows = dict(extract_diagnostic_rows(flat_results))

    # Merge: user's add_rows wins on collisions, auto rows fill gaps.
    # Backwards-compat: if user supplies the legacy "Fixed Effects" row,
    # suppress the auto-emitted per-variable FE rows ("Firm FE", "Year FE",
    # …) so old tables don't suddenly show a single row + auto rows stacked.
    user_rows = dict(add_rows) if add_rows else {}
    if "Fixed Effects" in user_rows:
        auto_rows = {k: v for k, v in auto_rows.items() if not k.endswith(" FE")}
    merged_add_rows: Dict[str, List[str]] = {}
    for label, vals in auto_rows.items():
        if label not in user_rows:
            merged_add_rows[label] = list(vals)
    for label, vals in user_rows.items():
        merged_add_rows[label] = list(vals)

    # --- Resolve multi_se -----------------------------------------------
    multi_se_norm = _resolve_multi_se(multi_se, total_models)

    # --- Resolve reproducibility note -----------------------------------
    final_notes = list(notes) if notes else []
    final_notes.extend(template_notes)
    if repro:
        repro_kwargs = dict(repro) if isinstance(repro, dict) else {}
        repro_note = build_repro_note(**repro_kwargs)
        if repro_note:
            final_notes.append(repro_note)

    result = RegtableResult(
        panels=panels,
        panel_labels=panel_labels,
        model_labels=model_labels,
        dep_var_labels=dep_var_labels,
        coef_labels=coef_labels,
        keep=list(keep) if keep else None,
        drop=list(drop) if drop else None,
        order=list(order) if order else None,
        se_type=se_type,
        stars=stars,
        star_levels=tuple(star_levels),
        fmt=fmt,
        title=title,
        notes=final_notes,
        add_rows=merged_add_rows,
        stats=list(stats) if stats else None,
        output=output,
        alpha=alpha,
        multi_se=multi_se_norm,
        se_label=se_label_override,
        template=template,
        quarto_label=quarto_label,
        quarto_caption=quarto_caption,
        eform_flags=eform_flags,
        column_spanners=list(column_spanners) if column_spanners else None,
    )

    # --- Output handling ---
    # Do NOT auto-print: Jupyter renders via _repr_html_, REPL via __repr__.
    # Scripts that want the rendered text should `print(regtable(...))`.
    if filename:
        result.save(filename)
    elif output in ("word", "excel"):
        warnings.warn(
            f"output='{output}' requires a filename. "
            f"Use filename='table.{'docx' if output == 'word' else 'xlsx'}'"
        )

    return result


# ═══════════════════════════════════════════════════════════════════════════
# MeanComparisonResult
# ═══════════════════════════════════════════════════════════════════════════

class MeanComparisonResult:
    """Rich result object for balance / mean comparison tables."""

    def __init__(
        self,
        data: pd.DataFrame,
        variables: List[str],
        group: str,
        group_labels: Tuple[str, str],
        test: str,
        fmt: str,
        title: str,
        weights: Optional[str],
        output: str = "text",
    ):
        self.variables = variables
        self.group = group
        self.group_labels = group_labels
        self.test = test
        self.fmt = fmt
        self.title = title
        self.weights = weights
        self._output = output

        # Compute all stats
        self._compute(data)

    def _compute(self, data: pd.DataFrame) -> None:
        g0_label, g1_label = self.group_labels
        mask0 = data[self.group] == 0
        mask1 = data[self.group] == 1

        # If not binary 0/1, use the two unique values
        unique_vals = sorted(data[self.group].dropna().unique())
        if len(unique_vals) == 2:
            mask0 = data[self.group] == unique_vals[0]
            mask1 = data[self.group] == unique_vals[1]

        d0 = data.loc[mask0]
        d1 = data.loc[mask1]
        self.n0 = len(d0)
        self.n1 = len(d1)

        rows: List[Dict[str, Any]] = []
        for var in self.variables:
            v0 = d0[var].dropna()
            v1 = d1[var].dropna()

            mean0, sd0 = v0.mean(), v0.std()
            mean1, sd1 = v1.mean(), v1.std()
            diff = mean1 - mean0

            # Test
            if self.test == "ranksum":
                try:
                    stat, pval = sp_stats.mannwhitneyu(v0, v1, alternative="two-sided")
                except Exception:
                    pval = np.nan
            elif self.test == "chi2":
                try:
                    ct = pd.crosstab(data[self.group], data[var])
                    chi2, pval, _, _ = sp_stats.chi2_contingency(ct)
                except Exception:
                    pval = np.nan
            else:
                # Default: Welch t-test
                try:
                    stat, pval = sp_stats.ttest_ind(v0, v1, equal_var=False)
                except Exception:
                    pval = np.nan

            stars = _format_stars(pval, (0.10, 0.05, 0.01))

            rows.append({
                "variable": var,
                "mean0": mean0,
                "sd0": sd0,
                "mean1": mean1,
                "sd1": sd1,
                "diff": diff,
                "pvalue": pval,
                "stars": stars,
            })

        self._rows = rows

    # ═══════════════════════════════════════════════════════════════════════
    # TEXT
    # ═══════════════════════════════════════════════════════════════════════

    def to_text(self) -> str:
        g0_label, g1_label = self.group_labels
        var_w = max(max(len(v) for v in self.variables), 10) + 2
        col_w = 16
        total_w = var_w + col_w * 4 + 2

        thick = "\u2501" * total_w
        thin = "\u2500" * total_w
        lines: List[str] = []

        lines.append(f"  {self.title}")
        lines.append(thick)

        # Header
        hdr = f"{'':>{var_w}}"
        hdr += f"{g0_label:>{col_w}}"
        hdr += f"{g1_label:>{col_w}}"
        hdr += f"{'Diff':>{col_w}}"
        hdr += f"{'p-value':>{col_w}}"
        lines.append(hdr)

        sub = f"{'':>{var_w}}"
        sub += f"{'Mean (SD)':>{col_w}}"
        sub += f"{'Mean (SD)':>{col_w}}"
        sub += " " * col_w
        sub += " " * col_w
        lines.append(sub)

        lines.append(thin)

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            fmt_diff = (self.fmt % row["diff"]) + row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            col0 = f"{fmt_mean0} ({fmt_sd0})"
            col1 = f"{fmt_mean1} ({fmt_sd1})"

            line = f"{row['variable']:<{var_w}}"
            line += f"{col0:>{col_w}}"
            line += f"{col1:>{col_w}}"
            line += f"{fmt_diff:>{col_w}}"
            line += f"{fmt_pval:>{col_w}}"
            lines.append(line)

        lines.append(thin)
        lines.append(f"{'N':<{var_w}}{self.n0:>{col_w},}{self.n1:>{col_w},}")
        lines.append(thick)
        lines.append(f"* p<0.10, ** p<0.05, *** p<0.01")

        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # HTML
    # ═══════════════════════════════════════════════════════════════════════

    def to_html(self) -> str:
        g0_label, g1_label = self.group_labels
        lines: List[str] = []
        lines.append(
            '<table class="balance-table" style="border-collapse:collapse; '
            'font-family:\'Times New Roman\', serif; font-size:13px;">'
        )
        lines.append(
            f'<caption style="font-weight:bold; font-size:14px; '
            f'margin-bottom:8px;">{_html_escape(self.title)}</caption>'
        )

        # Header
        lines.append("<thead>")
        lines.append("<tr>")
        for h in ["", g0_label, g1_label, "Diff", "p-value"]:
            border = "border-top:3px solid black; border-bottom:1px solid black; "
            align = "text-align:center;" if h else "text-align:left;"
            lines.append(f'<th style="{border}{align} padding:4px 10px;">{_html_escape(h)}</th>')
        lines.append("</tr>")
        lines.append("<tr>")
        lines.append('<th style="text-align:left;"></th>')
        lines.append(f'<th style="text-align:center; font-weight:normal; font-size:11px;">Mean (SD)</th>')
        lines.append(f'<th style="text-align:center; font-weight:normal; font-size:11px;">Mean (SD)</th>')
        lines.append('<th></th><th></th>')
        lines.append("</tr>")
        lines.append("</thead>")
        lines.append("<tbody>")

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            fmt_diff = (self.fmt % row["diff"]) + row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            lines.append("<tr>")
            lines.append(f'<td style="text-align:left; padding:2px 10px;">{_html_escape(row["variable"])}</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_mean0} ({fmt_sd0})</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_mean1} ({fmt_sd1})</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_diff}</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_pval}</td>')
            lines.append("</tr>")

        # N row
        lines.append(
            f'<tr><td colspan="5" style="border-top:1px solid black; padding:0;"></td></tr>'
        )
        lines.append("<tr>")
        lines.append(f'<td style="text-align:left; padding:2px 10px;">N</td>')
        lines.append(f'<td style="text-align:center; padding:2px 10px;">{self.n0:,}</td>')
        lines.append(f'<td style="text-align:center; padding:2px 10px;">{self.n1:,}</td>')
        lines.append('<td></td><td></td>')
        lines.append("</tr>")
        lines.append(
            f'<tr><td colspan="5" style="border-top:3px solid black; padding:0;"></td></tr>'
        )

        lines.append("</tbody>")
        lines.append("<tfoot>")
        lines.append(
            f'<tr><td colspan="5" style="text-align:left; font-size:11px; padding:4px 10px;">'
            f'* p&lt;0.10, ** p&lt;0.05, *** p&lt;0.01</td></tr>'
        )
        lines.append("</tfoot>")
        lines.append("</table>")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # LaTeX
    # ═══════════════════════════════════════════════════════════════════════

    def to_latex(self) -> str:
        g0_label, g1_label = self.group_labels
        lines: List[str] = []
        lines.append("\\begin{table}[htbp]")
        lines.append("\\centering")
        lines.append(f"\\caption{{{_latex_escape(self.title)}}}")
        lines.append("\\begin{tabular}{lcccc}")
        lines.append("\\hline\\hline")
        lines.append(
            f" & {_latex_escape(g0_label)} & {_latex_escape(g1_label)} "
            f"& Diff & p-value \\\\"
        )
        lines.append(
            " & Mean (SD) & Mean (SD) & & \\\\"
        )
        lines.append("\\hline")

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            diff_str = self.fmt % row["diff"]
            stars_str = row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            var_esc = _latex_escape(row["variable"])
            lines.append(
                f"{var_esc} & {fmt_mean0} ({fmt_sd0}) & "
                f"{fmt_mean1} ({fmt_sd1}) & "
                f"{diff_str}{stars_str} & {fmt_pval} \\\\"
            )

        lines.append("\\hline")
        lines.append(f"N & {self.n0:,} & {self.n1:,} & & \\\\")
        lines.append("\\hline\\hline")
        lines.append(
            "\\multicolumn{5}{l}{\\footnotesize * p<0.10, ** p<0.05, *** p<0.01} \\\\"
        )
        lines.append("\\end{tabular}")
        lines.append("\\end{table}")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Markdown
    # ═══════════════════════════════════════════════════════════════════════

    def to_markdown(self) -> str:
        g0_label, g1_label = self.group_labels
        lines: List[str] = []
        lines.append(f"**{self.title}**")
        lines.append("")
        lines.append(f"| | {g0_label} | {g1_label} | Diff | p-value |")
        lines.append("|---|---:|---:|---:|---:|")

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            fmt_diff = (self.fmt % row["diff"]) + row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            lines.append(
                f"| {row['variable']} | {fmt_mean0} ({fmt_sd0}) | "
                f"{fmt_mean1} ({fmt_sd1}) | {fmt_diff} | {fmt_pval} |"
            )

        lines.append(f"| N | {self.n0:,} | {self.n1:,} | | |")
        lines.append("")
        lines.append("*\\* p<0.10, \\*\\* p<0.05, \\*\\*\\* p<0.01*")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # DataFrame
    # ═══════════════════════════════════════════════════════════════════════

    def to_dataframe(self) -> pd.DataFrame:
        g0_label, g1_label = self.group_labels
        records = []
        for row in self._rows:
            records.append({
                "Variable": row["variable"],
                f"{g0_label} Mean": row["mean0"],
                f"{g0_label} SD": row["sd0"],
                f"{g1_label} Mean": row["mean1"],
                f"{g1_label} SD": row["sd1"],
                "Difference": row["diff"],
                "p-value": row["pvalue"],
                "Significance": row["stars"],
            })
        df = pd.DataFrame(records)
        df = df.set_index("Variable")
        return df

    # ═══════════════════════════════════════════════════════════════════════
    # Excel / Word / Save
    # ═══════════════════════════════════════════════════════════════════════

    def to_excel(self, filename: str) -> None:
        """Export balance table to Excel as a book-tab three-line table."""
        try:
            import openpyxl  # noqa: F401
        except ImportError:
            warnings.warn("openpyxl required for Excel export: pip install openpyxl")
            return
        from ._excel_style import render_dataframe_to_xlsx
        render_dataframe_to_xlsx(
            self.to_dataframe(),
            filename,
            title=getattr(self, "title", None),
            sheet_name="Balance Table",
            index_label="Variable",
        )

    def to_word(self, filename: str) -> None:
        """Export balance table to Word in AER/QJE book-tab style."""
        try:
            from docx import Document
        except ImportError:
            warnings.warn("python-docx required for Word export: pip install python-docx")
            return

        from ._aer_style import (
            apply_word_booktab_rules,
            style_word_table_typography,
            add_word_notes_paragraph,
        )

        doc = Document()
        doc.add_heading(self.title, level=2)
        df = self.to_dataframe().reset_index()
        n_rows = len(df) + 1
        n_cols = len(df.columns)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.autofit = True

        for j, col in enumerate(df.columns):
            table.rows[0].cells[j].text = str(col)

        for i, (_, row_data) in enumerate(df.iterrows(), 1):
            for j, val in enumerate(row_data):
                cell = table.rows[i].cells[j]
                if isinstance(val, float):
                    cell.text = self.fmt % val if not np.isnan(val) else ""
                else:
                    cell.text = str(val)

        style_word_table_typography(table, header_rows=(0,))
        apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)
        add_word_notes_paragraph(doc, "* p<0.10, ** p<0.05, *** p<0.01")
        doc.save(filename)

    def save(self, filename: str) -> None:
        """Auto-detect format from extension and save."""
        path = Path(filename)
        ext = path.suffix.lower()
        if ext in (".xlsx", ".xls"):
            self.to_excel(filename)
        elif ext == ".docx":
            self.to_word(filename)
        elif ext == ".tex":
            path.write_text(self.to_latex(), encoding="utf-8")
        elif ext in (".html", ".htm"):
            path.write_text(self.to_html(), encoding="utf-8")
        elif ext == ".md":
            path.write_text(self.to_markdown(), encoding="utf-8")
        elif ext == ".csv":
            self.to_dataframe().to_csv(filename)
        else:
            path.write_text(self.to_text(), encoding="utf-8")

    # ═══════════════════════════════════════════════════════════════════════
    # Dunder
    # ═══════════════════════════════════════════════════════════════════════

    def _render(self, fmt: str) -> str:
        return {
            "text": self.to_text,
            "latex": self.to_latex,
            "tex": self.to_latex,
            "html": self.to_html,
            "markdown": self.to_markdown,
            "md": self.to_markdown,
        }.get(fmt, self.to_text)()

    def __str__(self) -> str:
        return self._render(self._output)

    def __repr__(self) -> str:
        return self.__str__()

    def _repr_html_(self) -> str:
        return self.to_html()


# ═══════════════════════════════════════════════════════════════════════════
# mean_comparison() — public API
# ═══════════════════════════════════════════════════════════════════════════

def mean_comparison(
    data: pd.DataFrame,
    variables: List[str],
    group: str,
    *,
    weights: Optional[str] = None,
    test: str = "ttest",
    fmt: str = "%.2f",
    title: str = "Balance Table",
    group_labels: Optional[Tuple[str, str]] = None,
    output: str = "text",
    filename: Optional[str] = None,
) -> MeanComparisonResult:
    """
    Compare means across two groups with statistical tests.

    Creates a publication-quality balance table showing means, standard
    deviations, differences, and p-values for each variable.

    Parameters
    ----------
    data : pd.DataFrame
        Input data.
    variables : list of str
        Column names to compare.
    group : str
        Binary grouping variable name.
    weights : str, optional
        Column name for weights (reserved for future use).
    test : str, default ``"ttest"``
        Test type: ``"ttest"`` (Welch's), ``"ranksum"`` (Mann-Whitney),
        or ``"chi2"`` (chi-squared).
    fmt : str, default ``"%.2f"``
        Format string for numeric values.
    title : str, default ``"Balance Table"``
        Table title.
    group_labels : tuple of str, optional
        Labels for (control, treated) groups. Defaults to
        ``("Control", "Treated")``.
    output : str, default ``"text"``
        Output format: ``"text"``, ``"latex"``, ``"html"``, ``"markdown"``.
    filename : str, optional
        Save to file (format auto-detected from extension).

    Returns
    -------
    MeanComparisonResult
        Object with ``.to_text()``, ``.to_latex()``, ``.to_html()``,
        ``.to_markdown()``, ``.to_excel(filename)``, ``.to_word(filename)``,
        ``.to_dataframe()``, ``.save(filename)`` methods.
        Renders as rich HTML in Jupyter notebooks via ``_repr_html_()``.

    Examples
    --------
    >>> import statspai as sp
    >>> sp.mean_comparison(df, ["age", "income", "education"], group="treated")
    """
    if group_labels is None:
        group_labels = ("Control", "Treated")

    _VALID_OUTPUTS = {
        "text", "latex", "tex", "html", "markdown", "md", "word", "excel",
    }
    if output not in _VALID_OUTPUTS:
        raise ValueError(
            f"output={output!r} is invalid. Must be one of: "
            f"{sorted(_VALID_OUTPUTS)}"
        )

    result = MeanComparisonResult(
        data=data,
        variables=variables,
        group=group,
        group_labels=group_labels,
        test=test,
        fmt=fmt,
        title=title,
        weights=weights,
        output=output,
    )

    if filename:
        result.save(filename)

    # No auto-print: Jupyter uses _repr_html_, REPL uses __repr__.
    # Scripts: explicit `print(mean_comparison(...))`.
    return result
