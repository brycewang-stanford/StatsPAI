"""
Publication-quality multi-model comparison tables.

Equivalent to R's modelsummary() / fixest::etable() / Stata's esttab.
Accepts any mix of EconometricResults and CausalResult objects and
produces side-by-side tables in text, LaTeX, HTML, or DataFrame format.

References
----------
Arel-Bundock, V. (2022). "modelsummary: Data and Model Summaries in R."
*Journal of Statistical Software*, 103(1), 1-23. [@arelbundock2022modelsummary]
"""

from typing import Optional, List, Dict, Any, Union, Sequence
import numpy as np
import pandas as pd


# ======================================================================
# Public API
# ======================================================================

def modelsummary(
    *models,
    model_names: Optional[List[str]] = None,
    stars: Union[bool, Dict[str, float]] = True,
    se_type: str = 'parentheses',
    show_ci: bool = False,
    stats: Optional[List[str]] = None,
    output: str = 'text',
    title: str = '',
    notes: Optional[List[str]] = None,
    fmt: str = '%.4f',
    coef_map: Optional[Dict[str, str]] = None,
    add_rows: Optional[Dict[str, List[str]]] = None,
) -> Union[str, pd.DataFrame]:
    """
    Generate publication-quality multi-model comparison table.

    Accepts any combination of EconometricResults and CausalResult objects
    and produces a formatted comparison table.

    Parameters
    ----------
    *models
        Model result objects (EconometricResults, CausalResult, or anything
        with ``.params`` and ``.std_errors`` Series attributes).
    model_names : list of str, optional
        Column headers. Defaults to ``(1), (2), ...``.
    stars : bool or dict, default True
        Significance stars. True for default thresholds.
        Dict maps symbol to threshold: ``{'*': 0.1, '**': 0.05, '***': 0.01}``.
    se_type : str, default 'parentheses'
        ``'parentheses'``, ``'brackets'``, or ``'none'``.
    show_ci : bool, default False
        Show 95% confidence intervals instead of standard errors.
    stats : list of str, optional
        Bottom-row statistics to include. Available keys:
        ``'nobs'``, ``'r_squared'``, ``'adj_r_squared'``, ``'f_stat'``,
        ``'aic'``, ``'bic'``, ``'method'``, ``'bandwidth'``, ``'estimand'``.
        Defaults to ``['nobs', 'r_squared']``.
    output : str, default 'text'
        ``'text'``, ``'latex'``, ``'html'``, or ``'dataframe'``.
    title : str
        Table title / caption.
    notes : list of str, optional
        Footnotes. A significance note is auto-appended when ``stars=True``.
    fmt : str, default '%.4f'
        Number format for coefficients and SEs.
    coef_map : dict, optional
        Rename / reorder coefficients: ``{'x1': 'Education'}``.
        Variables not in the map are kept in original order; set value
        to ``None`` to drop a variable.
    add_rows : dict, optional
        Extra rows: ``{'FE: Year': ['Yes', 'No', 'Yes']}``.

    Returns
    -------
    str or pd.DataFrame
        Formatted table (text/LaTeX/HTML string, or DataFrame).

    Examples
    --------
    >>> import statspai as sp
    >>> r1 = sp.regress("y ~ x1 + x2", data=df)
    >>> r2 = sp.regress("y ~ x1 + x2 + x3", data=df, robust='hc1')
    >>> r3 = sp.did(df, y='y', treat='d', time='t')
    >>> print(sp.modelsummary(r1, r2, r3))
    """
    if len(models) == 0:
        raise ValueError("At least one model required.")

    # Defaults
    if model_names is None:
        model_names = [f'({i + 1})' for i in range(len(models))]
    if len(model_names) != len(models):
        raise ValueError("Length of model_names must match number of models.")
    if stats is None:
        stats = ['nobs', 'r_squared']
    if notes is None:
        notes = []

    star_thresholds = (
        {'***': 0.01, '**': 0.05, '*': 0.1}
        if stars is True
        else (stars if isinstance(stars, dict) else {})
    )

    # 1. Extract coefficient data from each model
    coef_data = [_extract_coefs(m) for m in models]

    # 2. Determine variable ordering
    all_vars = _collect_variables(coef_data, coef_map)

    # 3. Build coefficient rows
    rows = _build_coef_rows(
        all_vars, coef_data, model_names,
        star_thresholds, se_type, show_ci, fmt, coef_map,
    )

    # 4. Build statistics rows
    stat_rows = _build_stat_rows(models, model_names, stats, fmt)

    # 5. Build custom rows
    custom_rows = []
    if add_rows:
        for label, vals in add_rows.items():
            row = [label] + list(vals)
            while len(row) < len(model_names) + 1:
                row.append('')
            custom_rows.append(row)

    # 6. Format output
    all_rows = rows + [None] + stat_rows + custom_rows  # None = separator
    columns = [''] + model_names

    if output == 'dataframe':
        return _to_dataframe(all_rows, columns)
    elif output == 'latex':
        return _to_latex(all_rows, columns, title, notes, star_thresholds)
    elif output == 'html':
        return _to_html(all_rows, columns, title, notes, star_thresholds)
    elif output.endswith('.xlsx'):
        _to_excel(all_rows, columns, title, notes, star_thresholds, output)
        return f"Table exported to: {output}"
    elif output.endswith('.docx'):
        _to_word(all_rows, columns, title, notes, star_thresholds, output)
        return f"Table exported to: {output}"
    else:  # text
        return _to_text(all_rows, columns, title, notes, star_thresholds)


def coefplot(
    *models,
    model_names: Optional[List[str]] = None,
    variables: Optional[List[str]] = None,
    ax=None,
    figsize: tuple = (8, 6),
    colors: Optional[List[str]] = None,
    title: Optional[str] = None,
    alpha: float = 0.05,
):
    """
    Forest plot comparing coefficients across models.

    Parameters
    ----------
    *models
        Model result objects.
    model_names : list of str, optional
    variables : list of str, optional
        Which variables to plot. Default: all shared variables.
    ax : matplotlib Axes, optional
    figsize : tuple
    colors : list of str, optional
    title : str, optional
    alpha : float
        Significance level for CIs.

    Returns
    -------
    (fig, ax)
    """
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        raise ImportError("matplotlib required. Install: pip install matplotlib")

    if model_names is None:
        model_names = [f'Model {i + 1}' for i in range(len(models))]
    if colors is None:
        colors = ['#2C3E50', '#E74C3C', '#3498DB', '#2ECC71',
                  '#9B59B6', '#F39C12', '#1ABC9C', '#E67E22']

    from scipy import stats as sp_stats
    z_crit = sp_stats.norm.ppf(1 - alpha / 2)

    coef_data = [_extract_coefs(m) for m in models]

    # Variables to plot
    if variables is None:
        all_v = set()
        for cd in coef_data:
            all_v.update(cd.keys())
        variables = sorted(all_v)

    if ax is None:
        fig, ax = plt.subplots(figsize=figsize)
    else:
        fig = ax.get_figure()

    n_vars = len(variables)
    n_models = len(models)
    offsets = np.linspace(-0.15 * (n_models - 1), 0.15 * (n_models - 1), n_models)

    for m_idx, (cd, name) in enumerate(zip(coef_data, model_names)):
        color = colors[m_idx % len(colors)]
        positions = []
        estimates = []
        ci_lo = []
        ci_hi = []

        for v_idx, var in enumerate(variables):
            if var in cd:
                coef, se, _ = cd[var]
                positions.append(v_idx + offsets[m_idx])
                estimates.append(coef)
                ci_lo.append(coef - z_crit * se)
                ci_hi.append(coef + z_crit * se)

        if positions:
            pos = np.array(positions)
            est = np.array(estimates)
            lo = np.array(ci_lo)
            hi = np.array(ci_hi)
            ax.scatter(est, pos, color=color, s=40, zorder=5, label=name)
            ax.errorbar(
                est, pos, xerr=[est - lo, hi - est],
                fmt='none', color=color, capsize=3, linewidth=1, zorder=3,
            )

    ax.axvline(x=0, color='gray', linestyle='--', linewidth=0.8)
    ax.set_yticks(range(n_vars))
    ax.set_yticklabels(variables)
    ax.invert_yaxis()
    ax.set_xlabel('Coefficient Estimate')
    ax.set_title(title or 'Coefficient Plot')
    ax.legend(fontsize=9, frameon=False)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()

    return fig, ax


# ======================================================================
# Extraction helpers
# ======================================================================

def _extract_coefs(model) -> Dict[str, tuple]:
    """Extract {var_name: (coef, se, pvalue)} from a model object."""
    result = {}

    params = getattr(model, 'params', None)
    std_errors = getattr(model, 'std_errors', None)
    pvalues = getattr(model, 'pvalues', None)

    if params is None:
        return result

    if isinstance(params, pd.Series):
        for var in params.index:
            coef = float(params[var])
            se = float(std_errors[var]) if std_errors is not None and var in std_errors.index else np.nan
            if pvalues is not None:
                if isinstance(pvalues, pd.Series) and var in pvalues.index:
                    pv = float(pvalues[var])
                elif isinstance(pvalues, np.ndarray):
                    idx = list(params.index).index(var)
                    pv = float(pvalues[idx]) if idx < len(pvalues) else np.nan
                else:
                    pv = np.nan
            else:
                pv = np.nan
            result[var] = (coef, se, pv)
    elif isinstance(params, (int, float)):
        name = getattr(model, 'estimand', 'estimate')
        se = float(std_errors) if std_errors is not None else np.nan
        pv = float(pvalues) if pvalues is not None else np.nan
        result[name] = (float(params), se, pv)

    return result


def _extract_stat(model, stat_name: str) -> Any:
    """Extract a named statistic from a model."""
    # Try diagnostics (EconometricResults)
    diag = getattr(model, 'diagnostics', {})
    if isinstance(diag, dict):
        mapping = {
            'r_squared': ['R-squared', 'r_squared'],
            'adj_r_squared': ['Adj. R-squared', 'adj_r_squared'],
            'f_stat': ['F-statistic', 'f_stat'],
            'aic': ['AIC', 'aic'],
            'bic': ['BIC', 'bic'],
        }
        for key in mapping.get(stat_name, [stat_name]):
            if key in diag:
                return diag[key]

    # Try model_info (CausalResult)
    mi = getattr(model, 'model_info', {})
    if isinstance(mi, dict):
        direct_map = {
            'r_squared': 'r_squared',
            'bandwidth': 'bandwidth_h',
            'method': 'method',
            'estimand': 'estimand',
        }
        key = direct_map.get(stat_name, stat_name)
        if key in mi:
            return mi[key]

    # Try data_info
    di = getattr(model, 'data_info', {})
    if isinstance(di, dict) and stat_name == 'nobs' and 'nobs' in di:
        return di['nobs']

    # Try direct attribute
    if stat_name == 'nobs':
        return getattr(model, 'n_obs', None)

    return None


# ======================================================================
# Table construction
# ======================================================================

def _collect_variables(
    coef_data: List[Dict[str, tuple]],
    coef_map: Optional[Dict[str, str]],
) -> List[str]:
    """Collect and order variable names across models."""
    seen = {}  # var → first appearance order
    for cd in coef_data:
        for var in cd:
            if var not in seen:
                seen[var] = len(seen)

    if coef_map:
        # Reorder: mapped variables first (in map order), then remainder
        ordered = []
        for old_name in coef_map:
            if old_name in seen and coef_map[old_name] is not None:
                ordered.append(old_name)
        for var in sorted(seen, key=seen.get):
            if var not in ordered:
                if coef_map is None or var not in coef_map or coef_map[var] is not None:
                    ordered.append(var)
        return ordered
    else:
        return sorted(seen, key=seen.get)


def _format_num(value: float, fmt: str) -> str:
    if np.isnan(value):
        return ''
    return fmt % value


def _stars_str(pvalue: float, thresholds: Dict[str, float]) -> str:
    if not thresholds or np.isnan(pvalue):
        return ''
    for symbol in sorted(thresholds, key=lambda s: thresholds[s]):
        if pvalue < thresholds[symbol]:
            best = symbol
    else:
        best = ''
    # Find the most stars that apply
    best = ''
    for symbol, thresh in sorted(thresholds.items(), key=lambda x: x[1]):
        if pvalue < thresh:
            if len(symbol) > len(best):
                best = symbol
    return best


def _build_coef_rows(
    variables, coef_data, model_names,
    star_thresholds, se_type, show_ci, fmt, coef_map,
):
    """Build coefficient + SE rows."""
    rows = []
    for var in variables:
        display_name = var
        if coef_map and var in coef_map:
            display_name = coef_map[var]
            if display_name is None:
                continue

        # Coefficient row
        coef_row = [display_name]
        for cd in coef_data:
            if var in cd:
                coef, se, pv = cd[var]
                s = _stars_str(pv, star_thresholds)
                coef_row.append(_format_num(coef, fmt) + s)
            else:
                coef_row.append('')
        rows.append(coef_row)

        # SE / CI row
        if se_type != 'none':
            se_row = ['']
            for cd in coef_data:
                if var in cd:
                    _, se, _ = cd[var]
                    if show_ci:
                        from scipy import stats as sp_stats
                        z = sp_stats.norm.ppf(0.975)
                        coef = cd[var][0]
                        lo, hi = coef - z * se, coef + z * se
                        cell = f'[{_format_num(lo, fmt)}, {_format_num(hi, fmt)}]'
                    else:
                        wrap = '({})' if se_type == 'parentheses' else '[{}]'
                        cell = wrap.format(_format_num(se, fmt))
                    se_row.append(cell)
                else:
                    se_row.append('')
            rows.append(se_row)

    return rows


def _build_stat_rows(models, model_names, stats, fmt):
    """Build bottom statistics rows."""
    rows = []

    stat_labels = {
        'nobs': 'Observations',
        'r_squared': 'R²',
        'adj_r_squared': 'Adj. R²',
        'f_stat': 'F-statistic',
        'aic': 'AIC',
        'bic': 'BIC',
        'method': 'Method',
        'bandwidth': 'Bandwidth',
        'estimand': 'Estimand',
    }

    for stat in stats:
        label = stat_labels.get(stat, stat)
        row = [label]
        for model in models:
            val = _extract_stat(model, stat)
            if val is None:
                row.append('')
            elif isinstance(val, float):
                if stat == 'nobs':
                    row.append(f'{int(val):,}')
                else:
                    row.append(_format_num(val, fmt))
            elif isinstance(val, int):
                row.append(f'{val:,}')
            else:
                row.append(str(val))
        rows.append(row)

    return rows


# ======================================================================
# Output formatters
# ======================================================================

def _to_dataframe(all_rows, columns):
    """Convert to pandas DataFrame."""
    clean = [r for r in all_rows if r is not None]
    df = pd.DataFrame(clean, columns=columns)
    return df


def _to_text(all_rows, columns, title, notes, star_thresholds):
    """Format as plain text table."""
    # Compute column widths
    n_cols = len(columns)
    widths = [0] * n_cols
    for row in all_rows:
        if row is None:
            continue
        for j, cell in enumerate(row):
            widths[j] = max(widths[j], len(str(cell)))
    for j, col in enumerate(columns):
        widths[j] = max(widths[j], len(col))

    # Add padding
    widths = [w + 2 for w in widths]
    total_width = sum(widths) + n_cols - 1

    lines = []
    if title:
        lines.append(title)
        lines.append('')

    lines.append('=' * total_width)

    # Header
    header = ''
    for j, col in enumerate(columns):
        if j == 0:
            header += col.ljust(widths[j])
        else:
            header += col.rjust(widths[j])
    lines.append(header)
    lines.append('-' * total_width)

    # Rows
    for row in all_rows:
        if row is None:
            lines.append('-' * total_width)
            continue
        line = ''
        for j, cell in enumerate(row):
            cell_str = str(cell)
            if j == 0:
                line += cell_str.ljust(widths[j])
            else:
                line += cell_str.rjust(widths[j])
        lines.append(line)

    lines.append('=' * total_width)

    # Notes
    if star_thresholds:
        star_note = '; '.join(
            f'{sym} p<{thresh}' for sym, thresh in
            sorted(star_thresholds.items(), key=lambda x: -len(x[0]))
        )
        lines.append(star_note)
    for note in notes:
        lines.append(note)

    return '\n'.join(lines)


def _to_latex(all_rows, columns, title, notes, star_thresholds):
    """Format as LaTeX table."""
    n_cols = len(columns)
    spec = 'l' + 'c' * (n_cols - 1)

    lines = [
        '\\begin{table}[htbp]',
        '\\centering',
    ]
    if title:
        lines.append(f'\\caption{{{title}}}')

    lines += [
        f'\\begin{{tabular}}{{{spec}}}',
        '\\hline\\hline',
    ]

    # Header
    lines.append(' & '.join(columns) + ' \\\\')
    lines.append('\\hline')

    # Rows
    for row in all_rows:
        if row is None:
            lines.append('\\hline')
            continue
        cells = []
        for cell in row:
            s = str(cell)
            # Escape LaTeX special chars in variable names
            s = s.replace('_', '\\_').replace('%', '\\%').replace('&', '\\&')
            cells.append(s)
        lines.append(' & '.join(cells) + ' \\\\')

    lines += [
        '\\hline\\hline',
        '\\end{tabular}',
    ]

    # Notes
    note_lines = []
    if star_thresholds:
        star_note = '; '.join(
            f'{sym} p<{thresh}' for sym, thresh in
            sorted(star_thresholds.items(), key=lambda x: -len(x[0]))
        )
        note_lines.append(star_note)
    note_lines.extend(notes or [])
    if note_lines:
        lines.append('\\begin{tablenotes}')
        lines.append('\\footnotesize')
        for n in note_lines:
            lines.append(f'\\item {n}')
        lines.append('\\end{tablenotes}')

    lines.append('\\end{table}')
    return '\n'.join(lines)


def _to_html(all_rows, columns, title, notes, star_thresholds):
    """Format as HTML table."""
    lines = ['<table style="border-collapse:collapse; font-family:serif;">']
    if title:
        lines.append(f'<caption style="font-weight:bold; font-size:14px;">{title}</caption>')

    # Header
    lines.append('<thead><tr>')
    for j, col in enumerate(columns):
        align = 'left' if j == 0 else 'center'
        lines.append(f'  <th style="text-align:{align}; border-bottom:2px solid black; '
                     f'padding:4px 8px;">{col}</th>')
    lines.append('</tr></thead>')

    # Body
    lines.append('<tbody>')
    for row in all_rows:
        if row is None:
            lines.append('<tr><td colspan="{}" style="border-bottom:1px solid black;"></td></tr>'
                         .format(len(columns)))
            continue
        lines.append('<tr>')
        for j, cell in enumerate(row):
            align = 'left' if j == 0 else 'center'
            lines.append(f'  <td style="text-align:{align}; padding:2px 8px;">{cell}</td>')
        lines.append('</tr>')
    lines.append('</tbody>')

    # Footer
    lines.append('<tfoot><tr><td colspan="{}" style="border-top:2px solid black; '
                 'font-size:11px; padding-top:4px;">'.format(len(columns)))
    if star_thresholds:
        star_note = '; '.join(
            f'{sym} p&lt;{thresh}' for sym, thresh in
            sorted(star_thresholds.items(), key=lambda x: -len(x[0]))
        )
        lines.append(star_note + '<br/>')
    for n in (notes or []):
        lines.append(n + '<br/>')
    lines.append('</td></tr></tfoot>')

    lines.append('</table>')
    return '\n'.join(lines)


def _to_excel(all_rows, columns, title, notes, star_thresholds, filename):
    """Export to publication-quality Excel."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Results'

    row_idx = 1

    # Title
    if title:
        cell = ws.cell(row=row_idx, column=1, value=title)
        cell.font = Font(bold=True, size=12)
        row_idx += 2

    # Header
    header_font = Font(bold=True, size=10)
    thin_bottom = Border(bottom=Side(style='thin'))
    for j, col in enumerate(columns, 1):
        cell = ws.cell(row=row_idx, column=j, value=col)
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_bottom
    row_idx += 1

    # Data
    for row in all_rows:
        if row is None:
            # Separator
            for j in range(1, len(columns) + 1):
                ws.cell(row=row_idx, column=j).border = Border(
                    bottom=Side(style='thin'))
            row_idx += 1
            continue
        for j, val in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=j, value=str(val))
            cell.font = Font(size=9)
            cell.alignment = Alignment(
                horizontal='left' if j == 1 else 'center')
        row_idx += 1

    # Bottom border
    for j in range(1, len(columns) + 1):
        ws.cell(row=row_idx - 1, column=j).border = Border(
            bottom=Side(style='double'))

    # Notes
    if star_thresholds:
        star_note = '; '.join(
            f'{sym} p<{thresh}' for sym, thresh in
            sorted(star_thresholds.items(), key=lambda x: -len(x[0])))
        row_idx += 1
        cell = ws.cell(row=row_idx, column=1, value=star_note)
        cell.font = Font(italic=True, size=8)
    for note in (notes or []):
        row_idx += 1
        cell = ws.cell(row=row_idx, column=1, value=note)
        cell.font = Font(italic=True, size=8)

    # Auto-width
    for col_idx in range(1, len(columns) + 1):
        letter = openpyxl.utils.get_column_letter(col_idx)
        ws.column_dimensions[letter].width = 15

    wb.save(filename)


def _to_word(all_rows, columns, title, notes, star_thresholds, filename):
    """Export to publication-quality Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx required for Word export. "
            "Install: pip install python-docx")

    doc = Document()

    # Title
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Count actual rows (skip None separators, use border instead)
    data_rows = [r for r in all_rows if r is not None]
    sep_indices = set()
    count = 0
    for i, r in enumerate(all_rows):
        if r is None:
            sep_indices.add(count)
        else:
            count += 1

    n_rows = len(data_rows) + 1  # +1 for header
    n_cols = len(columns)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, col in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val) if val else ''
            for para in cell.paragraphs:
                para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                                  else WD_ALIGN_PARAGRAPH.CENTER)
                for run in para.runs:
                    run.font.size = Pt(9)

    # APA borders
    def _set_border(cell, top=None, bottom=None):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge, val in [('top', top), ('bottom', bottom)]:
            if val:
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), val.get('val', 'single'))
                el.set(qn('w:sz'), val.get('sz', '4'))
                el.set(qn('w:color'), '000000')
                el.set(qn('w:space'), '0')
                borders.append(el)
        for edge in ['start', 'end']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')
            borders.append(el)
        tcPr.append(borders)

    thick = {'val': 'single', 'sz': '12'}
    thin = {'val': 'single', 'sz': '4'}

    for j in range(n_cols):
        # Top of table
        _set_border(table.rows[0].cells[j], top=thick, bottom=thin)
        # Bottom of table
        _set_border(table.rows[-1].cells[j], bottom=thick)
        # Separator rows
        for si in sep_indices:
            if si + 1 < n_rows:
                _set_border(table.rows[si + 1].cells[j], top=thin)

    # Notes
    if star_thresholds or notes:
        p = doc.add_paragraph()
        if star_thresholds:
            star_note = '; '.join(
                f'{sym} p<{thresh}' for sym, thresh in
                sorted(star_thresholds.items(), key=lambda x: -len(x[0])))
            run = p.add_run(star_note + '\n')
            run.italic = True
            run.font.size = Pt(8)
        for note in (notes or []):
            run = p.add_run(note + '\n')
            run.italic = True
            run.font.size = Pt(8)

    doc.save(filename)
