"""
Output regression results to Excel in publication-ready format
Similar to Stata's outreg2 command
"""

import pandas as pd
import numpy as np
from typing import List, Dict, Any, Optional, Union
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from pathlib import Path
import warnings

from ..core.results import EconometricResults


class OutReg2:
    """
    Export regression results to Excel in publication-ready format
    
    Mimics Stata's outreg2 functionality with additional customization options.
    """
    
    def __init__(self):
        self.results = []
        self.model_names = []
        self.title = "Regression Results"
        self.notes = []
        
    def add_model(self, results: EconometricResults, name: Optional[str] = None):
        """
        Add a regression model to the output table
        
        Parameters
        ----------
        results : EconometricResults
            Fitted model results
        name : str, optional
            Name for the model column
        """
        self.results.append(results)
        if name is None:
            name = f"Model {len(self.results)}"
        self.model_names.append(name)
        
    def clear(self):
        """Clear all stored results"""
        self.results = []
        self.model_names = []
        self.notes = []
    
    def set_title(self, title: str):
        """Set the table title"""
        self.title = title
        
    def add_note(self, note: str):
        """Add a note to be displayed below the table"""
        self.notes.append(note)
    
    def _format_number(self, value: float, decimal_places: int = 3) -> str:
        """Format numbers for display"""
        if pd.isna(value) or np.isnan(value):
            return ""
        return f"{value:.{decimal_places}f}"
    
    def _format_pvalue(self, pvalue: float) -> str:
        """Format p-values with stars"""
        if pd.isna(pvalue) or np.isnan(pvalue):
            return ""

        stars = ""
        if pvalue < 0.01:
            stars = "***"
        elif pvalue < 0.05:
            stars = "**"
        elif pvalue < 0.1:
            stars = "*"

        return stars

    def _is_causal_forest(self, result) -> bool:
        """Check if result is a Causal Forest model"""
        return hasattr(result, 'effect') and hasattr(result, '_forest')
    
    def _create_regression_table(
        self, 
        show_stars: bool = True,
        show_se: bool = True,
        show_tstat: bool = False,
        decimal_places: int = 3,
        variable_labels: Optional[Dict[str, str]] = None
    ) -> pd.DataFrame:
        """
        Create the main regression table
        
        Parameters
        ----------
        show_stars : bool, default True
            Whether to show significance stars
        show_se : bool, default True
            Whether to show standard errors in parentheses
        show_tstat : bool, default False
            Whether to show t-statistics instead of standard errors
        decimal_places : int, default 3
            Number of decimal places for formatting
        variable_labels : dict, optional
            Custom labels for variables
            
        Returns
        -------
        pd.DataFrame
            Formatted regression table
        """
        if not self.results:
            raise ValueError("No regression results added. Use add_model() first.")
        
        # Check if we have mixed model types
        has_regression = any(not self._is_causal_forest(r) for r in self.results)
        has_causal_forest = any(self._is_causal_forest(r) for r in self.results)
        
        if has_regression and has_causal_forest:
            warnings.warn(
                "Mixing regression and Causal Forest results in the same table. "
                "Table format optimized for regression models."
            )
        
        # If all models are Causal Forest, create specialized table
        if has_causal_forest and not has_regression:
            return self._create_causal_forest_table(decimal_places, variable_labels)
        
        # Standard regression table (original code)
        # Get all unique variables across all models
        all_vars = set()
        for result in self.results:
            if not self._is_causal_forest(result):
                all_vars.update(result.params.index)
        all_vars = sorted(list(all_vars))
        
        # Create the table structure
        table_data = []
        
        for var in all_vars:
            # Variable label
            var_label = variable_labels.get(var, var) if variable_labels else var
            
            # Coefficient row
            coef_row = [var_label]
            for result in self.results:
                if self._is_causal_forest(result):
                    coef_row.append("N/A (Causal Forest)")
                    continue
                    
                if var in result.params.index:
                    coef = result.params[var]
                    
                    # Handle pvalues - convert to Series if it's an array
                    if isinstance(result.pvalues, np.ndarray):
                        pvalues_series = pd.Series(result.pvalues, index=result.params.index)
                        pval = pvalues_series[var] if var in pvalues_series.index else np.nan
                    else:
                        pval = result.pvalues[var] if var in result.pvalues.index else np.nan
                    
                    coef_str = self._format_number(coef, decimal_places)
                    if show_stars:
                        coef_str += self._format_pvalue(pval)
                    
                    coef_row.append(coef_str)
                else:
                    coef_row.append("")
            
            table_data.append(coef_row)
            
            # Standard error / t-statistic row
            if show_se or show_tstat:
                se_row = [""]
                for result in self.results:
                    if self._is_causal_forest(result):
                        se_row.append("")
                        continue
                        
                    if var in result.params.index:
                        if show_tstat and var in result.tvalues.index:
                            value = result.tvalues[var]
                            se_str = f"({self._format_number(value, decimal_places)})"
                        elif show_se and var in result.std_errors.index:
                            value = result.std_errors[var]
                            se_str = f"({self._format_number(value, decimal_places)})"
                        else:
                            se_str = ""
                        se_row.append(se_str)
                    else:
                        se_row.append("")
                
                table_data.append(se_row)
        
        # Add model diagnostics
        table_data.append([""])  # Empty row
        
        # R-squared
        r2_row = ["R-squared"]
        for result in self.results:
            if self._is_causal_forest(result):
                r2_row.append("N/A")
            else:
                r2 = result.diagnostics.get('R-squared', np.nan)
                r2_row.append(self._format_number(r2, decimal_places))
        table_data.append(r2_row)
        
        # Adjusted R-squared
        adj_r2_row = ["Adj. R-squared"]
        for result in self.results:
            if self._is_causal_forest(result):
                adj_r2_row.append("N/A")
            else:
                adj_r2 = result.diagnostics.get('Adj. R-squared', np.nan)
                adj_r2_row.append(self._format_number(adj_r2, decimal_places))
        table_data.append(adj_r2_row)
        
        # Number of observations
        n_row = ["Observations"]
        for result in self.results:
            if self._is_causal_forest(result):
                n_obs = result.data_info.get('nobs', 'N/A')
            else:
                n_obs = result.data_info.get('nobs', 'N/A')
            n_row.append(str(n_obs))
        table_data.append(n_row)
        
        # F-statistic / Number of trees
        f_row = ["F-statistic / Trees"]
        for result in self.results:
            if self._is_causal_forest(result):
                n_trees = result.diagnostics.get('n_estimators', 'N/A')
                f_row.append(str(n_trees))
            else:
                f_stat = result.diagnostics.get('F-statistic', np.nan)
                f_row.append(self._format_number(f_stat, decimal_places))
        table_data.append(f_row)
        
        # Create DataFrame
        columns = ["Variables"] + self.model_names
        df = pd.DataFrame(table_data, columns=columns)
        
        return df
    
    def _create_causal_forest_table(
        self, 
        decimal_places: int = 3,
        variable_labels: Optional[Dict[str, str]] = None
    ) -> pd.DataFrame:
        """
        Create specialized table for Causal Forest results
        
        Parameters
        ----------
        decimal_places : int, default 3
            Number of decimal places for formatting
        variable_labels : dict, optional
            Custom labels for variables
            
        Returns
        -------
        pd.DataFrame
            Formatted Causal Forest table
        """
        table_data = []
        
        # Average Treatment Effect row
        ate_row = ["Average Treatment Effect"]
        for result in self.results:
            ate = result.diagnostics.get('average_treatment_effect', np.nan)
            ate_row.append(self._format_number(ate, decimal_places))
        table_data.append(ate_row)
        
        # Empty row for separation
        table_data.append([""])
        
        # Model specifications
        method_row = ["Method"]
        for result in self.results:
            method_row.append("Causal Forest")
        table_data.append(method_row)
        
        trees_row = ["Number of Trees"]
        for result in self.results:
            n_trees = result.diagnostics.get('n_estimators', 'N/A')
            trees_row.append(str(n_trees))
        table_data.append(trees_row)
        
        features_row = ["Number of Features"]
        for result in self.results:
            n_features = result.data_info.get('n_features', 'N/A')
            features_row.append(str(n_features))
        table_data.append(features_row)
        
        treatment_row = ["Treatment Type"]
        for result in self.results:
            t_type = result.diagnostics.get('treatment_type', 'Unknown')
            treatment_row.append(t_type)
        table_data.append(treatment_row)
        
        # Number of observations
        n_row = ["Observations"]
        for result in self.results:
            n_obs = result.data_info.get('nobs', 'N/A')
            n_row.append(str(n_obs))
        table_data.append(n_row)
        
        # Create DataFrame
        columns = ["Model Statistics"] + self.model_names
        df = pd.DataFrame(table_data, columns=columns)
        
        return df
    
    def to_excel(
        self,
        filename: str,
        sheet_name: str = "Regression Results",
        show_stars: bool = True,
        show_se: bool = True,
        show_tstat: bool = False,
        decimal_places: int = 3,
        variable_labels: Optional[Dict[str, str]] = None,
        add_formatting: bool = True
    ):
        """
        Export regression results to Excel file
        
        Parameters
        ----------
        filename : str
            Output Excel filename
        sheet_name : str, default "Regression Results"
            Name of the Excel sheet
        show_stars : bool, default True
            Whether to show significance stars (*, **, ***)
        show_se : bool, default True
            Whether to show standard errors in parentheses
        show_tstat : bool, default False
            Whether to show t-statistics instead of standard errors
        decimal_places : int, default 3
            Number of decimal places for numbers
        variable_labels : dict, optional
            Custom labels for variables {var_name: display_name}
        add_formatting : bool, default True
            Whether to add Excel formatting for publication quality
            
        Examples
        --------
        >>> outreg = OutReg2()
        >>> outreg.add_model(results1, "Model 1")
        >>> outreg.add_model(results2, "Model 2")
        >>> outreg.set_title("Main Results")
        >>> outreg.add_note("Standard errors in parentheses")
        >>> outreg.add_note("* p<0.1, ** p<0.05, *** p<0.01")
        >>> outreg.to_excel("results.xlsx")
        """
        # Create the regression table
        df = self._create_regression_table(
            show_stars=show_stars,
            show_se=show_se,
            show_tstat=show_tstat,
            decimal_places=decimal_places,
            variable_labels=variable_labels
        )
        
        # Ensure filename has .xlsx extension
        if not filename.endswith('.xlsx'):
            filename += '.xlsx'
        
        if add_formatting:
            self._export_with_formatting(df, filename, sheet_name)
        else:
            # Simple export without formatting
            df.to_excel(filename, sheet_name=sheet_name, index=False)
        
        print(f"Regression results exported to: {filename}")
    
    def _export_with_formatting(self, df: pd.DataFrame, filename: str, sheet_name: str):
        """Export with publication-quality Excel formatting (book-tab style).

        Replaces the legacy four-edge ``thin_border`` (Excel grid style)
        with the three-line book-tab convention used everywhere else in
        :mod:`statspai.output`. Font is Times New Roman to match
        ``regtable`` / ``sumstats`` / ``paper_tables`` / ``collection``.
        """
        from ._excel_style import (
            TIMES, BODY_PT, HEADER_PT, NOTES_PT,
            apply_booktab_borders, autofit_columns, write_title,
        )

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        title_font = Font(bold=True, name=TIMES, size=12)
        header_font = Font(bold=True, name=TIMES, size=HEADER_PT)
        normal_font = Font(name=TIMES, size=BODY_PT)
        italic_font = Font(italic=True, name=TIMES, size=NOTES_PT)
        center_align = Alignment(horizontal='center', vertical='center')
        left_align = Alignment(horizontal='left', vertical='center')

        n_cols = len(df.columns)
        current_row = 1

        if self.title:
            title_cell = ws.cell(row=current_row, column=1, value=self.title)
            title_cell.font = title_font
            title_cell.alignment = center_align
            if n_cols > 1:
                ws.merge_cells(
                    start_row=current_row, start_column=1,
                    end_row=current_row, end_column=n_cols,
                )
            current_row += 1

        # Header row
        header_top_row = current_row
        for col_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(row=current_row, column=col_idx, value=col_name)
            cell.font = header_font
            cell.alignment = center_align
        header_bot_row = current_row
        current_row += 1

        # Body rows — no per-cell borders, book-tab adds the three rules
        body_top_row = current_row
        last_data_row = current_row - 1
        for row_idx, row_data in df.iterrows():
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=current_row, column=col_idx, value=value)
                cell.font = normal_font
                cell.alignment = left_align if col_idx == 1 else center_align
            last_data_row = current_row
            current_row += 1
        body_bot_row = last_data_row

        apply_booktab_borders(
            ws,
            header_top_row=header_top_row,
            header_bot_row=header_bot_row,
            body_top_row=body_top_row,
            body_bot_row=body_bot_row,
            n_cols=n_cols,
        )

        # Notes
        if self.notes:
            current_row += 1
            for note in self.notes:
                note_cell = ws.cell(row=current_row, column=1, value=note)
                note_cell.font = italic_font
                note_cell.alignment = left_align
                current_row += 1

        autofit_columns(ws, n_cols, max_width=20)
        wb.save(filename)
    
    def to_latex(
        self,
        filename: Optional[str] = None,
        show_stars: bool = True,
        show_se: bool = True,
        decimal_places: int = 3,
        variable_labels: Optional[Dict[str, str]] = None
    ) -> str:
        """
        Export regression results to LaTeX table format
        
        Parameters
        ----------
        filename : str, optional
            Output LaTeX filename. If None, returns LaTeX string
        show_stars : bool, default True
            Whether to show significance stars
        show_se : bool, default True
            Whether to show standard errors in parentheses
        decimal_places : int, default 3
            Number of decimal places for numbers
        variable_labels : dict, optional
            Custom labels for variables
            
        Returns
        -------
        str
            LaTeX table code
        """
        df = self._create_regression_table(
            show_stars=show_stars,
            show_se=show_se,
            decimal_places=decimal_places,
            variable_labels=variable_labels
        )
        
        # Create LaTeX table
        n_cols = len(df.columns)
        col_spec = "l" + "c" * (n_cols - 1)
        
        latex_lines = [
            "\\begin{table}[htbp]",
            "\\centering",
            f"\\caption{{{self.title}}}",
            f"\\begin{{tabular}}{{{col_spec}}}",
            "\\hline\\hline"
        ]
        
        # Add header
        header = " & ".join(df.columns) + " \\\\"
        latex_lines.append(header)
        latex_lines.append("\\hline")
        
        # Add data rows
        for _, row in df.iterrows():
            row_str = " & ".join([str(x) if x != "" else "" for x in row]) + " \\\\"
            latex_lines.append(row_str)
        
        latex_lines.extend([
            "\\hline\\hline",
            "\\end{tabular}"
        ])
        
        # Add notes
        if self.notes:
            latex_lines.append("\\begin{tablenotes}")
            latex_lines.append("\\footnotesize")
            for note in self.notes:
                latex_lines.append(f"\\item {note}")
            latex_lines.append("\\end{tablenotes}")
        
        latex_lines.append("\\end{table}")
        
        latex_code = "\n".join(latex_lines)
        
        if filename:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(latex_code)
            print(f"LaTeX table exported to: {filename}")
        
        return latex_code


    def to_word(
        self,
        filename: str,
        show_stars: bool = True,
        show_se: bool = True,
        show_tstat: bool = False,
        decimal_places: int = 3,
        variable_labels: Optional[Dict[str, str]] = None,
    ):
        """
        Export regression results to Word (.docx) file.

        Produces a publication-ready table with APA-style formatting,
        equivalent to Stata's ``outreg2 using results.doc, word``.

        Parameters
        ----------
        filename : str
            Output Word filename (.docx).
        show_stars : bool, default True
            Whether to show significance stars.
        show_se : bool, default True
            Whether to show standard errors in parentheses.
        show_tstat : bool, default False
            Whether to show t-statistics instead of standard errors.
        decimal_places : int, default 3
            Number of decimal places.
        variable_labels : dict, optional
            Custom labels for variables.

        Examples
        --------
        >>> outreg = OutReg2()
        >>> outreg.add_model(results1, "OLS")
        >>> outreg.add_model(results2, "IV")
        >>> outreg.to_word("table1.docx")
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError(
                "python-docx required for Word export. "
                "Install: pip install python-docx"
            )

        df = self._create_regression_table(
            show_stars=show_stars,
            show_se=show_se,
            show_tstat=show_tstat,
            decimal_places=decimal_places,
            variable_labels=variable_labels,
        )

        if not filename.endswith('.docx'):
            filename += '.docx'

        doc = Document()

        # Title
        if self.title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(self.title)
            title_run.bold = True
            title_run.font.size = Pt(12)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Table
        n_rows = len(df)
        n_cols = len(df.columns)
        table = doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col_name)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for i, (_, row) in enumerate(df.iterrows()):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(val) if val != "" else ""
                for paragraph in cell.paragraphs:
                    if j == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # APA-style borders: top/bottom thick, header bottom thin
        self._apply_apa_borders(table, n_rows + 1, n_cols)

        # Notes
        if self.notes:
            notes_para = doc.add_paragraph()
            notes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for note in self.notes:
                run = notes_para.add_run(note + '\n')
                run.italic = True
                run.font.size = Pt(8)

        doc.save(filename)
        print(f"Regression results exported to: {filename}")

    @staticmethod
    def _apply_apa_borders(table, n_rows, n_cols):
        """Apply APA-style borders (top, header-bottom, bottom only)."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge, val in [('top', top), ('bottom', bottom),
                              ('start', start), ('end', end)]:
                if val is not None:
                    element = OxmlElement(f'w:{edge}')
                    element.set(qn('w:val'), val.get('val', 'single'))
                    element.set(qn('w:sz'), val.get('sz', '4'))
                    element.set(qn('w:color'), val.get('color', '000000'))
                    element.set(qn('w:space'), '0')
                    tcBorders.append(element)
            tcPr.append(tcBorders)

        thick = {'val': 'single', 'sz': '12', 'color': '000000'}
        thin = {'val': 'single', 'sz': '4', 'color': '000000'}
        none_b = {'val': 'none', 'sz': '0', 'color': '000000'}

        for i in range(n_rows):
            for j in range(n_cols):
                cell = table.rows[i].cells[j]
                top = thick if i == 0 else (thin if i == 1 else none_b)
                bottom = thick if i == n_rows - 1 else (thin if i == 0 else none_b)
                set_cell_border(cell, top=top, bottom=bottom,
                                start=none_b, end=none_b)


def outreg2(
    *results: EconometricResults,
    filename: str,
    model_names: Optional[List[str]] = None,
    title: str = "Regression Results",
    notes: Optional[List[str]] = None,
    show_stars: bool = True,
    show_se: bool = True,
    show_tstat: bool = False,
    decimal_places: int = 3,
    variable_labels: Optional[Dict[str, str]] = None,
    format: str = "auto"
) -> Optional[str]:
    """
    Convenient function to export multiple regression results.

    Parameters
    ----------
    *results : EconometricResults
        Multiple regression results to include.
    filename : str
        Output filename. The format is auto-detected from the
        extension: ``.xlsx`` → Excel, ``.docx`` → Word, ``.tex`` →
        LaTeX.  Override with ``format=``.
    model_names : list of str, optional
        Names for each model
    title : str, default "Regression Results"
        Table title
    notes : list of str, optional
        Notes to include below the table
    show_stars : bool, default True
        Whether to show significance stars
    show_se : bool, default True
        Whether to show standard errors in parentheses
    show_tstat : bool, default False
        Whether to show t-statistics instead of standard errors
    decimal_places : int, default 3
        Number of decimal places
    variable_labels : dict, optional
        Custom variable labels
    format : str, default "auto"
        Output format. ``"auto"`` (default) detects from the filename
        extension: ``.xlsx`` → Excel, ``.docx`` → Word, ``.tex`` →
        LaTeX. Unknown extensions default to Excel. Override with
        ``"excel"``, ``"word"``, or ``"latex"`` explicitly.

    Returns
    -------
    str or None
        LaTeX code if format="latex", otherwise None
        
    Examples
    --------
    >>> # Export single model
    >>> outreg2(results1, filename="model1.xlsx")
    
    >>> # Export multiple models with custom names
    >>> outreg2(results1, results2, results3,
    ...         filename="all_models.xlsx",
    ...         model_names=["Baseline", "With Controls", "Full Model"],
    ...         title="Main Regression Results",
    ...         notes=["Standard errors in parentheses",
    ...                "* p<0.1, ** p<0.05, *** p<0.01"])
    """
    reg_table = OutReg2()
    reg_table.set_title(title)
    
    # Add models
    for i, result in enumerate(results):
        name = model_names[i] if model_names and i < len(model_names) else f"Model {i+1}"
        reg_table.add_model(result, name)
    
    # Add notes
    if notes:
        for note in notes:
            reg_table.add_note(note)
    
    # Add default significance note if showing stars and no custom notes
    if show_stars and not notes:
        reg_table.add_note("* p<0.1, ** p<0.05, *** p<0.01")
    
    # Auto-detect format from filename extension
    fmt = format.lower()
    if fmt == 'auto':
        if filename.endswith('.docx') or filename.endswith('.doc'):
            fmt = 'word'
        elif filename.endswith('.tex'):
            fmt = 'latex'
        else:
            fmt = 'excel'

    # Export
    if fmt == "latex":
        return reg_table.to_latex(
            filename=filename if filename.endswith('.tex') else filename + '.tex',
            show_stars=show_stars,
            show_se=show_se,
            decimal_places=decimal_places,
            variable_labels=variable_labels
        )
    elif fmt == "word":
        reg_table.to_word(
            filename=filename,
            show_stars=show_stars,
            show_se=show_se,
            show_tstat=show_tstat,
            decimal_places=decimal_places,
            variable_labels=variable_labels,
        )
        return None
    else:
        reg_table.to_excel(
            filename=filename,
            show_stars=show_stars,
            show_se=show_se,
            show_tstat=show_tstat,
            decimal_places=decimal_places,
            variable_labels=variable_labels
        )
        return None
