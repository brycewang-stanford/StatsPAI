"""``MeanComparisonResult`` and the ``mean_comparison()`` public API.

Extracted from ``regression_table.py`` (which had ballooned to 3,335
lines and held two unrelated result classes). This module is the single
home for two-group balance / mean-comparison tables; the regression
table renderer in ``regression_table.py`` no longer needs to keep this
~510-line block on its hot path.

The class is re-exported from ``regression_table`` for backwards
compatibility with code that did
``from statspai.output.regression_table import MeanComparisonResult``.
"""

from __future__ import annotations

import warnings
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
from scipy import stats as sp_stats

from ._format import format_stars as _format_stars
from .estimates import _latex_escape, _html_escape

__all__ = ["MeanComparisonResult", "mean_comparison"]


class MeanComparisonResult:
    """Rich result object for balance / mean comparison tables."""

    def __init__(
        self,
        data: pd.DataFrame,
        variables: List[str],
        group: str,
        group_labels: Tuple[str, str],
        test: str,
        fmt: str,
        title: str,
        weights: Optional[str],
        output: str = "text",
    ):
        self.variables = variables
        self.group = group
        self.group_labels = group_labels
        self.test = test
        self.fmt = fmt
        self.title = title
        self.weights = weights
        self._output = output

        # Compute all stats
        self._compute(data)

    def _compute(self, data: pd.DataFrame) -> None:
        g0_label, g1_label = self.group_labels
        mask0 = data[self.group] == 0
        mask1 = data[self.group] == 1

        # If not binary 0/1, use the two unique values
        unique_vals = sorted(data[self.group].dropna().unique())
        if len(unique_vals) == 2:
            mask0 = data[self.group] == unique_vals[0]
            mask1 = data[self.group] == unique_vals[1]

        d0 = data.loc[mask0]
        d1 = data.loc[mask1]
        self.n0 = len(d0)
        self.n1 = len(d1)

        rows: List[Dict[str, Any]] = []
        for var in self.variables:
            v0 = d0[var].dropna()
            v1 = d1[var].dropna()

            mean0, sd0 = v0.mean(), v0.std()
            mean1, sd1 = v1.mean(), v1.std()
            diff = mean1 - mean0

            # Test
            if self.test == "ranksum":
                try:
                    stat, pval = sp_stats.mannwhitneyu(v0, v1, alternative="two-sided")
                except Exception:
                    pval = np.nan
            elif self.test == "chi2":
                try:
                    ct = pd.crosstab(data[self.group], data[var])
                    chi2, pval, _, _ = sp_stats.chi2_contingency(ct)
                except Exception:
                    pval = np.nan
            else:
                # Default: Welch t-test
                try:
                    stat, pval = sp_stats.ttest_ind(v0, v1, equal_var=False)
                except Exception:
                    pval = np.nan

            stars = _format_stars(pval, (0.10, 0.05, 0.01))

            rows.append({
                "variable": var,
                "mean0": mean0,
                "sd0": sd0,
                "mean1": mean1,
                "sd1": sd1,
                "diff": diff,
                "pvalue": pval,
                "stars": stars,
            })

        self._rows = rows

    # ═══════════════════════════════════════════════════════════════════════
    # TEXT
    # ═══════════════════════════════════════════════════════════════════════

    def to_text(self) -> str:
        g0_label, g1_label = self.group_labels
        var_w = max(max(len(v) for v in self.variables), 10) + 2
        col_w = 16
        total_w = var_w + col_w * 4 + 2

        thick = "━" * total_w
        thin = "─" * total_w
        lines: List[str] = []

        lines.append(f"  {self.title}")
        lines.append(thick)

        # Header
        hdr = f"{'':>{var_w}}"
        hdr += f"{g0_label:>{col_w}}"
        hdr += f"{g1_label:>{col_w}}"
        hdr += f"{'Diff':>{col_w}}"
        hdr += f"{'p-value':>{col_w}}"
        lines.append(hdr)

        sub = f"{'':>{var_w}}"
        sub += f"{'Mean (SD)':>{col_w}}"
        sub += f"{'Mean (SD)':>{col_w}}"
        sub += " " * col_w
        sub += " " * col_w
        lines.append(sub)

        lines.append(thin)

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            fmt_diff = (self.fmt % row["diff"]) + row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            col0 = f"{fmt_mean0} ({fmt_sd0})"
            col1 = f"{fmt_mean1} ({fmt_sd1})"

            line = f"{row['variable']:<{var_w}}"
            line += f"{col0:>{col_w}}"
            line += f"{col1:>{col_w}}"
            line += f"{fmt_diff:>{col_w}}"
            line += f"{fmt_pval:>{col_w}}"
            lines.append(line)

        lines.append(thin)
        lines.append(f"{'N':<{var_w}}{self.n0:>{col_w},}{self.n1:>{col_w},}")
        lines.append(thick)
        lines.append(f"* p<0.10, ** p<0.05, *** p<0.01")

        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # HTML
    # ═══════════════════════════════════════════════════════════════════════

    def to_html(self) -> str:
        g0_label, g1_label = self.group_labels
        lines: List[str] = []
        lines.append(
            '<table class="balance-table" style="border-collapse:collapse; '
            'font-family:\'Times New Roman\', serif; font-size:13px;">'
        )
        lines.append(
            f'<caption style="font-weight:bold; font-size:14px; '
            f'margin-bottom:8px;">{_html_escape(self.title)}</caption>'
        )

        # Header
        lines.append("<thead>")
        lines.append("<tr>")
        for h in ["", g0_label, g1_label, "Diff", "p-value"]:
            border = "border-top:3px solid black; border-bottom:1px solid black; "
            align = "text-align:center;" if h else "text-align:left;"
            lines.append(f'<th style="{border}{align} padding:4px 10px;">{_html_escape(h)}</th>')
        lines.append("</tr>")
        lines.append("<tr>")
        lines.append('<th style="text-align:left;"></th>')
        lines.append(f'<th style="text-align:center; font-weight:normal; font-size:11px;">Mean (SD)</th>')
        lines.append(f'<th style="text-align:center; font-weight:normal; font-size:11px;">Mean (SD)</th>')
        lines.append('<th></th><th></th>')
        lines.append("</tr>")
        lines.append("</thead>")
        lines.append("<tbody>")

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            fmt_diff = (self.fmt % row["diff"]) + row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            lines.append("<tr>")
            lines.append(f'<td style="text-align:left; padding:2px 10px;">{_html_escape(row["variable"])}</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_mean0} ({fmt_sd0})</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_mean1} ({fmt_sd1})</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_diff}</td>')
            lines.append(f'<td style="text-align:center; padding:2px 10px;">{fmt_pval}</td>')
            lines.append("</tr>")

        # N row
        lines.append(
            f'<tr><td colspan="5" style="border-top:1px solid black; padding:0;"></td></tr>'
        )
        lines.append("<tr>")
        lines.append(f'<td style="text-align:left; padding:2px 10px;">N</td>')
        lines.append(f'<td style="text-align:center; padding:2px 10px;">{self.n0:,}</td>')
        lines.append(f'<td style="text-align:center; padding:2px 10px;">{self.n1:,}</td>')
        lines.append('<td></td><td></td>')
        lines.append("</tr>")
        lines.append(
            f'<tr><td colspan="5" style="border-top:3px solid black; padding:0;"></td></tr>'
        )

        lines.append("</tbody>")
        lines.append("<tfoot>")
        lines.append(
            f'<tr><td colspan="5" style="text-align:left; font-size:11px; padding:4px 10px;">'
            f'* p&lt;0.10, ** p&lt;0.05, *** p&lt;0.01</td></tr>'
        )
        lines.append("</tfoot>")
        lines.append("</table>")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # LaTeX
    # ═══════════════════════════════════════════════════════════════════════

    def to_latex(self) -> str:
        g0_label, g1_label = self.group_labels
        lines: List[str] = []
        lines.append("\\begin{table}[htbp]")
        lines.append("\\centering")
        lines.append(f"\\caption{{{_latex_escape(self.title)}}}")
        lines.append("\\begin{tabular}{lcccc}")
        lines.append("\\hline\\hline")
        lines.append(
            f" & {_latex_escape(g0_label)} & {_latex_escape(g1_label)} "
            f"& Diff & p-value \\\\"
        )
        lines.append(
            " & Mean (SD) & Mean (SD) & & \\\\"
        )
        lines.append("\\hline")

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            diff_str = self.fmt % row["diff"]
            stars_str = row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            var_esc = _latex_escape(row["variable"])
            lines.append(
                f"{var_esc} & {fmt_mean0} ({fmt_sd0}) & "
                f"{fmt_mean1} ({fmt_sd1}) & "
                f"{diff_str}{stars_str} & {fmt_pval} \\\\"
            )

        lines.append("\\hline")
        lines.append(f"N & {self.n0:,} & {self.n1:,} & & \\\\")
        lines.append("\\hline\\hline")
        lines.append(
            "\\multicolumn{5}{l}{\\footnotesize * p<0.10, ** p<0.05, *** p<0.01} \\\\"
        )
        lines.append("\\end{tabular}")
        lines.append("\\end{table}")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Markdown
    # ═══════════════════════════════════════════════════════════════════════

    def to_markdown(self) -> str:
        g0_label, g1_label = self.group_labels
        lines: List[str] = []
        lines.append(f"**{self.title}**")
        lines.append("")
        lines.append(f"| | {g0_label} | {g1_label} | Diff | p-value |")
        lines.append("|---|---:|---:|---:|---:|")

        for row in self._rows:
            fmt_mean0 = self.fmt % row["mean0"]
            fmt_sd0 = self.fmt % row["sd0"]
            fmt_mean1 = self.fmt % row["mean1"]
            fmt_sd1 = self.fmt % row["sd1"]
            fmt_diff = (self.fmt % row["diff"]) + row["stars"]
            fmt_pval = "%.3f" % row["pvalue"] if not np.isnan(row["pvalue"]) else ""

            lines.append(
                f"| {row['variable']} | {fmt_mean0} ({fmt_sd0}) | "
                f"{fmt_mean1} ({fmt_sd1}) | {fmt_diff} | {fmt_pval} |"
            )

        lines.append(f"| N | {self.n0:,} | {self.n1:,} | | |")
        lines.append("")
        lines.append("*\\* p<0.10, \\*\\* p<0.05, \\*\\*\\* p<0.01*")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # DataFrame
    # ═══════════════════════════════════════════════════════════════════════

    def to_dataframe(self) -> pd.DataFrame:
        g0_label, g1_label = self.group_labels
        records = []
        for row in self._rows:
            records.append({
                "Variable": row["variable"],
                f"{g0_label} Mean": row["mean0"],
                f"{g0_label} SD": row["sd0"],
                f"{g1_label} Mean": row["mean1"],
                f"{g1_label} SD": row["sd1"],
                "Difference": row["diff"],
                "p-value": row["pvalue"],
                "Significance": row["stars"],
            })
        df = pd.DataFrame(records)
        df = df.set_index("Variable")
        return df

    # ═══════════════════════════════════════════════════════════════════════
    # Excel / Word / Save
    # ═══════════════════════════════════════════════════════════════════════

    def to_excel(self, filename: str) -> None:
        """Export balance table to Excel as a book-tab three-line table."""
        try:
            import openpyxl  # noqa: F401
        except ImportError:
            warnings.warn("openpyxl required for Excel export: pip install openpyxl")
            return
        from ._excel_style import render_dataframe_to_xlsx
        render_dataframe_to_xlsx(
            self.to_dataframe(),
            filename,
            title=getattr(self, "title", None),
            sheet_name="Balance Table",
            index_label="Variable",
        )

    def to_word(self, filename: str) -> None:
        """Export balance table to Word in AER/QJE book-tab style."""
        try:
            from docx import Document
        except ImportError:
            warnings.warn("python-docx required for Word export: pip install python-docx")
            return

        from ._aer_style import (
            apply_word_booktab_rules,
            style_word_table_typography,
            add_word_notes_paragraph,
        )

        doc = Document()
        doc.add_heading(self.title, level=2)
        df = self.to_dataframe().reset_index()
        n_rows = len(df) + 1
        n_cols = len(df.columns)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.autofit = True

        for j, col in enumerate(df.columns):
            table.rows[0].cells[j].text = str(col)

        for i, (_, row_data) in enumerate(df.iterrows(), 1):
            for j, val in enumerate(row_data):
                cell = table.rows[i].cells[j]
                if isinstance(val, float):
                    cell.text = self.fmt % val if not np.isnan(val) else ""
                else:
                    cell.text = str(val)

        style_word_table_typography(table, header_rows=(0,))
        apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)
        add_word_notes_paragraph(doc, "* p<0.10, ** p<0.05, *** p<0.01")
        doc.save(filename)

    def save(self, filename: str) -> None:
        """Auto-detect format from extension and save."""
        path = Path(filename)
        ext = path.suffix.lower()
        if ext in (".xlsx", ".xls"):
            self.to_excel(filename)
        elif ext == ".docx":
            self.to_word(filename)
        elif ext == ".tex":
            path.write_text(self.to_latex(), encoding="utf-8")
        elif ext in (".html", ".htm"):
            path.write_text(self.to_html(), encoding="utf-8")
        elif ext == ".md":
            path.write_text(self.to_markdown(), encoding="utf-8")
        elif ext == ".csv":
            self.to_dataframe().to_csv(filename)
        else:
            path.write_text(self.to_text(), encoding="utf-8")

    # ═══════════════════════════════════════════════════════════════════════
    # Dunder
    # ═══════════════════════════════════════════════════════════════════════

    def _render(self, fmt: str) -> str:
        return {
            "text": self.to_text,
            "latex": self.to_latex,
            "tex": self.to_latex,
            "html": self.to_html,
            "markdown": self.to_markdown,
            "md": self.to_markdown,
        }.get(fmt, self.to_text)()

    def __str__(self) -> str:
        return self._render(self._output)

    def __repr__(self) -> str:
        return self.__str__()

    def _repr_html_(self) -> str:
        return self.to_html()


# ═══════════════════════════════════════════════════════════════════════════
# mean_comparison() — public API
# ═══════════════════════════════════════════════════════════════════════════

def mean_comparison(
    data: pd.DataFrame,
    variables: List[str],
    group: str,
    *,
    weights: Optional[str] = None,
    test: str = "ttest",
    fmt: str = "%.2f",
    title: str = "Balance Table",
    group_labels: Optional[Tuple[str, str]] = None,
    output: str = "text",
    filename: Optional[str] = None,
) -> MeanComparisonResult:
    """
    Compare means across two groups with statistical tests.

    Creates a publication-quality balance table showing means, standard
    deviations, differences, and p-values for each variable.

    Parameters
    ----------
    data : pd.DataFrame
        Input data.
    variables : list of str
        Column names to compare.
    group : str
        Binary grouping variable name.
    weights : str, optional
        Column name for weights (reserved for future use).
    test : str, default ``"ttest"``
        Test type: ``"ttest"`` (Welch's), ``"ranksum"`` (Mann-Whitney),
        or ``"chi2"`` (chi-squared).
    fmt : str, default ``"%.2f"``
        Format string for numeric values.
    title : str, default ``"Balance Table"``
        Table title.
    group_labels : tuple of str, optional
        Labels for (control, treated) groups. Defaults to
        ``("Control", "Treated")``.
    output : str, default ``"text"``
        Output format: ``"text"``, ``"latex"``, ``"html"``, ``"markdown"``.
    filename : str, optional
        Save to file (format auto-detected from extension).

    Returns
    -------
    MeanComparisonResult
        Object with ``.to_text()``, ``.to_latex()``, ``.to_html()``,
        ``.to_markdown()``, ``.to_excel(filename)``, ``.to_word(filename)``,
        ``.to_dataframe()``, ``.save(filename)`` methods.
        Renders as rich HTML in Jupyter notebooks via ``_repr_html_()``.

    Examples
    --------
    >>> import statspai as sp
    >>> sp.mean_comparison(df, ["age", "income", "education"], group="treated")
    """
    if group_labels is None:
        group_labels = ("Control", "Treated")

    _VALID_OUTPUTS = {
        "text", "latex", "tex", "html", "markdown", "md", "word", "excel",
    }
    if output not in _VALID_OUTPUTS:
        raise ValueError(
            f"output={output!r} is invalid. Must be one of: "
            f"{sorted(_VALID_OUTPUTS)}"
        )

    result = MeanComparisonResult(
        data=data,
        variables=variables,
        group=group,
        group_labels=group_labels,
        test=test,
        fmt=fmt,
        title=title,
        weights=weights,
        output=output,
    )

    if filename:
        result.save(filename)

    # No auto-print: Jupyter uses _repr_html_, REPL uses __repr__.
    # Scripts: explicit `print(mean_comparison(...))`.
    return result
