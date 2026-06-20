"""
Cross-tabulation with statistical tests.

Equivalent to Stata's ``tab var1 var2, chi2 exact``.
Exports to text, LaTeX, Excel, Word.
"""

from typing import Optional, Any, Union
import pandas as pd
from scipy import stats


def tab(
    data: pd.DataFrame,
    row: str,
    col: Optional[str] = None,
    output: str = "text",
    test: bool = True,
    margins: bool = True,
    normalize: Optional[str] = None,
    title: Optional[str] = None,
) -> Union[str, pd.DataFrame]:
    """
    Cross-tabulation with chi-squared / Fisher's exact test.

    Parameters
    ----------
    data : pd.DataFrame
        Input data.
    row : str
        Row variable.
    col : str, optional
        Column variable. If None, produces a one-way frequency table.
    output : str, default 'text'
        'text', 'dataframe', 'latex', or filepath (.xlsx/.docx).
    test : bool, default True
        Include chi-squared test (and Fisher's exact for 2x2).
    margins : bool, default True
        Show row/column totals.
    normalize : str, optional
        'row', 'col', 'all', or None. Normalize to proportions.
    title : str, optional
        Table title.

    Returns
    -------
    str or pd.DataFrame

    Examples
    --------
    >>> import numpy as np
    >>> import pandas as pd
    >>> import statspai as sp
    >>> rng = np.random.default_rng(0)
    >>> df = pd.DataFrame({
    ...     'treatment': rng.integers(0, 2, size=200),
    ...     'outcome': rng.integers(0, 2, size=200),
    ... })
    >>> txt = sp.tab(df, 'treatment', 'outcome')
    >>> isinstance(txt, str)
    True
    >>> ct = sp.tab(df, 'treatment', 'outcome', output='dataframe')
    >>> isinstance(ct, pd.DataFrame)
    True
    >>> _ = sp.tab(df, 'treatment')  # one-way frequency
    >>> sp.tab(df, 'treatment', 'outcome', output='crosstab.docx')  # doctest: +SKIP
    """
    if col is None:
        return _one_way_tab(data, row, output, title)

    # Two-way cross-tabulation
    ct = pd.crosstab(
        data[row],
        data[col],
        margins=margins,
        margins_name="Total",
    )

    if normalize:
        norm_map = {"row": "index", "col": "columns", "all": "all"}
        ct_norm = pd.crosstab(
            data[row],
            data[col],
            normalize=norm_map.get(normalize, normalize),
            margins=margins,
            margins_name="Total",
        )
        # Format as percentages
        ct_display = ct_norm.map(lambda x: f"{x:.1%}")
    else:
        ct_display = ct

    # Statistical test
    test_result = None
    if test:
        # Use raw counts (no margins)
        ct_raw = pd.crosstab(data[row], data[col])
        chi2, p_chi2, dof, expected = stats.chi2_contingency(ct_raw)
        test_result = {
            "chi2": chi2,
            "pvalue": p_chi2,
            "df": dof,
        }
        # Fisher's exact for 2x2
        if ct_raw.shape == (2, 2):
            _, p_fisher = stats.fisher_exact(ct_raw)
            test_result["fisher_pvalue"] = p_fisher

    if title is None:
        title = f"Tabulation: {row} × {col}"

    return _format_tab(ct_display, test_result, output, title)


def _one_way_tab(data: Any, var: str, output: str, title: Optional[str]) -> Any:
    """One-way frequency table."""
    counts = data[var].value_counts().sort_index()
    pcts = data[var].value_counts(normalize=True).sort_index()

    df = pd.DataFrame(
        {
            "Freq": counts,
            "Percent": pcts.map(lambda x: f"{x:.1%}"),
            "Cum.": pcts.cumsum().map(lambda x: f"{x:.1%}"),
        }
    )
    df.loc["Total"] = [int(counts.sum()), "100.0%", "100.0%"]

    if title is None:
        title = f"Tabulation: {var}"

    return _format_tab(df, None, output, title)


def _format_tab(df: Any, test_result: Any, output: str, title: Optional[str]) -> Any:
    """Route to output format."""
    if output == "dataframe":
        return df

    if output.endswith(".xlsx"):
        _tab_to_excel(df, test_result, output, title)
        return f"Exported to: {output}"
    elif output.endswith(".docx"):
        _tab_to_word(df, test_result, output, title)
        return f"Exported to: {output}"
    elif output == "latex":
        return _tab_to_latex(df, test_result, title)

    # Default: text
    lines = []
    if title:
        lines.append(title)
    lines.append("=" * 60)
    lines.append(df.to_string())
    lines.append("=" * 60)

    if test_result:
        lines.append(
            f"  Pearson chi2({test_result['df']}) = "
            f"{test_result['chi2']:.4f}   Pr = {test_result['pvalue']:.4f}"
        )
        if "fisher_pvalue" in test_result:
            lines.append(
                f"  Fisher's exact                     "
                f"Pr = {test_result['fisher_pvalue']:.4f}"
            )

    return "\n".join(lines)


def _tab_to_latex(df: Any, test_result: Any, title: Optional[str]) -> str:
    latex = df.to_latex(caption=title)
    if test_result:
        latex += f"\n% chi2({test_result['df']}) = {test_result['chi2']:.4f}, p = {test_result['pvalue']:.4f}"
    return str(latex)


def _tab_to_excel(
    df: Any, test_result: Any, filename: str, title: Optional[str]
) -> Any:
    """Cross-tabulation -> book-tab xlsx.

    Layout follows the shared three-line convention. The chi-square /
    p-value test row is appended as an italic note below the table.
    """
    import openpyxl

    from ._excel_style import (
        apply_booktab_borders,
        apply_publication_sheet_defaults,
        autofit_columns,
        write_body,
        write_header,
        write_notes,
        write_title,
    )

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Crosstab"

    n_cols = len(df.columns) + 1
    row = 1
    if title:
        row = write_title(ws, row, n_cols, title)

    header_top, header_bot = write_header(ws, row, df, index_label="")
    row = header_bot + 1
    body_top, body_bot = write_body(ws, row, df)

    apply_booktab_borders(
        ws,
        header_top_row=header_top,
        header_bot_row=header_bot,
        body_top_row=body_top,
        body_bot_row=body_bot,
        n_cols=n_cols,
    )

    notes = []
    if test_result:
        notes.append(
            f"chi2({test_result['df']}) = {test_result['chi2']:.4f}, "
            f"p = {test_result['pvalue']:.4f}"
        )
    next_row = body_bot + 1
    if notes:
        next_row = write_notes(ws, next_row, notes, n_cols=n_cols)

    autofit_columns(ws, n_cols)
    apply_publication_sheet_defaults(
        ws,
        header_top_row=header_top,
        header_bot_row=header_bot,
        final_row=max(next_row - 1, body_bot),
        n_cols=n_cols,
    )
    wb.save(filename)


def _tab_to_word(df: Any, test_result: Any, filename: str, title: Optional[str]) -> Any:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx required. Install: pip install python-docx")

    from ._aer_style import (
        apply_word_document_defaults,
        apply_word_booktab_rules,
        style_word_table_typography,
        add_word_notes_paragraph,
    )

    doc = Document()
    apply_word_document_defaults(doc)

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    n_rows = len(df) + 1
    n_cols = len(df.columns) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header
    table.rows[0].cells[0].text = ""
    for j, col in enumerate(df.columns, 1):
        table.rows[0].cells[j].text = str(col)
    # Data
    for i, (idx, row) in enumerate(df.iterrows()):
        table.rows[i + 1].cells[0].text = str(idx)
        for j, val in enumerate(row, 1):
            table.rows[i + 1].cells[j].text = str(val)

    style_word_table_typography(
        table,
        header_rows=(0,),
        header_pt=10,
        body_pt=9,
        align_first_col="left",
        align_data_cols="center",
    )
    apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)

    if test_result:
        add_word_notes_paragraph(
            doc,
            f"Pearson chi2({test_result['df']}) = {test_result['chi2']:.4f}, "
            f"p = {test_result['pvalue']:.4f}",
        )

    doc.save(filename)
