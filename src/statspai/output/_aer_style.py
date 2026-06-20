"""
AER/QJE-style booktab styling helpers for DOCX and XLSX exports.

Implements the three-rule book-tabs convention favoured by economics journals
(*American Economic Review*, *Quarterly Journal of Economics*, *Review of
Economic Studies*, *JPE*): a heavy top rule above column headers, a thin
mid rule separating headers from the body, and a heavy bottom rule above
table notes — with **no** internal vertical or horizontal borders.

These helpers are deliberately tolerant of missing dependencies — they
expect callers to have already imported ``python-docx`` / ``openpyxl``
successfully.

References
----------
The booktab style traces to the LaTeX ``booktabs`` package
(https://ctan.org/pkg/booktabs) and is the de-facto rule for AER /
QJE / RES regression tables.
"""

from __future__ import annotations

from typing import Any, Iterable, Optional, Sequence

import pandas as pd

# OOXML border eighths-of-a-point sizes (sz attribute):
#  - 4 = 0.5pt (thin)
#  - 8 = 1pt (medium)
#  - 12 = 1.5pt (heavy / book-tab top/bottom)
TOP_RULE_SZ = "12"  # heavy
MID_RULE_SZ = "4"  # thin
BOTTOM_RULE_SZ = "12"  # heavy

DEFAULT_FONT = "Times New Roman"
HEADER_PT = 10
BODY_PT = 10
NOTE_PT = 8


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------


def _set_style_font_name(style: Any, font_name: str) -> None:
    """Set a Word style font across OOXML font slots."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def apply_word_document_defaults(
    doc: Any,
    *,
    font_name: str = DEFAULT_FONT,
) -> None:
    """Apply economics-journal document defaults to a DOCX document.

    Table helpers handle cell-level typography and book-tab rules; this
    function sets the surrounding Word document to the same publication
    baseline: Times New Roman, black headings, compact paragraph spacing,
    and standard one-inch manuscript margins.
    """
    from docx.shared import Inches, Pt, RGBColor

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style_specs = (
        ("Normal", BODY_PT, False, 0),
        ("Heading 1", 14, True, 6),
        ("Heading 2", 12, True, 6),
        ("Heading 3", 11, True, 4),
        ("Title", 14, True, 6),
    )
    for style_name, pt_size, bold, space_after in style_specs:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        _set_style_font_name(style, font_name)
        style.font.size = Pt(pt_size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.line_spacing = 1.0


def _docx_border_xml(
    *,
    top: Optional[dict[str, str]] = None,
    bottom: Optional[dict[str, str]] = None,
    left: Optional[dict[str, str]] = None,
    right: Optional[dict[str, str]] = None,
    insideH: Optional[dict[str, str]] = None,
    insideV: Optional[dict[str, str]] = None,
) -> Any:
    """Build a ``<w:tcBorders>`` element overriding the four cell edges.

    Each kwarg, when supplied, is a dict like ``{"sz": "12", "val": "single",
    "color": "000000"}``; when omitted that edge is set to ``val="nil"``
    (no border).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_borders = OxmlElement("w:tcBorders")
    edges = {
        "top": top,
        "left": left,
        "bottom": bottom,
        "right": right,
        "insideH": insideH,
        "insideV": insideV,
    }
    for edge, attrs in edges.items():
        elem = OxmlElement(f"w:{edge}")
        if attrs is None:
            elem.set(qn("w:val"), "nil")
        else:
            elem.set(qn("w:val"), attrs.get("val", "single"))
            elem.set(qn("w:sz"), attrs.get("sz", "4"))
            elem.set(qn("w:space"), attrs.get("space", "0"))
            elem.set(qn("w:color"), attrs.get("color", "000000"))
        tc_borders.append(elem)
    return tc_borders


def _set_cell_borders(cell: Any, **edges: Optional[dict[str, str]]) -> None:
    """Replace a cell's ``tcBorders`` element. Edges not passed → ``nil``."""
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(_docx_border_xml(**edges))


def _clear_table_borders(table: Any) -> None:
    """Strip every cell of every border (book-tab default state)."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)


def apply_word_booktab_rules(
    table: Any,
    *,
    header_top_idx: int = 0,
    header_bot_idx: int = 0,
    body_last_idx: Optional[int] = None,
) -> None:
    """Apply the three book-tab rules to a DOCX table.

    Parameters
    ----------
    table : docx.table.Table
    header_top_idx : int
        Row index that should receive the heavy *top rule* on its top edge.
        Typically ``0``.
    header_bot_idx : int
        Row index whose **bottom edge** receives the thin *mid rule*. The
        body starts on ``header_bot_idx + 1``. For a multi-row header,
        pass the last header row's index.
    body_last_idx : int, optional
        Row whose bottom edge receives the heavy *bottom rule*. Defaults
        to the table's final row.
    """
    rows = table.rows
    n = len(rows)
    if n == 0:
        return
    if body_last_idx is None:
        body_last_idx = n - 1

    _clear_table_borders(table)

    top = {"val": "single", "sz": TOP_RULE_SZ}
    mid = {"val": "single", "sz": MID_RULE_SZ}
    bot = {"val": "single", "sz": BOTTOM_RULE_SZ}

    for cell in rows[header_top_idx].cells:
        _set_cell_borders(cell, top=top)
    for cell in rows[header_bot_idx].cells:
        # Preserve top rule when top and bottom header rows coincide.
        kwargs = {"bottom": mid}
        if header_top_idx == header_bot_idx:
            kwargs["top"] = top
        _set_cell_borders(cell, **kwargs)
    for cell in rows[body_last_idx].cells:
        kwargs = {"bottom": bot}
        # When the bottom-rule row is also the top-rule row (single-row
        # table), preserve the heavy top rule that was set above — each
        # _set_cell_borders call replaces the entire tcBorders element.
        if body_last_idx == header_top_idx:
            kwargs["top"] = top
        elif body_last_idx == header_bot_idx and header_top_idx == header_bot_idx:
            kwargs["top"] = top
            kwargs["bottom"] = mid
        _set_cell_borders(cell, **kwargs)


def style_word_table_typography(
    table: Any,
    *,
    header_rows: Iterable[int] = (0,),
    font_name: str = DEFAULT_FONT,
    header_pt: int = HEADER_PT,
    body_pt: int = BODY_PT,
    align_first_col: str = "left",
    align_data_cols: str = "center",
) -> None:
    """Apply book-tab typography to every cell of a DOCX table.

    Header cells are bold + centered, body cells use the configured
    alignment (default: row labels left, numeric columns centered).
    """
    from docx.shared import Pt
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    first_align = align_map[align_first_col]
    data_align = align_map[align_data_cols]
    header_set = set(header_rows)

    for r_idx, row in enumerate(table.rows):
        is_header = r_idx in header_set
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                if is_header:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para.alignment = first_align if c_idx == 0 else data_align
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(header_pt if is_header else body_pt)
                    if is_header:
                        run.font.bold = True


def add_word_notes_paragraph(
    doc: Any,
    text: str,
    *,
    font_name: str = DEFAULT_FONT,
    pt_size: int = NOTE_PT,
) -> None:
    """Append an italic small-font notes paragraph below the table."""
    from docx.shared import Pt

    if not text:
        return
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(pt_size)
    run.font.italic = True


def _word_cell_text(value: Any) -> str:
    """Return stable display text for a scalar Word table cell."""
    if value is None:
        return ""
    try:
        missing = pd.isna(value)
        if bool(missing):
            return ""
    except (TypeError, ValueError):
        pass
    return str(value)


def render_dataframe_to_word_table(
    doc: Any,
    df: Any,
    *,
    index_label: Optional[str] = "",
    notes: Optional[str | Sequence[str]] = None,
    header_pt: int = HEADER_PT,
    body_pt: int = BODY_PT,
    align_first_col: str = "left",
    align_data_cols: str = "center",
) -> Any:
    """Append a DataFrame as a Times/book-tab Word table.

    The caller owns document-level title/heading/page-break decisions; this
    helper centralizes the table body, typography, book-tab rules, missing
    value handling, and optional notes paragraph.
    """
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        WD_TABLE_ALIGNMENT = None  # type: ignore[assignment,misc]

    if index_label is None:
        header = list(df.columns)
    else:
        header = [index_label] + list(df.columns)
    if not header:
        header = [""]

    body_rows = max(len(df), 1)
    table = doc.add_table(rows=body_rows + 1, cols=len(header))
    if WD_TABLE_ALIGNMENT is not None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for j, col in enumerate(header):
        table.rows[0].cells[j].text = str(col)

    if len(df) == 0:
        for j in range(len(header)):
            table.rows[1].cells[j].text = ""
    else:
        for i, (idx, row_data) in enumerate(df.iterrows(), 1):
            values = list(row_data)
            if index_label is not None:
                values = [idx] + values
            for j, val in enumerate(values):
                table.rows[i].cells[j].text = _word_cell_text(val)

    style_word_table_typography(
        table,
        header_rows=(0,),
        header_pt=header_pt,
        body_pt=body_pt,
        align_first_col=align_first_col,
        align_data_cols=align_data_cols,
    )
    apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)

    if notes:
        if isinstance(notes, str):
            note_text = notes
        else:
            note_text = "\n".join(str(note) for note in notes if note)
        add_word_notes_paragraph(doc, note_text)

    return table


# ---------------------------------------------------------------------------
# XLSX helpers
# ---------------------------------------------------------------------------


def excel_booktab_borders() -> tuple[Any, Any, Any, Any]:
    """Return ``(top_rule, mid_rule, bottom_rule, no_border)`` borders."""
    from openpyxl.styles import Border, Side

    medium = Side(style="medium")
    thin = Side(style="thin")
    none = Side(style=None)
    return (
        Border(top=medium, bottom=none, left=none, right=none),
        Border(top=none, bottom=thin, left=none, right=none),
        Border(top=none, bottom=medium, left=none, right=none),
        Border(top=none, bottom=none, left=none, right=none),
    )
