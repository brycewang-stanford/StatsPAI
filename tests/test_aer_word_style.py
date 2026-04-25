"""
Tests for AER/QJE book-tab DOCX styling across regtable / sumstats / tab.

Verifies the three-rule structure (heavy top, thin mid, heavy bottom) and
the absence of internal vertical borders.
"""

import os

import numpy as np
import pandas as pd
import pytest

import statspai as sp

pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def two_models():
    rng = np.random.default_rng(7)
    n = 300
    x1 = rng.normal(0, 1, n)
    x2 = rng.normal(0, 1, n)
    y = 1.0 + 0.5 * x1 - 0.3 * x2 + rng.normal(0, 1, n)
    df = pd.DataFrame({"y": y, "x1": x1, "x2": x2})
    return sp.regress("y ~ x1", data=df), sp.regress("y ~ x1 + x2", data=df)


@pytest.fixture
def small_df():
    rng = np.random.default_rng(2)
    return pd.DataFrame({
        "y": rng.normal(size=80),
        "x": rng.normal(size=80),
        "z": rng.integers(0, 2, size=80),
    })


# ---------------------------------------------------------------------------
# Helpers — read border XML out of the .docx
# ---------------------------------------------------------------------------


def _cell_border_attrs(cell, edge):
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        return None
    elem = tcBorders.find(qn(f"w:{edge}"))
    if elem is None:
        return None
    return {
        "val": elem.get(qn("w:val")),
        "sz": elem.get(qn("w:sz")),
    }


def _every_cell(table):
    for row in table.rows:
        for cell in row.cells:
            yield cell


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_regtable_word_has_booktab_three_rules(tmp_path, two_models):
    m1, m2 = two_models
    out = tmp_path / "regtable.docx"
    sp.regtable(m1, m2).to_word(str(out))
    assert out.exists() and out.stat().st_size > 0

    doc = Document(str(out))
    table = doc.tables[0]
    assert len(table.rows) >= 3

    # Header top: heavy
    top = _cell_border_attrs(table.rows[0].cells[0], "top")
    assert top is not None
    assert top["val"] == "single"
    assert top["sz"] == "12"

    # Header bottom: thin
    mid = _cell_border_attrs(table.rows[0].cells[0], "bottom")
    assert mid is not None
    assert mid["val"] == "single"
    assert mid["sz"] == "4"

    # Last row bottom: heavy
    last = _cell_border_attrs(table.rows[-1].cells[0], "bottom")
    assert last is not None
    assert last["val"] == "single"
    assert last["sz"] == "12"


def test_regtable_word_has_no_inner_vertical_borders(tmp_path, two_models):
    m1, m2 = two_models
    out = tmp_path / "regtable_inner.docx"
    sp.regtable(m1, m2).to_word(str(out))

    doc = Document(str(out))
    table = doc.tables[0]
    # Mid-body rows should not have left/right borders set to single
    body_rows = table.rows[1:-1]
    bad = []
    for r_idx, row in enumerate(body_rows, start=1):
        for c_idx, cell in enumerate(row.cells):
            for edge in ("left", "right"):
                attrs = _cell_border_attrs(cell, edge)
                if attrs is None:
                    continue
                if attrs.get("val") and attrs["val"] != "nil":
                    bad.append((r_idx, c_idx, edge, attrs))
    assert not bad, f"Inner vertical borders should be nil; found: {bad}"


def test_regtable_word_notes_paragraph_is_italic(tmp_path, two_models):
    m1, m2 = two_models
    out = tmp_path / "regtable_notes.docx"
    sp.regtable(m1, m2, notes=["Country FE included"]).to_word(str(out))

    doc = Document(str(out))
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert paragraphs, "expected at least one notes paragraph"
    last = paragraphs[-1]
    assert any(run.italic for run in last.runs)


def test_balance_table_word_uses_booktab(tmp_path, small_df):
    out = tmp_path / "balance.docx"
    sp.mean_comparison(small_df, ["x", "y"], group="z").to_word(str(out))
    doc = Document(str(out))
    table = doc.tables[0]
    top = _cell_border_attrs(table.rows[0].cells[0], "top")
    bot = _cell_border_attrs(table.rows[-1].cells[0], "bottom")
    assert top and top["sz"] == "12"
    assert bot and bot["sz"] == "12"


def test_sumstats_word_uses_booktab(tmp_path, small_df):
    out = tmp_path / "sumstats.docx"
    sp.sumstats(small_df, vars=["x", "y"], output=str(out))
    assert out.exists()
    doc = Document(str(out))
    table = doc.tables[0]
    top = _cell_border_attrs(table.rows[0].cells[0], "top")
    bot = _cell_border_attrs(table.rows[-1].cells[0], "bottom")
    assert top and top["sz"] == "12"
    assert bot and bot["sz"] == "12"


def test_tab_word_uses_booktab(tmp_path, small_df):
    out = tmp_path / "tab.docx"
    sp.tab(small_df, row="z", output=str(out))
    assert out.exists()
    doc = Document(str(out))
    table = doc.tables[0]
    top = _cell_border_attrs(table.rows[0].cells[0], "top")
    bot = _cell_border_attrs(table.rows[-1].cells[0], "bottom")
    assert top and top["sz"] == "12"
    assert bot and bot["sz"] == "12"
