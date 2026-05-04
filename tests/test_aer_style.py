"""Tests for output/_aer_style.py — AER/QJE book-tab DOCX / XLSX helpers.

Covers
------
* ``_docx_border_xml`` builds correct OOXML border elements.
* ``apply_word_booktab_rules`` places heavy / thin / heavy rules correctly
  and handles edge cases (empty table, single row).
* ``style_word_table_typography`` sets font, size, bold, alignment per
  header / body role.
* ``add_word_notes_paragraph`` appends an italic notes paragraph.
* ``excel_booktab_borders`` returns the four Border instances.
"""

import sys
from io import BytesIO

import pytest

# ---------------------------------------------------------------------------
# DOCX tests  (python-docx is an optional dep — skip if missing)
# ---------------------------------------------------------------------------

pytest.importorskip("docx")

import docx  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches  # noqa: E402
from docx.shared import Pt  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

from statspai.output._aer_style import (  # noqa: E402
    _docx_border_xml,
    _set_cell_borders,
    _clear_table_borders,
    apply_word_booktab_rules,
    style_word_table_typography,
    add_word_notes_paragraph,
    excel_booktab_borders,
    TOP_RULE_SZ,
    MID_RULE_SZ,
    BOTTOM_RULE_SZ,
)


# ---- _docx_border_xml ---------------------------------------------------


def test_border_xml_defaults():
    """All edges omitted → each edge gets w:val="nil"."""
    elem = _docx_border_xml()
    assert elem.tag == qn("w:tcBorders")
    children = list(elem)
    assert len(children) == 6
    for child in children:
        assert child.get(qn("w:val")) == "nil"


def test_border_xml_top_set():
    """Passing top= supplies w:sz and w:val."""
    elem = _docx_border_xml(top={"sz": "12", "val": "single", "color": "000000"})
    top_elem = elem.find(qn("w:top"))
    assert top_elem is not None
    assert top_elem.get(qn("w:sz")) == "12"
    assert top_elem.get(qn("w:val")) == "single"
    # Others default to nil
    assert elem.find(qn("w:bottom")).get(qn("w:val")) == "nil"


def test_border_xml_all_set():
    """All six edges can be set simultaneously."""
    kwargs = {edge: {"sz": "8", "val": "single"} for edge in
              ("top", "bottom", "left", "right", "insideH", "insideV")}
    elem = _docx_border_xml(**kwargs)
    for edge in kwargs:
        assert elem.find(qn(f"w:{edge}")).get(qn("w:val")) == "single"


# ---- _set_cell_borders / _clear_table_borders ---------------------------


def _make_doc_table(num_rows=3, num_cols=2):
    """Helper: create a bare DOCX table."""
    d = docx.Document()
    table = d.add_table(rows=num_rows, cols=num_cols)
    return table


def test_set_cell_borders_replace():
    """Setting borders twice replaces — re-fetch to verify."""
    table = _make_doc_table()
    cell = table.rows[0].cells[0]
    _set_cell_borders(cell, top={"sz": "12"})
    # Verify first set
    borders = cell._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
    assert borders.find(qn("w:top")).get(qn("w:sz")) == "12"
    # Second call replaces the entire tcBorders element
    _set_cell_borders(cell, bottom={"sz": "4"})
    # Re-fetch (the old borders element was removed from the DOM)
    borders = cell._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
    assert borders.find(qn("w:top")).get(qn("w:val")) == "nil"
    assert borders.find(qn("w:bottom")).get(qn("w:sz")) == "4"


def test_clear_table_borders():
    """After clearing, every cell has nil borders on all six edges."""
    table = _make_doc_table(4, 3)
    # Pre-set a border on every cell
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, top={"sz": "12"})
    _clear_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                val = cell._tc.find(qn("w:tcPr")).find(
                    qn("w:tcBorders")).find(qn(f"w:{edge}")).get(qn("w:val"))
                assert val == "nil", f"{edge} is {val!r}"


# ---- apply_word_booktab_rules -------------------------------------------


class TestApplyWordBooktabRules:
    """Structural rule placement."""

    def test_empty_table(self):
        """Empty table → no error."""
        d = docx.Document()
        t = d.add_table(rows=0, cols=0)
        apply_word_booktab_rules(t)  # should not raise

    def test_single_row(self):
        """Single row: header_top=header_bot=body_last=row 0."""
        table = _make_doc_table(1, 2)
        apply_word_booktab_rules(table)
        cell = table.rows[0].cells[0]
        borders = cell._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        # Same row needs both top and bottom rule
        assert borders.find(qn("w:top")).get(qn("w:sz")) == TOP_RULE_SZ
        assert borders.find(qn("w:bottom")).get(qn("w:sz")) == BOTTOM_RULE_SZ

    def test_three_rows_standard(self):
        """3-row table: top on row 0 top; mid on row 0 bottom; bot on row 2."""
        table = _make_doc_table(3, 2)
        apply_word_booktab_rules(table)
        cells = {r: table.rows[r].cells[0] for r in range(3)}

        borders0 = cells[0]._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        assert borders0.find(qn("w:top")).get(qn("w:sz")) == TOP_RULE_SZ
        assert borders0.find(qn("w:bottom")).get(qn("w:sz")) == MID_RULE_SZ

        borders2 = cells[2]._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        assert borders2.find(qn("w:bottom")).get(qn("w:sz")) == BOTTOM_RULE_SZ

    def test_multi_row_header(self):
        """header_top=0, header_bot=1 → top on row 0, mid on row 1 bottom."""
        table = _make_doc_table(4, 2)
        apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=1)
        cells = {r: table.rows[r].cells[0] for r in range(4)}

        borders0 = cells[0]._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        assert borders0.find(qn("w:top")).get(qn("w:sz")) == TOP_RULE_SZ
        # Row 0 bottom should be nil (mid rule is on row 1)
        assert borders0.find(qn("w:bottom")).get(qn("w:val")) == "nil"

        borders1 = cells[1]._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        assert borders1.find(qn("w:bottom")).get(qn("w:sz")) == MID_RULE_SZ

    def test_custom_body_last(self):
        """body_last_idx overrides the final row."""
        table = _make_doc_table(5, 2)
        apply_word_booktab_rules(table, body_last_idx=2)
        cells = {r: table.rows[r].cells[0] for r in range(5)}

        borders2 = cells[2]._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        assert borders2.find(qn("w:bottom")).get(qn("w:sz")) == BOTTOM_RULE_SZ

        # Row 4 should NOT have a bottom rule
        borders4 = cells[4]._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        assert borders4.find(qn("w:bottom")).get(qn("w:val")) == "nil"

    def test_all_cells_in_row_get_rules(self):
        """Every cell in the header/mid/bottom row receives the border."""
        table = _make_doc_table(3, 5)  # 5 columns
        apply_word_booktab_rules(table)
        row0_cells = table.rows[0].cells
        for c in row0_cells:
            borders = c._tc.find(qn("w:tcPr")).find(qn("w:tcBorders"))
            assert borders.find(qn("w:top")).get(qn("w:sz")) == TOP_RULE_SZ


# ---- style_word_table_typography ----------------------------------------


class TestStyleWordTableTypography:
    def test_defaults(self):
        """Default: header bold+center, body left/center."""
        table = _make_doc_table(2, 3)
        # add some text so runs exist
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.text = f"r{r_idx}c{c_idx}"
        style_word_table_typography(table)

        # Header row 0
        run0 = table.rows[0].cells[0].paragraphs[0].runs[0]
        assert run0.font.bold
        assert run0.font.size == Pt(10)
        # Body row 1
        run1 = table.rows[1].cells[0].paragraphs[0].runs[0]
        assert not run1.font.bold

    def test_custom_font_and_sizes(self):
        """Custom font and pt sizes propagate correctly."""
        table = _make_doc_table(2, 2)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "x"
        style_word_table_typography(
            table, font_name="Arial", header_pt=12, body_pt=9,
        )
        run0 = table.rows[0].cells[0].paragraphs[0].runs[0]
        assert run0.font.name == "Arial"
        assert run0.font.size == Pt(12)
        run1 = table.rows[1].cells[0].paragraphs[0].runs[0]
        assert run1.font.size == Pt(9)

    def test_custom_header_rows(self):
        """Multiple header rows all get bold + centered."""
        table = _make_doc_table(4, 2)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "x"
        style_word_table_typography(table, header_rows=(0, 1))
        for idx in (0, 1):
            run = table.rows[idx].cells[0].paragraphs[0].runs[0]
            assert run.font.bold, f"row {idx} should be bold"
        run2 = table.rows[2].cells[0].paragraphs[0].runs[0]
        assert not run2.font.bold

    def test_alignment_options(self):
        """Left/center/right alignment works for first and data cols."""
        table = _make_doc_table(2, 3)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "x"
        style_word_table_typography(
            table,
            align_first_col="right",
            align_data_cols="left",
        )
        body_row = table.rows[1]
        # first column right-aligned
        para0 = body_row.cells[0].paragraphs[0]
        assert para0.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        # data columns left-aligned
        para1 = body_row.cells[1].paragraphs[0]
        assert para1.alignment == WD_ALIGN_PARAGRAPH.LEFT


# ---- add_word_notes_paragraph -------------------------------------------


def test_notes_paragraph():
    """Adds an italic notes paragraph at the end of the document."""
    d = docx.Document()
    add_word_notes_paragraph(d, "Notes: standard errors in parentheses.")
    assert len(d.paragraphs) >= 1
    para = d.paragraphs[-1]
    assert "Notes:" in para.text
    for run in para.runs:
        assert run.font.italic


def test_notes_empty_text():
    """Empty text → no paragraph added."""
    d = docx.Document()
    n_before = len(d.paragraphs)
    add_word_notes_paragraph(d, "")
    assert len(d.paragraphs) == n_before


def test_notes_custom_font_size():
    """Custom pt_size overrides default."""
    d = docx.Document()
    add_word_notes_paragraph(d, "note", pt_size=6)
    run = d.paragraphs[-1].runs[0]
    assert run.font.size == Pt(6)


# ---- XLSX helpers --------------------------------------------------------


class TestExcelBooktabBorders:
    def test_four_borders_returned(self):
        top, mid, bot, none = excel_booktab_borders()
        from openpyxl.styles import Border, Side
        assert isinstance(top, Border)
        assert isinstance(mid, Border)
        assert isinstance(bot, Border)
        assert isinstance(none, Border)
        # top has medium top
        assert top.top.style == "medium"
        # mid has thin bottom
        assert mid.bottom.style == "thin"
        # bot has medium bottom
        assert bot.bottom.style == "medium"
        # none has nil everywhere
        assert none.top.style is None

    def test_borders_round_trip(self):
        """Apply borders to real cells and verify via openpyxl."""
        from openpyxl import Workbook
        from openpyxl.styles import Border
        wb = Workbook()
        ws = wb.active
        top, mid, bot, none = excel_booktab_borders()
        ws["A1"].border = top
        ws["B1"].border = mid
        ws["C1"].border = bot
        ws["D1"].border = none
        assert ws["A1"].border.top.style == "medium"
        assert ws["B1"].border.bottom.style == "thin"
