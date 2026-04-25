"""
Tests for ``PaperTables.to_docx`` / ``PaperTables.to_xlsx``.

Verifies that a multi-panel paper-table bundle round-trips to .docx /
.xlsx with AER/QJE book-tab styling and that ``paper_tables(...,
docx_filename=..., xlsx_filename=...)`` writes both files.
"""

import numpy as np
import pandas as pd
import pytest

import statspai as sp


@pytest.fixture
def two_panels():
    rng = np.random.default_rng(11)
    n = 400
    d = rng.integers(0, 2, n)
    x = rng.normal(size=n)
    y = 1.0 + 0.6 * d + 0.3 * x + rng.normal(0, 1, n)
    df = pd.DataFrame({"y": y, "d": d, "x": x})
    m1 = sp.regress("y ~ d", data=df)
    m2 = sp.regress("y ~ d + x", data=df)
    m_male = sp.regress("y ~ d + x", data=df.iloc[:200])
    m_female = sp.regress("y ~ d + x", data=df.iloc[200:])
    return [m1, m2], [m_male, m_female]


def test_paper_tables_to_docx_writes_file(tmp_path, two_panels):
    main, het = two_panels
    pt = sp.paper_tables(main=main, heterogeneity=het, template="aer")
    out = tmp_path / "tables.docx"
    pt.to_docx(str(out))
    assert out.exists() and out.stat().st_size > 1000  # non-trivial size

    docx = pytest.importorskip("docx")
    from docx import Document

    doc = Document(str(out))
    assert len(doc.tables) == 2  # main + heterogeneity


def test_paper_tables_to_docx_uses_booktab_borders(tmp_path, two_panels):
    docx = pytest.importorskip("docx")
    from docx import Document
    from docx.oxml.ns import qn

    main, het = two_panels
    out = tmp_path / "booktab.docx"
    sp.paper_tables(main=main, heterogeneity=het).to_docx(str(out))

    doc = Document(str(out))
    for table in doc.tables:
        # First-row top edge: heavy
        tcPr = table.rows[0].cells[0]._tc.find(qn("w:tcPr"))
        assert tcPr is not None
        tcBorders = tcPr.find(qn("w:tcBorders"))
        assert tcBorders is not None
        top = tcBorders.find(qn("w:top"))
        assert top is not None
        assert top.get(qn("w:sz")) == "12"


def test_paper_tables_to_xlsx_writes_one_sheet_per_panel(tmp_path, two_panels):
    openpyxl = pytest.importorskip("openpyxl")
    main, het = two_panels
    out = tmp_path / "tables.xlsx"
    sp.paper_tables(main=main, heterogeneity=het).to_xlsx(str(out))
    assert out.exists()

    wb = openpyxl.load_workbook(str(out))
    assert set(wb.sheetnames) >= {"main", "heterogeneity"}


def test_paper_tables_xlsx_has_booktab_borders(tmp_path, two_panels):
    openpyxl = pytest.importorskip("openpyxl")
    main, _ = two_panels
    out = tmp_path / "borders.xlsx"
    sp.paper_tables(main=main).to_xlsx(str(out))

    wb = openpyxl.load_workbook(str(out))
    ws = wb["main"]
    # Locate the header row (first row with bold "(1)" or similar in col B)
    header_row = None
    for r in range(1, 6):
        v = ws.cell(row=r, column=2).value
        if v and str(v).startswith("("):
            header_row = r
            break
    assert header_row is not None, "expected header row with column labels '(1)' etc."
    # Header top should be a medium border
    top_side = ws.cell(row=header_row, column=2).border.top
    assert top_side.style == "medium"


def test_paper_tables_kwargs_write_both_files(tmp_path, two_panels):
    main, _ = two_panels
    docx_path = tmp_path / "kw.docx"
    xlsx_path = tmp_path / "kw.xlsx"
    pt = sp.paper_tables(
        main=main,
        docx_filename=str(docx_path),
        xlsx_filename=str(xlsx_path),
    )
    assert docx_path.exists()
    assert xlsx_path.exists()
    assert pt.main is not None
