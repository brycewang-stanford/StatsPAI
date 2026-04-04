"""
Tests for Excel and Word export functionality.

Tests outreg2 and modelsummary export to .xlsx and .docx formats.
"""

import pytest
import os
import tempfile
import numpy as np
import pandas as pd
from pathlib import Path

from statspai import regress, outreg2, modelsummary
from statspai.output.outreg2 import OutReg2


@pytest.fixture
def sample_results():
    """Two OLS regression results for testing."""
    np.random.seed(42)
    n = 100
    x1 = np.random.normal(0, 1, n)
    x2 = np.random.normal(0, 1, n)
    y = 1 + 2 * x1 + 3 * x2 + np.random.normal(0, 1, n)
    df = pd.DataFrame({'y': y, 'x1': x1, 'x2': x2})

    r1 = regress("y ~ x1", data=df)
    r2 = regress("y ~ x1 + x2", data=df)
    return r1, r2


@pytest.fixture
def tmp_dir():
    """Temporary directory for output files."""
    with tempfile.TemporaryDirectory() as d:
        yield d


# ================================================================
# outreg2 Excel export
# ================================================================

class TestOutreg2Excel:

    def test_excel_export(self, sample_results, tmp_dir):
        """Basic Excel export should produce a file."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "test.xlsx")
        outreg2(r1, r2, filename=path)
        assert os.path.exists(path)

    def test_excel_file_size(self, sample_results, tmp_dir):
        """Excel file should have reasonable size."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "test.xlsx")
        outreg2(r1, r2, filename=path)
        assert os.path.getsize(path) > 1000  # at least 1KB


# ================================================================
# outreg2 Word export
# ================================================================

class TestOutreg2Word:

    def test_word_export_via_class(self, sample_results, tmp_dir):
        """OutReg2 class .to_word() should produce a .docx file."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "test.docx")

        reg = OutReg2()
        reg.add_model(r1, "Model 1")
        reg.add_model(r2, "Model 2")
        reg.set_title("Test Table")
        reg.add_note("* p<0.1, ** p<0.05, *** p<0.01")
        reg.to_word(path)

        assert os.path.exists(path)
        assert os.path.getsize(path) > 1000

    def test_word_export_via_function(self, sample_results, tmp_dir):
        """outreg2() convenience function with format='word'."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "test.docx")
        outreg2(r1, r2, filename=path, format='word')
        assert os.path.exists(path)

    def test_word_auto_detect(self, sample_results, tmp_dir):
        """Auto-detect .docx extension."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "test.docx")
        outreg2(r1, r2, filename=path, format='auto')
        assert os.path.exists(path)

    def test_word_content(self, sample_results, tmp_dir):
        """Word file should contain expected content."""
        from docx import Document

        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "test.docx")
        outreg2(r1, r2, filename=path, format='word',
                title="Main Results")

        doc = Document(path)

        # Check title exists
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Main Results' in text

        # Check table exists with expected structure
        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Header should have "Variables", "Model 1", "Model 2"
        header_text = [cell.text for cell in table.rows[0].cells]
        assert 'Variables' in header_text


# ================================================================
# modelsummary Excel/Word export
# ================================================================

class TestModelsummaryExcel:

    def test_excel_export(self, sample_results, tmp_dir):
        """modelsummary with output='.xlsx' should produce file."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "summary.xlsx")
        result = modelsummary(r1, r2, output=path)
        assert os.path.exists(path)
        assert 'exported' in result.lower()

    def test_excel_readable(self, sample_results, tmp_dir):
        """Exported Excel should be readable by pandas."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "summary.xlsx")
        modelsummary(r1, r2, output=path)

        df = pd.read_excel(path)
        assert len(df) > 0


class TestModelsummaryWord:

    def test_word_export(self, sample_results, tmp_dir):
        """modelsummary with output='.docx' should produce file."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "summary.docx")
        result = modelsummary(r1, r2, output=path, title="My Table")
        assert os.path.exists(path)
        assert 'exported' in result.lower()

    def test_word_with_options(self, sample_results, tmp_dir):
        """Word export with all options."""
        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "full.docx")
        modelsummary(
            r1, r2,
            model_names=['Baseline', 'Full'],
            output=path,
            title="Regression Results",
            notes=["Robust SE in parentheses"],
            stars=True,
        )
        assert os.path.exists(path)

    def test_word_content(self, sample_results, tmp_dir):
        """Word file should have table with correct dimensions."""
        from docx import Document

        r1, r2 = sample_results
        path = os.path.join(tmp_dir, "check.docx")
        modelsummary(r1, r2, output=path)

        doc = Document(path)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Should have 3 columns: var name, model 1, model 2
        assert len(table.rows[0].cells) == 3


if __name__ == "__main__":
    pytest.main([__file__, '-v'])
